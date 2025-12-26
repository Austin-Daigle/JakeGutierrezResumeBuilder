import json
import os
import re
import tkinter as tk
from dataclasses import dataclass
from tkinter import colorchooser, filedialog, messagebox, simpledialog, ttk
import tkinter.font as tkfont
from typing import Any, Dict, List, Optional, Tuple


DATA_FILE_NAME = "resume_data.json"
EXPORT_FILE_NAME = "resume.tex"
EXPORT_DOCX_FILE_NAME = "resume.docx"
EXPORT_PDF_FILE_NAME = "resume.pdf"


def apply_modern_theme(root: tk.Tk) -> None:
    style = ttk.Style(root)
    try:
        windowing = root.tk.call("tk", "windowingsystem")
    except Exception:
        windowing = ""

    preferred: List[str] = []
    if windowing == "win32":
        preferred = ["vista", "xpnative"]
    elif windowing == "aqua":
        preferred = ["aqua"]
    else:
        preferred = ["clam", "alt"]

    for theme in preferred:
        if theme in style.theme_names():
            style.theme_use(theme)
            break

    try:
        default_font = tkfont.nametofont("TkDefaultFont")
        base_family = default_font.cget("family")
        base_size = int(default_font.cget("size"))
        style.configure("TButton", padding=(10, 6))
        style.configure("TEntry", padding=(6, 4))
        style.configure("TCombobox", padding=(6, 4))
        style.configure("TLabelframe.Label", font=(base_family, base_size, "bold"))
        style.configure("Heading.Treeview", font=(base_family, base_size, "bold"))
        style.configure("Treeview", rowheight=max(24, base_size + 12))
    except Exception:
        pass


DEFAULT_DATA = {
    "header": {
        "name": "",
        "phone": "",
        "email": "",
        "linkedin": "",
        "linkedin_display": "",
        "github": "",
        "github_display": "",
    },
    "sections": [
        {"id": "education", "title": "Education", "kind": "education", "entries": []},
        {"id": "experience", "title": "Experience", "kind": "experience", "entries": []},
        {"id": "projects", "title": "Projects", "kind": "projects", "entries": []},
        {"id": "technical_skills", "title": "Technical Skills", "kind": "skills", "entries": []},
    ],
}


DEMO_DATA = {
    "header": {
        "name": "Jake Ryan",
        "phone": "123-456-7890",
        "email": "jake@su.edu",
        "linkedin": "https://linkedin.com/in/...",
        "linkedin_display": "linkedin.com/in/jake",
        "github": "https://github.com/...",
        "github_display": "github.com/jake",
    },
    "sections": [
        {
            "id": "education",
            "title": "Education",
            "kind": "education",
            "entries": [
                {
                    "school": "Southwestern University",
                    "location": "Georgetown, TX",
                    "degree": "Bachelor of Arts in Computer Science, Minor in Business",
                    "dates": "Aug. 2018 -- May 2021",
                },
                {
                    "school": "Blinn College",
                    "location": "Bryan, TX",
                    "degree": "Associate's in Liberal Arts",
                    "dates": "Aug. 2014 -- May 2018",
                },
            ],
        },
        {
            "id": "experience",
            "title": "Experience",
            "kind": "experience",
            "entries": [
                {
                    "role": "Undergraduate Research Assistant",
                    "dates": "June 2020 -- Present",
                    "org": "Texas A&M University",
                    "location": "College Station, TX",
                    "bullets": [
                        [{"text": "Developed a REST API using FastAPI and PostgreSQL to store data from learning management systems"}],
                        [{"text": "Developed a full-stack web application using Flask, React, PostgreSQL and Docker to analyze GitHub data"}],
                        [{"text": "Explored ways to visualize GitHub collaboration in a classroom setting"}],
                    ],
                },
                {
                    "role": "Information Technology Support Specialist",
                    "dates": "Sep. 2018 -- Present",
                    "org": "Southwestern University",
                    "location": "Georgetown, TX",
                    "bullets": [
                        [{"text": "Communicate with managers to set up campus computers used on campus"}],
                        [{"text": "Assess and troubleshoot computer problems brought by students, faculty and staff"}],
                        [{"text": "Maintain upkeep of computers, classroom equipment, and 200 printers across campus"}],
                    ],
                },
                {
                    "role": "Artificial Intelligence Research Assistant",
                    "dates": "May 2019 -- July 2019",
                    "org": "Southwestern University",
                    "location": "Georgetown, TX",
                    "bullets": [
                        [
                            {
                                "text": "Explored methods to generate video game dungeons based off of ",
                            },
                            {"text": "The Legend of Zelda", "i": True},
                        ],
                        [{"text": "Developed a game in Java to test the generated dungeons"}],
                        [{"text": "Contributed 50K+ lines of code to an established codebase via Git"}],
                        [{"text": "Conducted a human subject study to determine which video game dungeon generation technique is enjoyable"}],
                        [{"text": "Wrote an 8-page paper and gave multiple presentations on-campus"}],
                        [{"text": "Presented virtually to the World Conference on Computational Intelligence"}],
                    ],
                },
            ],
        },
        {
            "id": "projects",
            "title": "Projects",
            "kind": "projects",
            "entries": [
                {
                    "title": "Gitlytics",
                    "stack": "Python, Flask, React, PostgreSQL, Docker",
                    "dates": "June 2020 -- Present",
                    "bullets": [
                        [{"text": "Developed a full-stack web application using with Flask serving a REST API with React as the frontend"}],
                        [{"text": "Implemented GitHub OAuth to get data from user's repositories"}],
                        [{"text": "Visualized GitHub data to show collaboration"}],
                        [{"text": "Used Celery and Redis for asynchronous tasks"}],
                    ],
                },
                {
                    "title": "Simple Paintball",
                    "stack": "Spigot API, Java, Maven, TravisCI, Git",
                    "dates": "May 2018 -- May 2020",
                    "bullets": [
                        [{"text": "Developed a Minecraft server plugin to entertain kids during free time for a previous job"}],
                        [{"text": "Published plugin to websites gaining 2K+ downloads and an average 4.5/5-star review"}],
                        [{"text": "Implemented continuous delivery using TravisCI to build the plugin upon new a release"}],
                        [{"text": "Collaborated with Minecraft server administrators to suggest features and get feedback about the plugin"}],
                    ],
                },
            ],
        },
        {
            "id": "technical_skills",
            "title": "Technical Skills",
            "kind": "skills",
            "entries": [
                {
                    "label": "Languages",
                    "value": [
                        {"text": "Java, Python, C/C++, SQL (Postgres), JavaScript, HTML/CSS, R"}
                    ],
                },
                {
                    "label": "Frameworks",
                    "value": [
                        {"text": "React, Node.js, Flask, JUnit, WordPress, Material-UI, FastAPI"}
                    ],
                },
                {
                    "label": "Developer Tools",
                    "value": [
                        {
                            "text": "Git, Docker, TravisCI, Google Cloud Platform, VS Code, Visual Studio, PyCharm, IntelliJ, Eclipse"
                        }
                    ],
                },
                {
                    "label": "Libraries",
                    "value": [{"text": "pandas, NumPy, Matplotlib"}],
                },
            ],
        },
    ],
}


LATEX_PREAMBLE = r"""%-------------------------
% Resume in Latex
% Author : Jake Gutierrez
% Based off of: https://github.com/sb2nov/resume
% License : MIT
%------------------------

\documentclass[letterpaper,11pt]{article}

\usepackage{latexsym}
\usepackage[empty]{fullpage}
\usepackage{titlesec}
\usepackage{marvosym}
\usepackage[usenames,dvipsnames]{color}
\usepackage[usenames,dvipsnames]{xcolor}
\usepackage{verbatim}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{fancyhdr}
\usepackage[english]{babel}
\usepackage{tabularx}
\input{glyphtounicode}


%----------FONT OPTIONS----------
% sans-serif
% \usepackage[sfdefault]{FiraSans}
% \usepackage[sfdefault]{roboto}
% \usepackage[sfdefault]{noto-sans}
% \usepackage[default]{sourcesanspro}

% serif
% \usepackage{CormorantGaramond}
% \usepackage{charter}


\pagestyle{fancy}
\fancyhf{} % clear all header and footer fields
\fancyfoot{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

% Adjust margins
\addtolength{\oddsidemargin}{-0.5in}
\addtolength{\evensidemargin}{-0.5in}
\addtolength{\textwidth}{1in}
\addtolength{\topmargin}{-.5in}
\addtolength{\textheight}{1.0in}

\urlstyle{same}

\raggedbottom
\raggedright
\setlength{\tabcolsep}{0in}

% Sections formatting
\titleformat{\section}{
  \vspace{-4pt}\scshape\raggedright\large
}{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]

% Ensure that generate pdf is machine readable/ATS parsable
\pdfgentounicode=1

%-------------------------
% Custom commands
\newcommand{\resumeItem}[1]{
  \item\small{
    {#1 \vspace{-2pt}}
  }
}

\newcommand{\resumeSubheading}[4]{
  \vspace{-2pt}\item
    \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
      \textbf{#1} & #2 \\
      \textit{\small#3} & \textit{\small #4} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubSubheading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \textit{\small#1} & \textit{\small #2} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeProjectHeading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \small#1 & #2 \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubItem}[1]{\resumeItem{#1}\vspace{-4pt}}

\renewcommand\labelitemii{$\vcenter{\hbox{\tiny$\bullet$}}$}

\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
\newcommand{\resumeItemListStart}{\begin{itemize}}
\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}

%-------------------------------------------
%%%%%%  RESUME STARTS HERE  %%%%%%%%%%%%%%%%%%%%%%%%%%%%


\begin{document}
"""


LATEX_END = r"""
%-------------------------------------------
\end{document}
"""


def _workspace_path() -> str:
    return os.path.dirname(os.path.abspath(__file__))


def _data_file_path() -> str:
    return os.path.join(_workspace_path(), DATA_FILE_NAME)


def _export_file_path() -> str:
    return os.path.join(_workspace_path(), EXPORT_FILE_NAME)


def deep_copy(obj):
    return json.loads(json.dumps(obj))


_LATEX_ESCAPE_REPLACEMENTS = {
    "\\": r"\textbackslash{}",
    "{": r"\{",
    "}": r"\}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def latex_escape(s: str) -> str:
    return "".join(_LATEX_ESCAPE_REPLACEMENTS.get(ch, ch) for ch in s)


def _hex_to_rgb_floats(hex_color: str) -> Optional[Tuple[float, float, float]]:
    m = re.fullmatch(r"#?([0-9a-fA-F]{6})", hex_color or "")
    if not m:
        return None
    v = m.group(1)
    r = int(v[0:2], 16) / 255.0
    g = int(v[2:4], 16) / 255.0
    b = int(v[4:6], 16) / 255.0
    return (r, g, b)


def rich_segments_to_latex(segments) -> str:
    out_parts: List[str] = []
    for seg in segments or []:
        raw_text = seg.get("text", "")
        if raw_text == "":
            continue
        t = latex_escape(raw_text)
        if seg.get("i"):
            t = rf"\emph{{{t}}}"
        if seg.get("u"):
            t = rf"\underline{{{t}}}"
        if seg.get("b"):
            t = rf"\textbf{{{t}}}"

        fg = seg.get("fg")
        bg = seg.get("bg")

        fg_rgb = _hex_to_rgb_floats(fg) if fg else None
        bg_rgb = _hex_to_rgb_floats(bg) if bg else None

        if fg_rgb is not None:
            t = rf"\textcolor[rgb]{{{fg_rgb[0]:.3f},{fg_rgb[1]:.3f},{fg_rgb[2]:.3f}}}{{{t}}}"
        if bg_rgb is not None:
            t = rf"\colorbox[rgb]{{{bg_rgb[0]:.3f},{bg_rgb[1]:.3f},{bg_rgb[2]:.3f}}}{{{t}}}"

        size = seg.get("size")
        if size is not None:
            try:
                s_int = int(size)
            except Exception:
                s_int = 0
            if s_int > 0:
                baseline = max(int(round(s_int * 1.2)), s_int + 1)
                t = rf"{{\fontsize{{{s_int}}}{{{baseline}}}\selectfont {t}}}"

        out_parts.append(t)
    return "".join(out_parts)


def load_resume_data() -> dict:
    path = _data_file_path()
    if not os.path.exists(path):
        return deep_copy(DEFAULT_DATA)
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return deep_copy(DEFAULT_DATA)


def save_resume_data(data: dict) -> None:
    path = _data_file_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def generate_latex(data: dict) -> str:
    header = data.get("header", {})
    name = latex_escape(header.get("name", ""))
    phone = latex_escape(header.get("phone", ""))
    email = header.get("email", "")
    linkedin = header.get("linkedin", "")
    linkedin_display_override = header.get("linkedin_display", "")
    github = header.get("github", "")
    github_display_override = header.get("github_display", "")

    email_display = latex_escape(email)
    email_link = latex_escape(email)

    linkedin_display = latex_escape(
        (linkedin_display_override or linkedin.replace("https://", "").replace("http://", ""))
    )
    github_display = latex_escape(
        (github_display_override or github.replace("https://", "").replace("http://", ""))
    )

    parts: List[str] = [LATEX_PREAMBLE]

    contact_parts: List[str] = []
    if phone:
        contact_parts.append(phone)
    if email:
        contact_parts.append(rf"\href{{mailto:{email_link}}}{{\underline{{{email_display}}}}}")
    if linkedin:
        contact_parts.append(
            rf"\href{{{latex_escape(linkedin)}}}{{\underline{{{linkedin_display}}}}}"
        )
    if github:
        contact_parts.append(rf"\href{{{latex_escape(github)}}}{{\underline{{{github_display}}}}}")

    contact_line = " $|$ ".join(contact_parts)
    if contact_line:
        contact_line = "    \\small " + contact_line
    else:
        contact_line = "    \\small"

    parts.append(
        "\n".join(
            [
                "\\begin{center}",
                f"    \\textbf{{\\Huge \\scshape {name}}} \\\\ \\vspace{{1pt}}",
                contact_line,
                "\\end{center}",
                "",
            ]
        )
    )

    for section in data.get("sections", []):
        title = latex_escape(section.get("title", ""))
        kind = section.get("kind")
        entries = section.get("entries", [])

        parts.append(f"\\section{{{title}}}")

        if kind in ("education", "experience", "projects"):
            parts.append("  \\resumeSubHeadingListStart")

        if kind == "education":
            for e in entries:
                parts.append("    \\resumeSubheading")
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('school',''))}}}"
                    + rf"{{{latex_escape(e.get('location',''))}}}"
                )
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('degree',''))}}}"
                    + rf"{{{latex_escape(e.get('dates',''))}}}"
                )

                body = e.get("body", [])
                if body:
                    parts.append("      \\resumeItemListStart")
                    parts.append("        " + rf"\resumeItem{{{rich_segments_to_latex(body)}}}")
                    parts.append("      \\resumeItemListEnd")
            parts.append("  \\resumeSubHeadingListEnd")
            parts.append("")
            continue

        if kind == "experience":
            for e in entries:
                parts.append("    \\resumeSubheading")
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('role',''))}}}"
                    + rf"{{{latex_escape(e.get('dates',''))}}}"
                )
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('org',''))}}}"
                    + rf"{{{latex_escape(e.get('location',''))}}}"
                )

                bullets = e.get("bullets", [])
                if bullets:
                    parts.append("      \\resumeItemListStart")
                    for b in bullets:
                        parts.append(
                            "        " + rf"\resumeItem{{{rich_segments_to_latex(b)}}}"
                        )
                    parts.append("      \\resumeItemListEnd")
                    parts.append("")
            parts.append("  \\resumeSubHeadingListEnd")
            parts.append("")
            continue

        if kind == "projects":
            for e in entries:
                title_text = latex_escape(e.get("title", ""))
                stack_text = latex_escape(e.get("stack", ""))
                dates_text = latex_escape(e.get("dates", ""))
                left = rf"\textbf{{{title_text}}} $|$ \emph{{{stack_text}}}"
                parts.append("      \\resumeProjectHeading")
                parts.append("          " + rf"{{{left}}}{{{dates_text}}}")

                bullets = e.get("bullets", [])
                if bullets:
                    parts.append("          \\resumeItemListStart")
                    for b in bullets:
                        parts.append(
                            "            "
                            + rf"\resumeItem{{{rich_segments_to_latex(b)}}}"
                        )
                    parts.append("          \\resumeItemListEnd")
            parts.append("    \\resumeSubHeadingListEnd")
            parts.append("")
            continue

        if kind == "skills":
            parts.append(" \\begin{itemize}[leftmargin=0.15in, label={}]")
            parts.append("    \\small{\\item{")
            for idx, e in enumerate(entries):
                label = latex_escape(e.get("label", ""))
                value = rich_segments_to_latex(e.get("value", []))
                line = rf"     \textbf{{{label}}}{{: {value}}}"
                if idx != len(entries) - 1:
                    line += r" \\\\"
                parts.append(line)
            parts.append("    }}")
            parts.append(" \\end{itemize}")
            parts.append("")
            continue

        for e in entries:
            title_text = latex_escape(e.get("title", ""))
            body_segments = e.get("body", [])
            if title_text:
                parts.append(rf"\textbf{{{title_text}}}\\")
            if body_segments:
                parts.append(rich_segments_to_latex(body_segments) + r"\\")
            parts.append("")

    parts.append(LATEX_END)
    return "\n".join(parts)


def export_latex(data: dict) -> None:
    out = generate_latex(data)
    path = _export_file_path()
    with open(path, "w", encoding="utf-8") as f:
        f.write(out)


@dataclass
class StyleState:
    family: str
    size: int
    bold: bool
    italic: bool
    underline: bool
    fg: Optional[str]
    bg: Optional[str]


class RichTextEditor(ttk.Frame):
    def __init__(self, master, *, default_family: Optional[str] = None, default_size: Optional[int] = None):
        super().__init__(master)

        try:
            df = tkfont.nametofont("TkDefaultFont")
            sys_family = df.cget("family")
            sys_size = int(df.cget("size"))
        except Exception:
            sys_family = "Arial"
            sys_size = 10

        self.default_family = default_family or sys_family
        self.default_size = int(default_size or sys_size)
        self.font_cache: Dict[Tuple[Any, ...], tkfont.Font] = {}

        toolbar = ttk.Frame(self)
        toolbar.pack(fill=tk.X, padx=6, pady=(6, 0))

        self.btn_bold = ttk.Button(toolbar, text="B", width=3, command=self.toggle_bold)
        self.btn_italic = ttk.Button(toolbar, text="I", width=3, command=self.toggle_italic)
        self.btn_underline = ttk.Button(toolbar, text="U", width=3, command=self.toggle_underline)

        self.btn_bold.pack(side=tk.LEFT, padx=(0, 4))
        self.btn_italic.pack(side=tk.LEFT, padx=(0, 4))
        self.btn_underline.pack(side=tk.LEFT, padx=(0, 8))

        families = sorted(
            {
                self.default_family,
                "Arial",
                "Helvetica",
                "Times New Roman",
                "Courier New",
                "DejaVu Sans",
            }
        )
        self.font_family_var = tk.StringVar(value=self.default_family)
        self.font_family_combo = ttk.Combobox(
            toolbar,
            textvariable=self.font_family_var,
            values=families,
            width=18,
            state="readonly",
        )
        self.font_family_combo.pack(side=tk.LEFT, padx=(0, 6))
        self.font_family_combo.bind("<<ComboboxSelected>>", lambda _e: self.apply_font_family())

        self.font_size_var = tk.IntVar(value=self.default_size)
        self.font_size_spin = ttk.Spinbox(
            toolbar,
            from_=8,
            to=48,
            textvariable=self.font_size_var,
            width=5,
            command=self.apply_font_size,
        )
        self.font_size_spin.pack(side=tk.LEFT, padx=(0, 6))

        ttk.Button(toolbar, text="Text Color", command=self.choose_fg).pack(
            side=tk.LEFT, padx=(0, 6)
        )
        ttk.Button(toolbar, text="Highlight", command=self.choose_bg).pack(
            side=tk.LEFT, padx=(0, 6)
        )

        text_frame = ttk.Frame(self)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

        text_scroll = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
        text_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self._last_selection: Optional[Tuple[str, str]] = None
        self.text = tk.Text(
            text_frame,
            wrap="word",
            undo=True,
            yscrollcommand=text_scroll.set,
            exportselection=False,
        )
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scroll.config(command=self.text.yview)

        self.text.bind("<ButtonRelease-1>", lambda _e: self._capture_selection())
        self.text.bind("<B1-Motion>", lambda _e: self._capture_selection())
        self.text.bind("<KeyRelease>", lambda _e: self._capture_selection())

    def _mark_modified(self) -> None:
        try:
            self.text.edit_modified(True)
        except Exception:
            return
        try:
            self.text.event_generate("<<Modified>>")
        except Exception:
            pass

    def _capture_selection(self) -> None:
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
            self._last_selection = (start, end)
        except tk.TclError:
            self._last_selection = None

    def _selection_range(self) -> Optional[Tuple[str, str]]:
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
            self._last_selection = (start, end)
            return (start, end)
        except tk.TclError:
            return self._last_selection

    def _index_for_style_lookup(self) -> str:
        sel = self._selection_range()
        if sel:
            return sel[0]
        return self.text.index("insert")

    def _font_key_tag(self, st: StyleState) -> str:
        b = 1 if st.bold else 0
        i = 1 if st.italic else 0
        u = 1 if st.underline else 0
        return f"f:{st.family}|s:{st.size}|b:{b}|i:{i}|u:{u}"

    def _get_or_create_font(self, family: str, size: int, bold: bool, italic: bool, underline: bool) -> tkfont.Font:
        key = (family, size, bold, italic, underline)
        if key in self.font_cache:
            return self.font_cache[key]
        weight = "bold" if bold else "normal"
        slant = "italic" if italic else "roman"
        f = tkfont.Font(family=family, size=size, weight=weight, slant=slant, underline=underline)
        self.font_cache[key] = f
        return f

    def _style_at(self, index: str) -> StyleState:
        tags = set(self.text.tag_names(index))
        family = self.default_family
        size = self.default_size
        bold = False
        italic = False
        underline = False
        fg = None
        bg = None

        for t in tags:
            if t.startswith("f:"):
                try:
                    m = re.fullmatch(r"f:(.*)\|s:(\d+)\|b:(\d)\|i:(\d)\|u:(\d)", t)
                    if m:
                        family = m.group(1)
                        size = int(m.group(2))
                        bold = m.group(3) == "1"
                        italic = m.group(4) == "1"
                        underline = m.group(5) == "1"
                except Exception:
                    pass
            elif t.startswith("fg:"):
                fg = t[3:]
            elif t.startswith("bg:"):
                bg = t[3:]

        return StyleState(family=family, size=size, bold=bold, italic=italic, underline=underline, fg=fg, bg=bg)

    def _remove_tag_prefix_in_range(self, prefix: str, start: str, end: str) -> None:
        for tag in self.text.tag_names():
            if tag.startswith(prefix):
                self.text.tag_remove(tag, start, end)

    def _apply_style(self, *, family=None, size=None, bold=None, italic=None, underline=None, fg=None, bg=None) -> None:
        sel = self._selection_range()
        if not sel:
            return
        start, end = sel

        try:
            if self.text.compare(start, "==", end):
                return
        except Exception:
            pass

        want_font_change = any(v is not None for v in (family, size, bold, italic, underline))
        want_fg_change = fg is not None
        want_bg_change = bg is not None

        runs: List[Tuple[str, str, StyleState]] = []
        idx = start
        run_start = start
        cur_style = self._style_at(start)

        while True:
            try:
                if not self.text.compare(idx, "<", end):
                    break
            except Exception:
                break

            next_idx = self.text.index(f"{idx}+1c")
            try:
                if self.text.compare(next_idx, ">", end):
                    next_idx = end
            except Exception:
                next_idx = end

            if next_idx == end:
                break

            st_next = self._style_at(next_idx)
            if st_next != cur_style:
                runs.append((run_start, next_idx, cur_style))
                run_start = next_idx
                cur_style = st_next

            idx = next_idx

        runs.append((run_start, end, cur_style))

        for rstart, rend, base in runs:
            if want_font_change:
                st = StyleState(
                    family=family if family is not None else base.family,
                    size=size if size is not None else base.size,
                    bold=bold if bold is not None else base.bold,
                    italic=italic if italic is not None else base.italic,
                    underline=underline if underline is not None else base.underline,
                    fg=base.fg,
                    bg=base.bg,
                )
                self._remove_tag_prefix_in_range("f:", rstart, rend)
                font_tag = self._font_key_tag(st)
                fnt = self._get_or_create_font(st.family, st.size, st.bold, st.italic, st.underline)
                self.text.tag_configure(font_tag, font=fnt)
                self.text.tag_add(font_tag, rstart, rend)

            if want_fg_change:
                self._remove_tag_prefix_in_range("fg:", rstart, rend)
                fg_tag = f"fg:{fg}"
                self.text.tag_configure(fg_tag, foreground=fg)
                self.text.tag_add(fg_tag, rstart, rend)

            if want_bg_change:
                self._remove_tag_prefix_in_range("bg:", rstart, rend)
                bg_tag = f"bg:{bg}"
                self.text.tag_configure(bg_tag, background=bg)
                self.text.tag_add(bg_tag, rstart, rend)

        self._mark_modified()

    def toggle_bold(self):
        idx = self._index_for_style_lookup()
        st = self._style_at(idx)
        self._apply_style(bold=not st.bold)

    def toggle_italic(self):
        idx = self._index_for_style_lookup()
        st = self._style_at(idx)
        self._apply_style(italic=not st.italic)

    def toggle_underline(self):
        idx = self._index_for_style_lookup()
        st = self._style_at(idx)
        self._apply_style(underline=not st.underline)

    def apply_font_family(self):
        self._apply_style(family=self.font_family_var.get())

    def apply_font_size(self):
        try:
            s = int(self.font_size_var.get())
        except Exception:
            return
        self._apply_style(size=s)

    def choose_fg(self):
        c = colorchooser.askcolor(title="Choose text color")
        if not c or not c[1]:
            return
        if not self._selection_range():
            try:
                start = self.text.index("insert wordstart")
                end = self.text.index("insert wordend")
                if self.text.compare(start, "<", end):
                    self.text.tag_add("sel", start, end)
                    self._last_selection = (start, end)
            except Exception:
                pass
        self._apply_style(fg=c[1])

    def choose_bg(self):
        c = colorchooser.askcolor(title="Choose highlight color")
        if not c or not c[1]:
            return
        if not self._selection_range():
            try:
                start = self.text.index("insert wordstart")
                end = self.text.index("insert wordend")
                if self.text.compare(start, "<", end):
                    self.text.tag_add("sel", start, end)
                    self._last_selection = (start, end)
            except Exception:
                pass
        self._apply_style(bg=c[1])

    def set_segments(self, segments):
        self.text.delete("1.0", "end")
        self.text.insert("1.0", "")
        if not segments:
            return

        for seg in segments:
            txt = seg.get("text", "")
            if txt == "":
                continue
            start_index = self.text.index("end-1c")
            self.text.insert("end", txt)
            end_index = self.text.index("end-1c")

            st = StyleState(
                family=seg.get("font") or self.default_family,
                size=int(seg.get("size") or self.default_size),
                bold=bool(seg.get("b")),
                italic=bool(seg.get("i")),
                underline=bool(seg.get("u")),
                fg=seg.get("fg"),
                bg=seg.get("bg"),
            )

            font_tag = self._font_key_tag(st)
            f = self._get_or_create_font(st.family, st.size, st.bold, st.italic, st.underline)
            self.text.tag_configure(font_tag, font=f)
            self.text.tag_add(font_tag, start_index, end_index)

            if st.fg:
                fg_tag = f"fg:{st.fg}"
                self.text.tag_configure(fg_tag, foreground=st.fg)
                self.text.tag_add(fg_tag, start_index, end_index)

            if st.bg:
                bg_tag = f"bg:{st.bg}"
                self.text.tag_configure(bg_tag, background=st.bg)
                self.text.tag_add(bg_tag, start_index, end_index)

    def get_segments(self):
        text = self.text.get("1.0", "end-1c")
        if text == "":
            return []

        intervals_add: Dict[int, List[Tuple[str, str]]] = {}
        intervals_remove: Dict[int, List[Tuple[str, str]]] = {}

        def _add_event(offset: int, kind: str, value: str):
            intervals_add.setdefault(offset, []).append((kind, value))

        def _remove_event(offset: int, kind: str, value: str):
            intervals_remove.setdefault(offset, []).append((kind, value))

        def _offset_for_index(idx: str) -> int:
            try:
                c = self.text.count("1.0", idx, "chars")
                if c and c[0] is not None:
                    return int(c[0])
            except Exception:
                pass

            try:
                return len(self.text.get("1.0", idx))
            except Exception:
                return 0

        for tag in self.text.tag_names():
            if tag.startswith("f:") or tag.startswith("fg:") or tag.startswith("bg:"):
                ranges = self.text.tag_ranges(tag)
                for i in range(0, len(ranges), 2):
                    start = str(ranges[i])
                    end = str(ranges[i + 1])
                    start_off = _offset_for_index(start)
                    end_off = _offset_for_index(end)
                    if start_off == end_off:
                        continue
                    if tag.startswith("f:"):
                        _add_event(start_off, "font", tag)
                        _remove_event(end_off, "font", tag)
                    elif tag.startswith("fg:"):
                        _add_event(start_off, "fg", tag[3:])
                        _remove_event(end_off, "fg", tag[3:])
                    elif tag.startswith("bg:"):
                        _add_event(start_off, "bg", tag[3:])
                        _remove_event(end_off, "bg", tag[3:])

        active_font_tag = None
        active_fg = None
        active_bg = None

        def _style_from_font_tag(font_tag: Optional[str]):
            if not font_tag:
                return (self.default_family, self.default_size, False, False, False)
            m = re.fullmatch(r"f:(.*)\|s:(\d+)\|b:(\d)\|i:(\d)\|u:(\d)", font_tag)
            if not m:
                return (self.default_family, self.default_size, False, False, False)
            family = m.group(1)
            size = int(m.group(2))
            b = m.group(3) == "1"
            i = m.group(4) == "1"
            u = m.group(5) == "1"
            return (family, size, b, i, u)

        segments: List[Dict[str, Any]] = []
        cur = {
            "family": None,
            "size": None,
            "b": None,
            "i": None,
            "u": None,
            "fg": None,
            "bg": None,
        }
        cur_text_parts: List[str] = []

        def _flush():
            if not cur_text_parts:
                return
            seg = {"text": "".join(cur_text_parts)}
            if cur["family"] is not None:
                seg["font"] = cur["family"]
            if cur["size"] is not None:
                seg["size"] = cur["size"]
            if cur["b"]:
                seg["b"] = True
            if cur["i"]:
                seg["i"] = True
            if cur["u"]:
                seg["u"] = True
            if cur["fg"]:
                seg["fg"] = cur["fg"]
            if cur["bg"]:
                seg["bg"] = cur["bg"]
            segments.append(seg)
            cur_text_parts.clear()

        def _set_style_from_active():
            family, size, b, i, u = _style_from_font_tag(active_font_tag)
            cur["family"] = family
            cur["size"] = size
            cur["b"] = b
            cur["i"] = i
            cur["u"] = u
            cur["fg"] = active_fg
            cur["bg"] = active_bg

        _set_style_from_active()

        for pos in range(0, len(text)):
            if pos in intervals_add:
                for kind, value in intervals_add[pos]:
                    if kind == "font":
                        active_font_tag = value
                    elif kind == "fg":
                        active_fg = value
                    elif kind == "bg":
                        active_bg = value

            new_family, new_size, new_b, new_i, new_u = _style_from_font_tag(active_font_tag)
            new_style = {
                "family": new_family,
                "size": new_size,
                "b": new_b,
                "i": new_i,
                "u": new_u,
                "fg": active_fg,
                "bg": active_bg,
            }

            if (
                cur_text_parts
                and (
                    cur["family"] != new_style["family"]
                    or cur["size"] != new_style["size"]
                    or cur["b"] != new_style["b"]
                    or cur["i"] != new_style["i"]
                    or cur["u"] != new_style["u"]
                    or cur["fg"] != new_style["fg"]
                    or cur["bg"] != new_style["bg"]
                )
            ):
                _flush()
                cur.update(new_style)

            if not cur_text_parts:
                cur.update(new_style)

            cur_text_parts.append(text[pos])

            if (pos + 1) in intervals_remove:
                for kind, value in intervals_remove[pos + 1]:
                    if kind == "font" and active_font_tag == value:
                        active_font_tag = None
                    elif kind == "fg" and active_fg == value:
                        active_fg = None
                    elif kind == "bg" and active_bg == value:
                        active_bg = None

        _flush()

        compacted = []
        for seg in segments:
            if seg.get("text") == "":
                continue
            compacted.append(seg)
        return compacted


class BulletEditorDialog(tk.Toplevel):
    def __init__(self, master, title: str, bullets: list, *, editor_height=12):
        super().__init__(master)
        self.title(title)
        self.transient(master)
        self.grab_set()

        self.result = None
        self.bullets = deep_copy(bullets or [])

        self.active_index = -1
        self.dirty = False
        self._loading = False

        main = ttk.Frame(self)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left = ttk.Frame(main)
        left.pack(side=tk.LEFT, fill=tk.Y)

        list_frame = ttk.Frame(left)
        list_frame.pack(fill=tk.BOTH, expand=True)

        list_scroll_v = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
        list_scroll_v.pack(side=tk.RIGHT, fill=tk.Y)

        list_scroll_h = ttk.Scrollbar(list_frame, orient=tk.HORIZONTAL)
        list_scroll_h.pack(side=tk.BOTTOM, fill=tk.X)

        self.listbox = tk.Listbox(
            list_frame,
            height=editor_height,
            width=38,
            yscrollcommand=list_scroll_v.set,
            xscrollcommand=list_scroll_h.set,
        )
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_scroll_v.config(command=self.listbox.yview)
        list_scroll_h.config(command=self.listbox.xview)

        btns = ttk.Frame(left)
        btns.pack(fill=tk.X, pady=(6, 0))
        self.btn_add_bullet = ttk.Button(btns, text="Add", command=self.add_bullet)
        self.btn_add_bullet.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_delete_bullet = ttk.Button(btns, text="Delete", command=self.delete_bullet)
        self.btn_delete_bullet.pack(side=tk.LEFT)

        right = ttk.Frame(main)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))

        self.editor = RichTextEditor(right)
        self.editor.pack(fill=tk.BOTH, expand=True)

        try:
            self.editor.text.edit_modified(False)
        except Exception:
            pass
        self.editor.text.bind("<<Modified>>", self.on_editor_modified)

        action = ttk.Frame(self)
        action.pack(fill=tk.X, padx=10, pady=(0, 10))
        self.btn_save = ttk.Button(action, text="Save", command=self.save_active)
        self.btn_save.pack(side=tk.LEFT)
        ttk.Button(action, text="OK", command=self.on_ok).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(action, text="Cancel", command=self.on_cancel).pack(side=tk.RIGHT)

        self.listbox.bind("<<ListboxSelect>>", self.on_select)

        self.refresh_list()
        if self.bullets:
            self.listbox.selection_set(0)
            self.listbox.activate(0)
            self.on_select(None)
        else:
            self._show_no_selection()

        self._update_save_button()
        self._update_bullet_buttons()

        self.minsize(700, 420)
        self.protocol("WM_DELETE_WINDOW", self.on_cancel)

    def refresh_list(self):
        self.listbox.delete(0, tk.END)
        for i, b in enumerate(self.bullets):
            preview = rich_segments_to_latex(b)
            if preview.strip() == "":
                preview = "(empty)"
            self.listbox.insert(tk.END, f"Bullet {i + 1}: {preview}")
        self._update_bullet_buttons()

    def _set_editor_enabled(self, enabled: bool) -> None:
        def _apply(w):
            try:
                if isinstance(w, tk.Text):
                    w.configure(state=("normal" if enabled else "disabled"))
                elif isinstance(w, ttk.Combobox):
                    w.configure(state=("readonly" if enabled else "disabled"))
                elif isinstance(w, ttk.Spinbox):
                    w.configure(state=("normal" if enabled else "disabled"))
                elif isinstance(w, (ttk.Button, ttk.Entry, ttk.Scrollbar)):
                    w.configure(state=("normal" if enabled else "disabled"))
            except Exception:
                pass

            for c in w.winfo_children():
                _apply(c)

        _apply(self.editor)

    def _show_no_selection(self) -> None:
        self.active_index = -1
        self.dirty = False
        self._update_save_button()

        self._loading = True
        try:
            self._set_editor_enabled(True)
            self.editor.set_segments([])
            try:
                self.editor.text.edit_modified(False)
            except Exception:
                pass
            self._set_editor_enabled(False)
        finally:
            self._loading = False
        self._update_bullet_buttons()

    def _update_bullet_buttons(self) -> None:
        if hasattr(self, "btn_delete_bullet"):
            self.btn_delete_bullet.configure(state=("normal" if len(self.bullets) > 0 else "disabled"))

    def selected_index(self) -> int:
        try:
            return int(self.listbox.curselection()[0])
        except Exception:
            return -1

    def _update_save_button(self) -> None:
        if hasattr(self, "btn_save"):
            can_save = self.dirty and (0 <= int(self.active_index) < len(self.bullets))
            self.btn_save.configure(state=("normal" if can_save else "disabled"))

    def save_active(self):
        idx = int(self.active_index)
        if not (0 <= idx < len(self.bullets)):
            return
        self.bullets[idx] = self.editor.get_segments()
        self.dirty = False
        self._update_save_button()

        sel = self.selected_index()
        self.refresh_list()
        if 0 <= sel < len(self.bullets):
            self.listbox.selection_set(sel)
            self.listbox.activate(sel)
        else:
            self.listbox.selection_set(idx)
            self.listbox.activate(idx)

    def on_select(self, _event):
        idx = self.selected_index()
        if idx < 0:
            if not self.bullets:
                self._show_no_selection()
            return
        idx = max(0, min(idx, len(self.bullets) - 1))

        if self.dirty:
            self.save_active()

        self.active_index = idx
        self._loading = True
        try:
            self._set_editor_enabled(True)
            self.editor.set_segments(self.bullets[idx])
            try:
                self.editor.text.edit_modified(False)
            except Exception:
                pass
        finally:
            self._loading = False

        self.dirty = False
        self._update_save_button()

    def on_editor_modified(self, _event):
        if self._loading:
            try:
                self.editor.text.edit_modified(False)
            except Exception:
                pass
            return

        try:
            if not self.editor.text.edit_modified():
                return
            self.editor.text.edit_modified(False)
        except Exception:
            pass

        self.dirty = True
        self._update_save_button()

    def add_bullet(self):
        if self.dirty:
            self.save_active()
        self.bullets.append([{"text": ""}])
        self.refresh_list()
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(len(self.bullets) - 1)
        self.listbox.activate(len(self.bullets) - 1)
        self.on_select(None)

    def delete_bullet(self):
        if len(self.bullets) <= 0:
            return
        idx = self.selected_index()
        if self.dirty:
            self.save_active()
        if idx < 0:
            idx = int(self.active_index)
        if 0 <= idx < len(self.bullets):
            self.bullets.pop(idx)
        self.refresh_list()

        if not self.bullets:
            self.listbox.selection_clear(0, tk.END)
            self._show_no_selection()
            return

        new_idx = min(max(idx, 0), len(self.bullets) - 1)
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(new_idx)
        self.listbox.activate(new_idx)
        self.on_select(None)

    def on_ok(self):
        if self.dirty:
            self.save_active()
        self.result = self.bullets
        self.destroy()

    def on_cancel(self):
        self.result = None
        self.destroy()


class PreviewDialog(tk.Toplevel):
    def __init__(self, master, data: dict, latex_source: str):
        super().__init__(master)
        self.title("Preview")
        self.transient(master)

        self.data = data
        self.latex_source = latex_source
        self.font_cache: Dict[Tuple[Any, ...], tkfont.Font] = {}

        root = ttk.Frame(self)
        root.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        nb = ttk.Notebook(root)
        nb.pack(fill=tk.BOTH, expand=True)

        rendered_tab = ttk.Frame(nb)
        latex_tab = ttk.Frame(nb)

        nb.add(rendered_tab, text="Rendered")
        nb.add(latex_tab, text="LaTeX")

        rendered_frame = ttk.Frame(rendered_tab)
        rendered_frame.pack(fill=tk.BOTH, expand=True)

        rendered_scroll = ttk.Scrollbar(rendered_frame)
        rendered_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.rendered_text = tk.Text(
            rendered_frame,
            wrap="word",
            yscrollcommand=rendered_scroll.set,
            padx=16,
            pady=16,
            borderwidth=0,
            highlightthickness=0,
        )
        self.rendered_text.pack(fill=tk.BOTH, expand=True)
        rendered_scroll.config(command=self.rendered_text.yview)

        self._build_rendered_preview()

        latex_text_frame = ttk.Frame(latex_tab)
        latex_text_frame.pack(fill=tk.BOTH, expand=True)

        latex_scroll = ttk.Scrollbar(latex_text_frame)
        latex_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.latex_text = tk.Text(latex_text_frame, wrap="none", yscrollcommand=latex_scroll.set)
        self.latex_text.pack(fill=tk.BOTH, expand=True)
        latex_scroll.config(command=self.latex_text.yview)

        self.latex_text.insert("1.0", latex_source)
        self.latex_text.configure(state="disabled")

        actions = ttk.Frame(self)
        actions.pack(fill=tk.X, padx=10, pady=(0, 10))
        ttk.Button(actions, text="Copy LaTeX Code to Clipboard", command=self.copy_latex_to_clipboard).pack(
            side=tk.LEFT
        )
        ttk.Button(actions, text="Close", command=self.on_close).pack(side=tk.RIGHT)

        self.minsize(900, 600)
        self.protocol("WM_DELETE_WINDOW", self.on_close)

    def copy_latex_to_clipboard(self):
        try:
            txt = self.latex_text.get("1.0", "end-1c")
        except Exception:
            txt = self.latex_source

        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
            self.update_idletasks()
        except Exception as e:
            messagebox.showerror("Copy Failed", str(e))
            return

        messagebox.showinfo("Copied", "LaTeX copied to clipboard")

    def _get_font(self, family: str, size: int, bold: bool, italic: bool, underline: bool) -> tkfont.Font:
        key = (family, size, bold, italic, underline)
        if key in self.font_cache:
            return self.font_cache[key]
        weight = "bold" if bold else "normal"
        slant = "italic" if italic else "roman"
        f = tkfont.Font(family=family, size=size, weight=weight, slant=slant, underline=underline)
        self.font_cache[key] = f
        return f

    def _insert_segments(self, segments: List[Dict[str, Any]], *, base_family: str, base_size: int):
        for idx, seg in enumerate(segments or []):
            txt = seg.get("text", "")
            if txt == "":
                continue
            family = seg.get("font") or base_family
            size = int(seg.get("size") or base_size)
            bold = bool(seg.get("b"))
            italic = bool(seg.get("i"))
            underline = bool(seg.get("u"))
            fg = seg.get("fg")
            bg = seg.get("bg")

            tag = f"seg_{self.rendered_text.index('end')}__{idx}"
            f = self._get_font(family, size, bold, italic, underline)
            self.rendered_text.tag_configure(tag, font=f)
            if fg:
                self.rendered_text.tag_configure(tag, foreground=fg)
            if bg:
                self.rendered_text.tag_configure(tag, background=bg)

            start = self.rendered_text.index("end-1c")
            self.rendered_text.insert("end", txt)
            end = self.rendered_text.index("end-1c")
            self.rendered_text.tag_add(tag, start, end)

    def _build_rendered_preview(self):
        self.rendered_text.configure(state="normal")
        self.rendered_text.delete("1.0", "end")

        try:
            df = tkfont.nametofont("TkDefaultFont")
            base_family = df.cget("family")
            base_size = int(df.cget("size"))
        except Exception:
            base_family = "Arial"
            base_size = 10

        name_font = self._get_font(base_family, base_size + 10, True, False, False)
        section_font = self._get_font(base_family, base_size + 2, True, False, False)
        subhead_font = self._get_font(base_family, base_size, True, False, False)
        italic_font = self._get_font(base_family, base_size, False, True, False)

        self.rendered_text.tag_configure("center", justify="center")
        self.rendered_text.tag_configure("name", font=name_font, justify="center")
        self.rendered_text.tag_configure("contact", font=self._get_font(base_family, base_size, False, False, False), justify="center")
        self.rendered_text.tag_configure("section", font=section_font)
        self.rendered_text.tag_configure("subhead", font=subhead_font)
        self.rendered_text.tag_configure("italic", font=italic_font)
        self.rendered_text.tag_configure("spacer", spacing1=8, spacing3=8)

        header = self.data.get("header", {})
        name = header.get("name", "")
        phone = header.get("phone", "")
        email = header.get("email", "")
        linkedin_text = header.get("linkedin_display", "") or header.get("linkedin", "")
        github_text = header.get("github_display", "") or header.get("github", "")

        contact_parts: List[str] = []
        if phone:
            contact_parts.append(phone)
        if email:
            contact_parts.append(email)
        if linkedin_text:
            contact_parts.append(linkedin_text)
        if github_text:
            contact_parts.append(github_text)

        self.rendered_text.insert("end", (name + "\n"), ("name", "center"))
        if contact_parts:
            self.rendered_text.insert("end", (" | ".join(contact_parts) + "\n"), ("contact", "center"))
        self.rendered_text.insert("end", "\n")

        for section in self.data.get("sections", []):
            title = section.get("title", "")
            kind = section.get("kind")
            entries = section.get("entries", [])

            self.rendered_text.insert("end", title + "\n", ("section",))
            self.rendered_text.insert("end", "―" * 60 + "\n")

            if kind == "education":
                for e in entries:
                    school = e.get("school", "")
                    location = e.get("location", "")
                    degree = e.get("degree", "")
                    dates = e.get("dates", "")
                    body = e.get("body", [])

                    line1 = school
                    if location:
                        line1 += "    " + location
                    self.rendered_text.insert("end", line1 + "\n", ("subhead",))

                    line2 = degree
                    if dates:
                        line2 += "    " + dates
                    self.rendered_text.insert("end", line2 + "\n", ("italic",))

                    if body:
                        self.rendered_text.insert("end", "  • ")
                        self._insert_segments(body, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

            elif kind == "experience":
                for e in entries:
                    role = e.get("role", "")
                    dates = e.get("dates", "")
                    org = e.get("org", "")
                    location = e.get("location", "")

                    line1 = role
                    if dates:
                        line1 += "    " + dates
                    self.rendered_text.insert("end", line1 + "\n", ("subhead",))

                    line2 = org
                    if location:
                        line2 += "    " + location
                    self.rendered_text.insert("end", line2 + "\n", ("italic",))

                    for b in e.get("bullets", []) or []:
                        self.rendered_text.insert("end", "  • ")
                        self._insert_segments(b, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

            elif kind == "projects":
                for e in entries:
                    ptitle = e.get("title", "")
                    stack = e.get("stack", "")
                    dates = e.get("dates", "")

                    header_line = ptitle
                    if stack:
                        header_line += " | " + stack
                    if dates:
                        header_line += "    " + dates
                    self.rendered_text.insert("end", header_line + "\n", ("subhead",))

                    for b in e.get("bullets", []) or []:
                        self.rendered_text.insert("end", "  • ")
                        self._insert_segments(b, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

            elif kind == "skills":
                for e in entries:
                    label = e.get("label", "")
                    self.rendered_text.insert("end", f"{label}: ", ("subhead",))
                    self._insert_segments(e.get("value", []), base_family=base_family, base_size=base_size)
                    self.rendered_text.insert("end", "\n")
                self.rendered_text.insert("end", "\n")

            else:
                for e in entries:
                    etitle = e.get("title", "")
                    if etitle:
                        self.rendered_text.insert("end", etitle + "\n", ("subhead",))
                    body = e.get("body", [])
                    if body:
                        self._insert_segments(body, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

        self.rendered_text.configure(state="disabled")

    def on_close(self):
        self.destroy()


class EntryEditorDialog(tk.Toplevel):
    def __init__(self, master, section_kind: str, entry: Optional[dict]):
        super().__init__(master)
        self.section_kind = section_kind
        self.original = deep_copy(entry) if entry is not None else None
        self.entry = deep_copy(entry) if entry is not None else {}
        self.result = None

        self.title("Edit Entry")
        self.transient(master)
        self.grab_set()

        root = ttk.Frame(self)
        root.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.fields_frame = ttk.Frame(root)
        self.fields_frame.pack(fill=tk.X)

        self.field_vars: Dict[str, tk.StringVar] = {}

        def add_field(label: str, key: str):
            row = ttk.Frame(self.fields_frame)
            row.pack(fill=tk.X, pady=3)
            ttk.Label(row, text=label, width=18).pack(side=tk.LEFT)
            var = tk.StringVar(value=str(self.entry.get(key, "")))
            self.field_vars[key] = var
            ttk.Entry(row, textvariable=var).pack(side=tk.LEFT, fill=tk.X, expand=True)

        if section_kind == "education":
            add_field("School", "school")
            add_field("Location", "location")
            add_field("Degree", "degree")
            add_field("Dates", "dates")
        elif section_kind == "experience":
            add_field("Role", "role")
            add_field("Dates", "dates")
            add_field("Organization", "org")
            add_field("Location", "location")
        elif section_kind == "projects":
            add_field("Title", "title")
            add_field("Stack", "stack")
            add_field("Dates", "dates")
        elif section_kind == "skills":
            add_field("Label", "label")
        else:
            add_field("Title", "title")

        self.body_frame = ttk.Frame(root)
        self.body_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

        self.editor = None
        self.bullets_button = None

        if section_kind in ("experience", "projects"):
            self.bullets = deep_copy(self.entry.get("bullets", []))
            if not self.bullets:
                self.bullets = [[{"text": ""}]]

            self.bullets_button = ttk.Button(
                self.body_frame,
                text=f"Edit Bullets ({len(self.bullets)})",
                command=self.edit_bullets,
            )
            self.bullets_button.pack(anchor="w")

        elif section_kind == "skills":
            ttk.Label(self.body_frame, text="Value").pack(anchor="w")
            self.editor = RichTextEditor(self.body_frame)
            self.editor.pack(fill=tk.BOTH, expand=True)
            self.editor.set_segments(self.entry.get("value", []))

        else:
            ttk.Label(self.body_frame, text="Body").pack(anchor="w")
            self.editor = RichTextEditor(self.body_frame)
            self.editor.pack(fill=tk.BOTH, expand=True)
            self.editor.set_segments(self.entry.get("body", []))

        actions = ttk.Frame(self)
        actions.pack(fill=tk.X, padx=10, pady=(0, 10))
        ttk.Button(actions, text="OK", command=self.on_ok).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(actions, text="Cancel", command=self.on_cancel).pack(side=tk.RIGHT)

        self.minsize(720, 420)
        self.protocol("WM_DELETE_WINDOW", self.on_cancel)

    def edit_bullets(self):
        dlg = BulletEditorDialog(self, "Edit Bullets", self.bullets)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self.bullets = dlg.result
        if self.bullets_button is not None:
            self.bullets_button.configure(text=f"Edit Bullets ({len(self.bullets)})")

    def on_ok(self):
        out = deep_copy(self.entry)
        for k, v in self.field_vars.items():
            out[k] = v.get().strip()

        if self.section_kind in ("experience", "projects"):
            out["bullets"] = deep_copy(self.bullets)
        elif self.section_kind == "skills":
            out["value"] = self.editor.get_segments() if self.editor else []
        else:
            out["body"] = self.editor.get_segments() if self.editor else []

        self.result = out
        self.destroy()

    def on_cancel(self):
        self.result = None
        self.destroy()


class ResumeApp(ttk.Frame):
    def __init__(self, master):
        super().__init__(master)
        self.master = master
        self.data = deep_copy(DEFAULT_DATA)

        self.undo_stack: List[dict] = []
        self.redo_stack: List[dict] = []
        self._suspend_undo = False
        self._header_typing_in_progress = False
        self._header_commit_job = None

        self.pack(fill=tk.BOTH, expand=True)

        self._build_ui()
        self._bind_undo_redo_shortcuts()
        self._update_undo_redo_buttons()
        self.refresh_sections()

    def _normalize_loaded_state(self, loaded: Any) -> dict:
        if not isinstance(loaded, dict):
            raise ValueError("Invalid file format (expected a JSON object at the top level).")

        state = deep_copy(DEFAULT_DATA)

        def _norm_key(k: Any) -> str:
            try:
                s = str(k)
            except Exception:
                return ""
            s = s.strip().lower()
            s = s.replace("_", " ")
            s = re.sub(r"\s+", " ", s)
            return s

        header_key_map = {
            "name": "name",
            "phone": "phone",
            "email": "email",
            "linkedin": "linkedin",
            "linked in": "linkedin",
            "li": "linkedin",
            "li url": "linkedin",
            "li link": "linkedin",
            "linkedin url": "linkedin",
            "linkedin link": "linkedin",
            "li text": "linkedin_display",
            "linkedin text": "linkedin_display",
            "linkedin display": "linkedin_display",
            "linkedin display text": "linkedin_display",
            "linkedin_display": "linkedin_display",
            "github": "github",
            "git hub": "github",
            "gh": "github",
            "gh url": "github",
            "gh link": "github",
            "github url": "github",
            "github link": "github",
            "gh text": "github_display",
            "github text": "github_display",
            "github display": "github_display",
            "github display text": "github_display",
            "github_display": "github_display",
        }

        def _apply_header_values(src: Any) -> None:
            if not isinstance(src, dict):
                return
            for k, v in src.items():
                nk = _norm_key(k)
                dest = header_key_map.get(nk)
                if dest is None and isinstance(k, str) and k in state.get("header", {}):
                    dest = k
                if not dest:
                    continue
                state["header"][dest] = str(v) if v is not None else ""

        header = loaded.get("header")
        _apply_header_values(header)
        if isinstance(loaded.get("data"), dict):
            _apply_header_values(loaded.get("data").get("header"))
        _apply_header_values(loaded)

        sections_any: Any = loaded.get("sections")
        sections_list: List[Any] = []
        if isinstance(sections_any, list):
            sections_list = sections_any
        elif isinstance(sections_any, dict):
            sections_list = list(sections_any.values())

        if not sections_list and isinstance(loaded.get("data"), dict):
            data_wrap = loaded.get("data")
            sections_any = data_wrap.get("sections")
            if isinstance(sections_any, list):
                sections_list = sections_any
            elif isinstance(sections_any, dict):
                sections_list = list(sections_any.values())

        defaults_by_id = {
            s.get("id"): s
            for s in deep_copy(DEFAULT_DATA).get("sections", [])
            if isinstance(s, dict) and s.get("id")
        }

        if not sections_list:
            inferred_sections: List[dict] = []
            for sec_id, default_sec in defaults_by_id.items():
                if sec_id in loaded:
                    v = loaded.get(sec_id)
                    if isinstance(v, dict):
                        s = deep_copy(v)
                        if not isinstance(s.get("entries"), list):
                            s["entries"] = []
                        inferred_sections.append(s)
                    elif isinstance(v, list):
                        s = deep_copy(default_sec)
                        s["entries"] = v
                        inferred_sections.append(s)
            if inferred_sections:
                sections_list = inferred_sections

        def _coerce_segments(val: Any) -> List[dict]:
            if val is None:
                return []
            if isinstance(val, str):
                return [{"text": val}]
            if isinstance(val, dict) and "text" in val:
                return [deep_copy(val)]
            if isinstance(val, list):
                if all(isinstance(x, dict) and "text" in x for x in val):
                    return deep_copy(val)
                if all(isinstance(x, str) for x in val):
                    return [{"text": "\n".join(val)}]
            return []

        def _coerce_bullets(val: Any) -> List[List[dict]]:
            if val is None:
                return []
            if isinstance(val, str):
                return [[{"text": val}]]
            if isinstance(val, list):
                if not val:
                    return []
                if all(isinstance(b, list) for b in val):
                    out: List[List[dict]] = []
                    for b in val:
                        segs = _coerce_segments(b)
                        if segs:
                            out.append(segs)
                        else:
                            out.append([{"text": ""}])
                    return out
                if all(isinstance(b, str) for b in val):
                    return [[{"text": b}] for b in val]
                if all(isinstance(b, dict) for b in val):
                    segs = _coerce_segments(val)
                    return [segs] if segs else []
            return []

        if sections_list:
            normalized_sections: List[dict] = []
            for sec in sections_list:
                if not isinstance(sec, dict):
                    continue
                s = deep_copy(sec)
                sec_id = s.get("id")
                if sec_id in defaults_by_id:
                    default_sec = defaults_by_id[sec_id]
                    for k in ("title", "kind"):
                        if not s.get(k) and default_sec.get(k):
                            s[k] = default_sec.get(k)

                if not isinstance(s.get("entries"), list):
                    s["entries"] = []

                kind = s.get("kind")
                normalized_entries: List[dict] = []
                for e in s.get("entries", []) or []:
                    if not isinstance(e, dict):
                        continue
                    ee = deep_copy(e)
                    if kind == "education":
                        ee["body"] = _coerce_segments(ee.get("body"))
                    elif kind in ("experience", "projects"):
                        if "bullets" in ee:
                            ee["bullets"] = _coerce_bullets(ee.get("bullets"))
                        elif "body" in ee:
                            body_segs = _coerce_segments(ee.get("body"))
                            ee["bullets"] = [body_segs] if body_segs else []
                            ee.pop("body", None)
                        else:
                            ee["bullets"] = _coerce_bullets([])
                    elif kind == "skills":
                        ee["value"] = _coerce_segments(ee.get("value"))
                    else:
                        ee["body"] = _coerce_segments(ee.get("body"))

                    normalized_entries.append(ee)

                s["entries"] = normalized_entries
                normalized_sections.append(s)
            if normalized_sections:
                state["sections"] = normalized_sections

        return state

    def _build_ui(self):
        self.master.title("Resume Template Maker")
        self.master.minsize(980, 600)

        topbar = ttk.Frame(self)
        topbar.pack(fill=tk.X, padx=10, pady=10)

        topbar_row1 = ttk.Frame(topbar)
        topbar_row1.pack(fill=tk.X)

        topbar_row2 = ttk.Frame(topbar)
        topbar_row2.pack(fill=tk.X, pady=(6, 0))

        file_btn = ttk.Menubutton(topbar_row1, text="File")
        file_menu = tk.Menu(file_btn, tearoff=False)
        file_menu.add_command(label="Save project as .json...", command=self.save_as)
        file_menu.add_command(label="Load project from .json file...", command=self.load_from_file)
        file_menu.add_separator()
        file_menu.add_command(label="Export LaTeX as .tex...", command=self.export_as_tex)
        file_menu.add_command(label="Export as Word Document...", command=self.export_as_docx)
        file_menu.add_command(label="Export as PDF...", command=self.export_as_pdf)
        file_btn.configure(menu=file_menu)
        file_btn.pack(side=tk.LEFT, padx=(0, 6))

        edit_btn = ttk.Menubutton(topbar_row1, text="Edit")
        edit_menu = tk.Menu(edit_btn, tearoff=False)
        edit_menu.add_command(label="Undo", command=self.undo)
        edit_menu.add_command(label="Redo", command=self.redo)
        edit_menu.add_separator()
        edit_menu.add_command(label="Delete all", command=self.delete_all)
        edit_btn.configure(menu=edit_menu)
        edit_btn.pack(side=tk.LEFT, padx=(0, 6))

        self.btn_undo = ttk.Button(topbar_row2, text="Undo", command=self.undo)
        self.btn_undo.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_redo = ttk.Button(topbar_row2, text="Redo", command=self.redo)
        self.btn_redo.pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(topbar_row2, text="Delete all", command=self.delete_all).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(topbar_row2, text="Export LaTeX", command=self.export).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(topbar_row2, text="Preview", command=self.preview).pack(side=tk.LEFT, padx=(0, 6))

        ttk.Separator(topbar_row2, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=8)

        ttk.Button(topbar_row2, text="Demo mode", command=self.demo_mode).pack(
            side=tk.LEFT, padx=(0, 6)
        )

        outer = ttk.Panedwindow(self, orient=tk.HORIZONTAL)
        outer.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))

        left = ttk.Frame(outer)
        right = ttk.Frame(outer)
        outer.add(left, weight=1)
        outer.add(right, weight=3)

        header = ttk.Labelframe(right, text="Header")
        header.pack(fill=tk.X, pady=(0, 10))

        self.header_vars: Dict[str, tk.StringVar] = {}

        def add_header_field(row_parent, label, key):
            row = ttk.Frame(row_parent)
            row.pack(fill=tk.X, padx=8, pady=4)
            ttk.Label(row, text=label, width=10).pack(side=tk.LEFT)
            var = tk.StringVar(value=str(self.data.get("header", {}).get(key, "")))
            self.header_vars[key] = var
            ttk.Entry(row, textvariable=var).pack(side=tk.LEFT, fill=tk.X, expand=True)
            var.trace_add("write", lambda *_a: self._on_header_var_changed())

        add_header_field(header, "Name", "name")
        add_header_field(header, "Phone", "phone")
        add_header_field(header, "Email", "email")
        add_header_field(header, "LinkedIn", "linkedin")
        add_header_field(header, "LI Text", "linkedin_display")
        add_header_field(header, "GitHub", "github")
        add_header_field(header, "GH Text", "github_display")

        sections_frame = ttk.Labelframe(left, text="Sections")
        sections_frame.pack(fill=tk.BOTH, expand=True)

        sec_list_frame = ttk.Frame(sections_frame)
        sec_list_frame.pack(fill=tk.BOTH, expand=True, padx=8, pady=(8, 0))

        sec_scroll = ttk.Scrollbar(sec_list_frame, orient=tk.VERTICAL)
        sec_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.sections_tree = ttk.Treeview(
            sec_list_frame,
            columns=("title", "kind"),
            show="headings",
            selectmode="browse",
            height=12,
            yscrollcommand=sec_scroll.set,
        )
        self.sections_tree.heading("title", text="Section")
        self.sections_tree.heading("kind", text="Type")
        self.sections_tree.column("title", width=220, stretch=True)
        self.sections_tree.column("kind", width=90, stretch=False)
        self.sections_tree.pack(fill=tk.BOTH, expand=True)
        sec_scroll.config(command=self.sections_tree.yview)
        self.sections_tree.bind("<<TreeviewSelect>>", self.on_section_selected)

        sec_btns = ttk.Frame(sections_frame)
        sec_btns.pack(fill=tk.X, padx=8, pady=8)
        ttk.Button(sec_btns, text="Add", command=self.add_section).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(sec_btns, text="Rename", command=self.rename_section).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(sec_btns, text="Delete", command=self.delete_section).pack(side=tk.LEFT)

        entries_frame = ttk.Labelframe(right, text="Entries")
        entries_frame.pack(fill=tk.BOTH, expand=True)

        ent_list_frame = ttk.Frame(entries_frame)
        ent_list_frame.pack(fill=tk.BOTH, expand=True, padx=8, pady=(8, 0))

        ent_scroll = ttk.Scrollbar(ent_list_frame, orient=tk.VERTICAL)
        ent_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.entries_tree = ttk.Treeview(
            ent_list_frame,
            columns=("entry",),
            show="headings",
            selectmode="browse",
            height=14,
            yscrollcommand=ent_scroll.set,
        )
        self.entries_tree.heading("entry", text="Entry")
        self.entries_tree.column("entry", width=640, stretch=True)
        self.entries_tree.pack(fill=tk.BOTH, expand=True)
        ent_scroll.config(command=self.entries_tree.yview)

        self.entries_tree.bind("<Double-Button-1>", lambda _e: self.edit_entry())
        self.entries_tree.bind("<Return>", lambda _e: self.edit_entry())
        self.entries_tree.bind("<Delete>", lambda _e: self.delete_entry())
        self.entries_tree.bind("<<TreeviewSelect>>", lambda _e: self._update_entry_action_buttons())

        entry_btns = ttk.Frame(entries_frame)
        entry_btns.pack(fill=tk.X, padx=8, pady=8)

        self.btn_entry_add = ttk.Button(entry_btns, text="Add", command=self.add_entry)
        self.btn_entry_add.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_entry_edit = ttk.Button(entry_btns, text="Edit", command=self.edit_entry)
        self.btn_entry_edit.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_entry_delete = ttk.Button(entry_btns, text="Delete", command=self.delete_entry)
        self.btn_entry_delete.pack(side=tk.LEFT)

        self._update_entry_action_buttons()

    def _selected_section_index(self) -> int:
        try:
            sel = self.sections_tree.selection()
            if not sel:
                return -1
            return int(sel[0])
        except Exception:
            return -1

    def _selected_entry_index(self) -> int:
        try:
            sel = self.entries_tree.selection()
            if not sel:
                return -1
            return int(sel[0])
        except Exception:
            return -1

    def refresh_sections(self):
        prev_idx = -1
        try:
            sel = self.sections_tree.selection()
            if sel:
                prev_idx = int(sel[0])
        except Exception:
            prev_idx = -1

        for iid in self.sections_tree.get_children(""):
            self.sections_tree.delete(iid)

        sections = self.data.get("sections", [])
        for idx, sec in enumerate(sections):
            kind = sec.get("kind", "")
            title = sec.get("title", "")
            self.sections_tree.insert("", "end", iid=str(idx), values=(title, kind))

        if sections:
            target = -1
            if 0 <= prev_idx < len(sections):
                target = prev_idx
            else:
                for i, sec in enumerate(sections):
                    if sec.get("entries"):
                        target = i
                        break
            if target < 0:
                target = 0

            if not sections[target].get("entries"):
                for i, sec in enumerate(sections):
                    if sec.get("entries"):
                        target = i
                        break

            self.sections_tree.selection_set(str(target))
            self.on_section_selected(None)
        else:
            for iid in self.entries_tree.get_children(""):
                self.entries_tree.delete(iid)
            self._update_entry_action_buttons()

    def refresh_entries(self):
        for iid in self.entries_tree.get_children(""):
            self.entries_tree.delete(iid)
        sidx = self._selected_section_index()
        if sidx < 0:
            self._update_entry_action_buttons()
            return
        sec = self.data.get("sections", [])[sidx]
        kind = sec.get("kind")
        entries = sec.get("entries", [])

        for idx, e in enumerate(entries):
            self.entries_tree.insert(
                "", "end", iid=str(idx), values=(self._entry_summary(kind, e),)
            )

        if self.entries_tree.get_children(""):
            sel = self.entries_tree.selection()
            if not sel:
                self.entries_tree.selection_set("0")

        self._update_entry_action_buttons()

    def _update_entry_action_buttons(self) -> None:
        if not hasattr(self, "btn_entry_edit") or not hasattr(self, "btn_entry_delete"):
            return

        has_entries = len(self.entries_tree.get_children("")) > 0
        has_selection = self._selected_entry_index() >= 0

        self.btn_entry_edit.configure(state=("normal" if has_selection else "disabled"))
        self.btn_entry_delete.configure(state=("normal" if has_entries else "disabled"))

    def _entry_summary(self, kind: str, e: dict) -> str:
        if kind == "education":
            return f"{e.get('school','')} - {e.get('degree','')}"
        if kind == "experience":
            return f"{e.get('role','')} - {e.get('org','')}"
        if kind == "projects":
            return f"{e.get('title','')} - {e.get('stack','')}"
        if kind == "skills":
            return f"{e.get('label','')}"
        return e.get("title", "(entry)")

    def on_section_selected(self, _event):
        self.refresh_entries()

    def sync_header_from_ui(self):
        header = self.data.setdefault("header", {})
        for k, var in self.header_vars.items():
            header[k] = var.get().strip()

    def _push_undo_snapshot(self, snapshot: dict) -> None:
        self.undo_stack.append(deep_copy(snapshot))
        if len(self.undo_stack) > 100:
            self.undo_stack.pop(0)
        self.redo_stack.clear()
        self._update_undo_redo_buttons()

    def _checkpoint_before_change(self) -> None:
        self._commit_header_typing()
        self.sync_header_from_ui()
        self._push_undo_snapshot(self.data)

    def _clear_header_typing_state(self) -> None:
        if self._header_commit_job is not None:
            try:
                self.master.after_cancel(self._header_commit_job)
            except Exception:
                pass
        self._header_commit_job = None
        self._header_typing_in_progress = False

    def _commit_header_typing(self) -> None:
        if self._header_commit_job is not None:
            try:
                self.master.after_cancel(self._header_commit_job)
            except Exception:
                pass
            self._header_commit_job = None

        if self._header_typing_in_progress:
            self.sync_header_from_ui()
            self._header_typing_in_progress = False

    def _on_header_var_changed(self) -> None:
        if self._suspend_undo:
            return

        if not self._header_typing_in_progress:
            self._push_undo_snapshot(self.data)
            self._header_typing_in_progress = True

        if self._header_commit_job is not None:
            try:
                self.master.after_cancel(self._header_commit_job)
            except Exception:
                pass
            self._header_commit_job = None

        self._header_commit_job = self.master.after(600, self._commit_header_typing)

    def _apply_state(self, state: dict) -> None:
        self._clear_header_typing_state()
        self._suspend_undo = True
        try:
            self.data = deep_copy(state)
            for k, var in self.header_vars.items():
                var.set(str(self.data.get("header", {}).get(k, "")))
            self.refresh_sections()
            try:
                self.master.update_idletasks()
            except Exception:
                pass
        finally:
            self._suspend_undo = False

    def _update_undo_redo_buttons(self) -> None:
        if hasattr(self, "btn_undo"):
            self.btn_undo.configure(state=("normal" if self.undo_stack else "disabled"))
        if hasattr(self, "btn_redo"):
            self.btn_redo.configure(state=("normal" if self.redo_stack else "disabled"))

    def undo(self) -> None:
        self._commit_header_typing()
        self.sync_header_from_ui()
        if not self.undo_stack:
            return
        current = deep_copy(self.data)
        prev = self.undo_stack.pop()
        self.redo_stack.append(current)
        self._apply_state(prev)
        self._update_undo_redo_buttons()

    def redo(self) -> None:
        self._commit_header_typing()
        self.sync_header_from_ui()
        if not self.redo_stack:
            return
        current = deep_copy(self.data)
        nxt = self.redo_stack.pop()
        self.undo_stack.append(current)
        self._apply_state(nxt)
        self._update_undo_redo_buttons()

    def _handle_undo_shortcut(self, _event, which: str):
        w = self.master.focus_get()
        if w is not None:
            try:
                if w.winfo_toplevel() is not self.master:
                    return
            except Exception:
                return
            if isinstance(w, tk.Text):
                return

        if which == "undo":
            self.undo()
        else:
            self.redo()
        return "break"

    def _bind_undo_redo_shortcuts(self) -> None:
        self.master.bind_all(
            "<Control-z>", lambda e: self._handle_undo_shortcut(e, "undo"), add="+"
        )
        self.master.bind_all(
            "<Control-y>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )
        self.master.bind_all(
            "<Control-Shift-z>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )
        self.master.bind_all(
            "<Command-z>", lambda e: self._handle_undo_shortcut(e, "undo"), add="+"
        )
        self.master.bind_all(
            "<Command-y>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )
        self.master.bind_all(
            "<Command-Shift-z>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )

    def save(self):
        self.sync_header_from_ui()
        save_resume_data(self.data)
        messagebox.showinfo("Saved", f"Saved to {DATA_FILE_NAME}")

    def save_as(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Save JSON",
            defaultextension=".json",
            initialdir=_workspace_path(),
            initialfile=DATA_FILE_NAME,
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            messagebox.showerror("Save Failed", str(e))
            return
        messagebox.showinfo("Saved", f"Saved to {path}")

    def load(self):
        self._checkpoint_before_change()
        data = load_resume_data()
        self._apply_state(data)
        self._update_undo_redo_buttons()
        messagebox.showinfo("Loaded", f"Loaded from {DATA_FILE_NAME}")

    def load_from_file(self):
        path = filedialog.askopenfilename(
            parent=self,
            title="Load JSON",
            initialdir=_workspace_path(),
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except Exception as e:
            messagebox.showerror("Load Failed", str(e))
            return

        try:
            data = self._normalize_loaded_state(raw)
        except Exception as e:
            messagebox.showerror("Load Failed", str(e))
            return

        self._checkpoint_before_change()

        try:
            for iid in self.sections_tree.selection():
                self.sections_tree.selection_remove(iid)
        except Exception:
            pass

        self._apply_state(data)
        self._update_undo_redo_buttons()
        self.refresh_sections()
        try:
            self.master.update_idletasks()
        except Exception:
            pass

        try:
            total_entries = sum(len(sec.get("entries", []) or []) for sec in (self.data.get("sections", []) or []) if isinstance(sec, dict))
        except Exception:
            total_entries = 0

        h = self.data.get("header", {}) if isinstance(self.data.get("header"), dict) else {}
        name_preview = str(h.get("name", ""))
        email_preview = str(h.get("email", ""))

        try:
            ui_sections = len(self.sections_tree.get_children(""))
        except Exception:
            ui_sections = -1
        try:
            ui_entries = len(self.entries_tree.get_children(""))
        except Exception:
            ui_entries = -1
        try:
            ui_name = self.header_vars.get("name").get() if self.header_vars.get("name") else ""
            ui_email = self.header_vars.get("email").get() if self.header_vars.get("email") else ""
        except Exception:
            ui_name = ""
            ui_email = ""

        messagebox.showinfo(
            "Loaded",
            f"Loaded from {path}\n\nName: {name_preview}\nEmail: {email_preview}\nSections: {len(self.data.get('sections', []) or [])}\nEntries: {total_entries}\n\nUI Name: {ui_name}\nUI Email: {ui_email}\nUI Sections: {ui_sections}\nUI Entries: {ui_entries}",
        )

    def export(self):
        self.sync_header_from_ui()
        export_latex(self.data)
        messagebox.showinfo("Exported", f"Wrote {EXPORT_FILE_NAME}")

    def export_as_tex(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Export LaTeX",
            defaultextension=".tex",
            initialdir=_workspace_path(),
            initialfile=EXPORT_FILE_NAME,
            filetypes=[("TeX files", "*.tex"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            latex_source = generate_latex(self.data)
            with open(path, "w", encoding="utf-8") as f:
                f.write(latex_source)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))
            return

        messagebox.showinfo("Exported", f"Wrote {path}")

    def export_as_docx(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Export Word Document",
            defaultextension=".docx",
            initialdir=_workspace_path(),
            initialfile=EXPORT_DOCX_FILE_NAME,
            filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            from docx import Document
            from docx.shared import Pt
        except Exception:
            messagebox.showerror(
                "Export Failed",
                "Missing dependency: python-docx\n\nInstall it with: pip install python-docx",
            )
            return

        def segments_to_plain(segments: Any) -> str:
            if not isinstance(segments, list):
                return ""
            parts: List[str] = []
            for seg in segments:
                if isinstance(seg, dict):
                    t = seg.get("text", "")
                    if isinstance(t, str) and t:
                        parts.append(t)
            return "".join(parts)

        def add_rich_paragraph(doc, segments: Any, *, style: Optional[str] = None, bullet: bool = False):
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            if bullet:
                try:
                    p.style = doc.styles["List Bullet"]
                except Exception:
                    pass
            if not isinstance(segments, list):
                if isinstance(segments, str) and segments:
                    p.add_run(segments)
                return p
            for seg in segments:
                if not isinstance(seg, dict):
                    continue
                t = seg.get("text", "")
                if not isinstance(t, str) or t == "":
                    continue
                run = p.add_run(t)
                run.bold = bool(seg.get("b"))
                run.italic = bool(seg.get("i"))
                run.underline = bool(seg.get("u"))
                if seg.get("font"):
                    try:
                        run.font.name = str(seg.get("font"))
                    except Exception:
                        pass
                if seg.get("size"):
                    try:
                        run.font.size = Pt(int(seg.get("size")))
                    except Exception:
                        pass
            return p

        header = self.data.get("header", {}) if isinstance(self.data.get("header"), dict) else {}
        name = str(header.get("name", "")).strip()
        phone = str(header.get("phone", "")).strip()
        email = str(header.get("email", "")).strip()
        linkedin = str(header.get("linkedin", "")).strip()
        github = str(header.get("github", "")).strip()

        doc = Document()
        if name:
            doc.add_heading(name, level=0)

        contact_parts: List[str] = []
        if phone:
            contact_parts.append(phone)
        if email:
            contact_parts.append(email)
        if linkedin:
            contact_parts.append(linkedin)
        if github:
            contact_parts.append(github)
        if contact_parts:
            doc.add_paragraph(" | ".join(contact_parts))

        for section in self.data.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title", "")).strip()
            kind = section.get("kind")
            entries = section.get("entries", [])

            doc.add_heading(title, level=1)

            for e in entries:
                if not isinstance(e, dict):
                    continue

                if kind == "education":
                    school = str(e.get("school", "")).strip()
                    location = str(e.get("location", "")).strip()
                    degree = str(e.get("degree", "")).strip()
                    dates = str(e.get("dates", "")).strip()
                    line1 = school
                    if location:
                        line1 += "    " + location
                    doc.add_paragraph(line1)
                    line2 = degree
                    if dates:
                        line2 += "    " + dates
                    doc.add_paragraph(line2)
                    body = e.get("body", [])
                    body_text = segments_to_plain(body)
                    if body_text:
                        add_rich_paragraph(doc, body, bullet=True)

                elif kind in ("experience", "projects"):
                    if kind == "experience":
                        role = str(e.get("role", "")).strip()
                        org = str(e.get("org", "")).strip()
                        location = str(e.get("location", "")).strip()
                        dates = str(e.get("dates", "")).strip()
                        header_line = " - ".join([x for x in [role, org, location, dates] if x])
                    else:
                        proj_title = str(e.get("title", "")).strip()
                        stack = str(e.get("stack", "")).strip()
                        dates = str(e.get("dates", "")).strip()
                        header_line = " - ".join([x for x in [proj_title, stack, dates] if x])
                    if header_line:
                        doc.add_paragraph(header_line)

                    bullets = e.get("bullets", [])
                    if isinstance(bullets, list):
                        for b in bullets:
                            add_rich_paragraph(doc, b, bullet=True)

                elif kind == "skills":
                    label = str(e.get("label", "")).strip()
                    value = e.get("value", [])
                    value_text = segments_to_plain(value)
                    if label or value_text:
                        if value_text:
                            doc.add_paragraph(f"{label}: {value_text}" if label else value_text)
                        else:
                            doc.add_paragraph(label)

                else:
                    entry_title = str(e.get("title", "")).strip()
                    if entry_title:
                        doc.add_paragraph(entry_title)
                    body = e.get("body", [])
                    body_text = segments_to_plain(body)
                    if body_text:
                        add_rich_paragraph(doc, body)

        try:
            doc.save(path)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))
            return

        messagebox.showinfo("Exported", f"Wrote {path}")

    def export_as_pdf(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Export PDF",
            defaultextension=".pdf",
            initialdir=_workspace_path(),
            initialfile=EXPORT_PDF_FILE_NAME,
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            from xml.sax.saxutils import escape

            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase.pdfmetrics import stringWidth
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.platypus import (
                HRFlowable,
                Indenter,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )
        except Exception:
            messagebox.showerror(
                "Export Failed",
                "Missing dependency: reportlab\n\nInstall it with: pip install reportlab",
            )
            return

        def _sanitize_text(t: Any) -> str:
            s = "" if t is None else str(t)
            try:
                return s.encode("cp1252", errors="replace").decode("cp1252")
            except Exception:
                return s

        def _nbsp(s: str) -> str:
            return (s or "").replace(" ", "\u00A0")

        def _map_font_face(face: Any) -> str:
            f = ("" if face is None else str(face)).strip().lower()
            if not f:
                return "Times-Roman"
            if "courier" in f or "mono" in f:
                return "Courier"
            if "arial" in f or "helvetica" in f or "sans" in f:
                return "Helvetica"
            if "times" in f or "serif" in f or "georgia" in f:
                return "Times-Roman"
            return "Times-Roman"

        def _map_font_variant(base: str, *, bold: bool, italic: bool) -> str:
            b = bool(bold)
            i = bool(italic)
            if base.startswith("Times"):
                if b and i:
                    return "Times-BoldItalic"
                if b:
                    return "Times-Bold"
                if i:
                    return "Times-Italic"
                return "Times-Roman"
            if base.startswith("Helvetica"):
                if b and i:
                    return "Helvetica-BoldOblique"
                if b:
                    return "Helvetica-Bold"
                if i:
                    return "Helvetica-Oblique"
                return "Helvetica"
            if base.startswith("Courier"):
                if b and i:
                    return "Courier-BoldOblique"
                if b:
                    return "Courier-Bold"
                if i:
                    return "Courier-Oblique"
                return "Courier"
            return base

        def _normalize_color(c: Any) -> Optional[str]:
            s = ("" if c is None else str(c)).strip()
            if not s:
                return None
            if re.fullmatch(r"#?[0-9a-fA-F]{6}", s):
                return ("#" + s.lstrip("#")).lower()
            return s

        def _to_rl_color(c: Any):
            s = _normalize_color(c)
            if not s:
                return None
            try:
                return colors.toColor(s)
            except Exception:
                try:
                    if isinstance(s, str) and s.startswith("#") and len(s) == 7:
                        return colors.HexColor(s)
                except Exception:
                    pass
            return None

        def _common_bg(segments: Any) -> Optional[str]:
            if not isinstance(segments, list) or not segments:
                return None
            bg: Optional[str] = None
            saw = False
            for seg in segments:
                if not isinstance(seg, dict):
                    continue
                c = _normalize_color(seg.get("bg"))
                if c is None:
                    continue
                if not saw:
                    bg = c
                    saw = True
                elif c != bg:
                    return None
            return bg if saw else None

        def segments_to_markup(segments: Any) -> str:
            if segments is None:
                return ""
            if isinstance(segments, str):
                return escape(_sanitize_text(segments)).replace("\n", "<br/>")
            if not isinstance(segments, list):
                return ""
            out: List[str] = []
            for seg in segments:
                if isinstance(seg, dict):
                    raw = seg.get("text", "")
                    if not isinstance(raw, str) or raw == "":
                        continue
                    t = escape(_sanitize_text(raw)).replace("\n", "<br/>")

                    font_attrs: List[str] = []
                    face = seg.get("font")
                    if face:
                        is_b = bool(seg.get("b"))
                        is_i = bool(seg.get("i"))
                        mapped = _map_font_variant(_map_font_face(face), bold=is_b, italic=is_i)
                        font_attrs.append(f" face=\"{escape(mapped)}\"")
                        seg = dict(seg)
                        seg["b"] = False
                        seg["i"] = False
                    size = seg.get("size")
                    if size is not None:
                        try:
                            s_int = int(size)
                        except Exception:
                            s_int = 0
                        if s_int > 0:
                            font_attrs.append(f" size=\"{s_int}\"")
                    fg = _normalize_color(seg.get("fg"))
                    if fg:
                        font_attrs.append(f" color=\"{escape(fg)}\"")

                    bg = _normalize_color(seg.get("bg"))
                    if bg and _to_rl_color(bg) is not None:
                        font_attrs.append(f" backcolor=\"{escape(bg)}\"")

                    if seg.get("u"):
                        t = f"<u>{t}</u>"
                    if seg.get("i"):
                        t = f"<i>{t}</i>"
                    if seg.get("b"):
                        t = f"<b>{t}</b>"

                    if font_attrs:
                        t = f"<font{''.join(font_attrs)}>{t}</font>"
                    out.append(t)
                elif isinstance(seg, str) and seg:
                    out.append(escape(_sanitize_text(seg)).replace("\n", "<br/>"))
            return "".join(out)

        def segments_to_paragraph(segments: Any, style: ParagraphStyle, *, bullet_text: Optional[str] = None):
            markup = segments_to_markup(segments)
            if not markup:
                return None
            bg = _common_bg(segments)
            bg_color = _to_rl_color(bg) if bg else None
            st = style
            if bg and bg_color is not None:
                markup = f"<para backColor=\"{escape(str(bg))}\" borderPadding=\"1\">{markup}</para>"
            if bullet_text is None:
                return Paragraph(markup, st)
            return Paragraph(markup, st, bulletText=bullet_text)

        def _strip_bullet_prefix(segments: Any) -> Any:
            prefixes = ("- ", "– ", "— ", "• ", "* ")

            if isinstance(segments, str):
                s = segments.lstrip()
                for p in prefixes:
                    if s.startswith(p):
                        return s[len(p) :].lstrip()
                return s

            if not isinstance(segments, list):
                return segments

            out: List[Any] = []
            stripped = False
            for seg in segments:
                if not stripped and isinstance(seg, dict):
                    t = seg.get("text", "")
                    if isinstance(t, str) and t:
                        s = t.lstrip()
                        for p in prefixes:
                            if s.startswith(p):
                                s = s[len(p) :].lstrip()
                                break
                        new_seg = dict(seg)
                        new_seg["text"] = s
                        out.append(new_seg)
                        stripped = True
                        continue
                if not stripped and isinstance(seg, str) and seg:
                    s = seg.lstrip()
                    for p in prefixes:
                        if s.startswith(p):
                            s = s[len(p) :].lstrip()
                            break
                    out.append(s)
                    stripped = True
                    continue
                out.append(seg)
            return out

        def _display_url(url: str, display_override: str) -> str:
            if display_override:
                return display_override.strip()
            return (url or "").replace("https://", "").replace("http://", "").strip()

        styles = getSampleStyleSheet()
        base = ParagraphStyle(
            "ResumeBase",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=10.5,
            leading=12.2,
            spaceBefore=0,
            spaceAfter=0,
        )
        name_style = ParagraphStyle(
            "ResumeName",
            parent=base,
            fontName="Times-Bold",
            fontSize=22,
            leading=24,
            alignment=1,
            spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "ResumeContact",
            parent=base,
            fontSize=9.5,
            leading=11,
            alignment=1,
            spaceAfter=10,
        )
        section_style = ParagraphStyle(
            "ResumeSection",
            parent=base,
            fontName="Times-Roman",
            fontSize=11,
            leading=13,
            spaceBefore=6,
            spaceAfter=1,
        )
        entry_left_bold = ParagraphStyle(
            "EntryLeftBold",
            parent=base,
            fontName="Times-Bold",
        )
        entry_left_italic = ParagraphStyle(
            "EntryLeftItalic",
            parent=base,
            fontName="Times-Italic",
            fontSize=9.8,
            leading=11.5,
        )
        entry_right = ParagraphStyle(
            "EntryRight",
            parent=base,
            alignment=2,
        )
        entry_right_italic = ParagraphStyle(
            "EntryRightItalic",
            parent=entry_right,
            fontName="Times-Italic",
            fontSize=9.8,
            leading=11.5,
        )
        bullet_style = ParagraphStyle(
            "ResumeBullet",
            parent=base,
            leftIndent=18,
            bulletIndent=10,
            spaceBefore=1,
            spaceAfter=0,
        )
        bullet_style.bulletFontName = "Helvetica"
        bullet_style.bulletFontSize = 7

        doc = SimpleDocTemplate(
            path,
            pagesize=letter,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36,
        )

        def _two_col(left: Paragraph, right: Paragraph):
            right_plain = ""
            try:
                right_plain = right.getPlainText()
            except Exception:
                right_plain = str(getattr(right, "text", "") or "")

            try:
                right_font = str(getattr(right.style, "fontName", base.fontName))
                right_size = float(getattr(right.style, "fontSize", base.fontSize))
            except Exception:
                right_font = base.fontName
                right_size = base.fontSize

            w = stringWidth(_sanitize_text(right_plain), right_font, right_size) + 6
            right_w = max(doc.width * 0.18, min(doc.width * 0.42, w))
            left_w = max(1, doc.width - right_w)

            t = Table(
                [[left, right]],
                colWidths=[left_w, right_w],
                hAlign="LEFT",
            )
            t.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ALIGN", (0, 0), (0, 0), "LEFT"),
                        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ]
                )
            )
            return t

        def _add_section_header(story: List[Any], title: str) -> None:
            t = (title or "").strip()
            if not t:
                return
            story.append(Paragraph(escape(_sanitize_text(t)), section_style))
            story.append(
                HRFlowable(
                    width="100%",
                    thickness=1,
                    color=colors.black,
                    spaceBefore=0,
                    spaceAfter=3,
                )
            )

        def _add_bullets_to(dest: List[Any], bullets: Any) -> None:
            if not isinstance(bullets, list):
                return
            for b in bullets:
                p = segments_to_paragraph(_strip_bullet_prefix(b), bullet_style, bullet_text="\u2022")
                if p is not None:
                    dest.append(p)

        story: List[Any] = []

        header = self.data.get("header", {}) if isinstance(self.data.get("header"), dict) else {}
        name = str(header.get("name", "")).strip()
        phone = str(header.get("phone", "")).strip()
        email = str(header.get("email", "")).strip()
        linkedin = str(header.get("linkedin", "")).strip()
        linkedin_display = str(header.get("linkedin_display", "")).strip()
        github = str(header.get("github", "")).strip()
        github_display = str(header.get("github_display", "")).strip()

        if name:
            story.append(Paragraph(escape(_sanitize_text(name)), name_style))

        contact_parts: List[str] = []
        if phone:
            contact_parts.append(_sanitize_text(phone))
        if email:
            contact_parts.append(_sanitize_text(email))
        if linkedin:
            contact_parts.append(_sanitize_text(_display_url(linkedin, linkedin_display)))
        if github:
            contact_parts.append(_sanitize_text(_display_url(github, github_display)))
        if contact_parts:
            story.append(Paragraph(escape(" | ".join(contact_parts)), contact_style))
        else:
            story.append(Spacer(1, 10))

        for section in self.data.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title", "")).strip()
            kind = str(section.get("kind", "") or "").strip().lower()
            raw_entries = section.get("entries", []) or []
            entries = raw_entries if isinstance(raw_entries, list) else []

            if title:
                _add_section_header(story, title)

            indent_started = False
            if kind in ("education", "experience", "projects", "skills"):
                story.append(Indenter(left=10.8, right=0))
                indent_started = True

            for e in entries:
                if not isinstance(e, dict):
                    continue

                block: List[Any] = []

                if kind == "education":
                    school = str(e.get("school", "")).strip()
                    location = str(e.get("location", "")).strip()
                    degree = str(e.get("degree", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(school)), entry_left_bold),
                            Paragraph(escape(_sanitize_text(_nbsp(location))), entry_right),
                        )
                    )
                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(degree)), entry_left_italic),
                            Paragraph(escape(_sanitize_text(_nbsp(dates))), entry_right_italic),
                        )
                    )

                    p = segments_to_paragraph(
                        _strip_bullet_prefix(e.get("body", [])),
                        bullet_style,
                        bullet_text="\u2022",
                    )
                    if p is not None:
                        block.append(p)

                    block.append(Spacer(1, 3))

                elif kind == "experience":
                    role = str(e.get("role", "")).strip()
                    org = str(e.get("org", "")).strip()
                    location = str(e.get("location", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(role)), entry_left_bold),
                            Paragraph(escape(_sanitize_text(_nbsp(dates))), entry_right),
                        )
                    )
                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(org)), entry_left_italic),
                            Paragraph(escape(_sanitize_text(_nbsp(location))), entry_right_italic),
                        )
                    )
                    _add_bullets_to(block, e.get("bullets", []))
                    block.append(Spacer(1, 4))

                elif kind == "projects":
                    proj_title = str(e.get("title", "")).strip()
                    stack = str(e.get("stack", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    left = f"<b>{escape(_sanitize_text(proj_title))}</b>"
                    if stack:
                        left += f" | <i>{escape(_sanitize_text(stack))}</i>"

                    block.append(
                        _two_col(
                            Paragraph(left, base),
                            Paragraph(escape(_sanitize_text(_nbsp(dates))), entry_right),
                        )
                    )
                    _add_bullets_to(block, e.get("bullets", []))
                    block.append(Spacer(1, 4))

                elif kind == "skills":
                    break

                else:
                    entry_title = str(e.get("title", "")).strip()
                    if entry_title:
                        block.append(Paragraph(escape(_sanitize_text(entry_title)), entry_left_bold))
                    p = segments_to_paragraph(e.get("body", []), base)
                    if p is not None:
                        block.append(p)
                    block.append(Spacer(1, 4))

                if block:
                    story.extend(block)

            if kind == "skills":
                lines: List[str] = []
                for e in entries:
                    if not isinstance(e, dict):
                        continue
                    label = str(e.get("label", "")).strip()
                    value_markup = segments_to_markup(e.get("value", []))
                    label_txt = escape(_sanitize_text(label))
                    if label_txt and value_markup:
                        lines.append(f"<b>{label_txt}</b>: {value_markup}")
                    elif label_txt:
                        lines.append(f"<b>{label_txt}</b>:")
                    elif value_markup:
                        lines.append(value_markup)
                if lines:
                    story.append(Paragraph("<br/>".join(lines), base))

            if indent_started:
                story.append(Indenter(left=-10.8, right=0))

        try:
            doc.build(story)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))
            return

        messagebox.showinfo("Exported", f"Wrote {path}")

    def preview(self):
        self.sync_header_from_ui()
        latex_source = generate_latex(self.data)
        dlg = PreviewDialog(self.master, self.data, latex_source)
        dlg.grab_set()

    def delete_all(self):
        ok = messagebox.askyesno(
            "Delete all",
            "Delete all entered data and reset to a clean start?\n\nThis cannot be undone unless you use Undo.",
        )
        if not ok:
            return
        self._checkpoint_before_change()
        self._apply_state(DEFAULT_DATA)
        self._update_undo_redo_buttons()

    def demo_mode(self):
        ok = messagebox.askyesno(
            "Demo mode",
            "Do you want to load demo data?\n\nIf you select Yes, all current data in the program will be overwritten.",
        )
        if not ok:
            return
        self._checkpoint_before_change()
        self._apply_state(DEMO_DATA)
        self._update_undo_redo_buttons()

    def add_section(self):
        title = simpledialog.askstring("Add Section", "Section title:", parent=self)
        if not title:
            return
        kind = simpledialog.askstring(
            "Add Section",
            "Section kind (education, experience, projects, skills, custom):",
            parent=self,
        )
        if not kind:
            return
        kind = kind.strip().lower()
        self._checkpoint_before_change()
        sec = {"id": f"sec_{len(self.data.get('sections', [])) + 1}", "title": title.strip(), "kind": kind, "entries": []}
        self.data.setdefault("sections", []).append(sec)
        self.refresh_sections()
        self._update_undo_redo_buttons()

    def rename_section(self):
        idx = self._selected_section_index()
        if idx < 0:
            return
        sec = self.data.get("sections", [])[idx]
        new_title = simpledialog.askstring(
            "Rename Section", "New title:", initialvalue=sec.get("title", ""), parent=self
        )
        if not new_title:
            return
        self._checkpoint_before_change()
        sec["title"] = new_title.strip()
        self.refresh_sections()
        self.sections_tree.selection_set(str(idx))
        self.on_section_selected(None)
        self._update_undo_redo_buttons()

    def delete_section(self):
        idx = self._selected_section_index()
        if idx < 0:
            return
        sec = self.data.get("sections", [])[idx]
        ok = messagebox.askyesno("Delete Section", f"Delete section '{sec.get('title','')}'?")
        if not ok:
            return
        self._checkpoint_before_change()
        self.data.get("sections", []).pop(idx)
        self.refresh_sections()
        self._update_undo_redo_buttons()

    def add_entry(self):
        sidx = self._selected_section_index()
        if sidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        kind = sec.get("kind")
        new_entry = {}
        if kind == "education":
            new_entry = {"school": "", "location": "", "degree": "", "dates": ""}
        elif kind == "experience":
            new_entry = {"role": "", "dates": "", "org": "", "location": "", "bullets": [[{"text": ""}]]}
        elif kind == "projects":
            new_entry = {"title": "", "stack": "", "dates": "", "bullets": [[{"text": ""}]]}
        elif kind == "skills":
            new_entry = {"label": "", "value": [{"text": ""}]}
        else:
            new_entry = {"title": "", "body": [{"text": ""}]}

        dlg = EntryEditorDialog(self, kind, new_entry)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self._checkpoint_before_change()
        sec.setdefault("entries", []).append(dlg.result)
        self.refresh_entries()
        self._update_undo_redo_buttons()

    def edit_entry(self):
        sidx = self._selected_section_index()
        eidx = self._selected_entry_index()
        if sidx < 0 or eidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        kind = sec.get("kind")
        entries = sec.get("entries", [])
        if eidx >= len(entries):
            return

        dlg = EntryEditorDialog(self, kind, entries[eidx])
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self._checkpoint_before_change()
        entries[eidx] = dlg.result
        self.refresh_entries()
        self.entries_tree.selection_set(str(eidx))
        self._update_undo_redo_buttons()

    def delete_entry(self):
        sidx = self._selected_section_index()
        eidx = self._selected_entry_index()
        if sidx < 0 or eidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        entries = sec.get("entries", [])
        if eidx >= len(entries):
            return
        ok = messagebox.askyesno("Delete Entry", "Delete selected entry?")
        if not ok:
            return
        self._checkpoint_before_change()
        entries.pop(eidx)
        self.refresh_entries()
        self._update_undo_redo_buttons()


def main():
    root = tk.Tk()
    apply_modern_theme(root)
    app = ResumeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
