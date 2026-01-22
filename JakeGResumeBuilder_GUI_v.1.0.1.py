import json
import os
import re
import sys
import tkinter as tk
from dataclasses import dataclass
from tkinter import colorchooser, filedialog, messagebox, simpledialog, ttk
import tkinter.font as tkfont
from typing import Any, Callable, Dict, List, Optional, Tuple


DATA_FILE_NAME = "resume_data.json"
EXPORT_FILE_NAME = "resume.tex"
EXPORT_DOCX_FILE_NAME = "resume.docx"
EXPORT_PDF_FILE_NAME = "resume.pdf"

APP_NAME = "JakeGResumeBuilder"
APP_VERSION = "1.0.1"


QUICK_START_GUIDE_TEXT = """# Quick Start Guide

## 1) Create / load a project
- Use File -> Save project as .json... to create a project file.
- Use File -> Load project from .json file... to open an existing project.

## 2) Fill out the Header
- Enter your contact information in the Header panel.
- The header updates the preview/export.
- Use the Platform dropdowns to choose what a link represents (LinkedIn, GitHub, TikTok, etc.).
- Choose None to omit that link from Preview/Export (and the URL field will be disabled).
- Choose Custom to type your own platform label.

## 3) Work with Sections
- The left panel lists your Sections.
- Select a section to see its Entries.
- Drag sections to reorder them.
- Add / delete sections using the buttons next to the list.

## 4) Work with Entries
- Select a section, then add entries.
- Use the entry editor to fill in fields.
- Drag entries to reorder them.

## 5) Bullets (Experience / Projects)
- Open an entry, then add bullets.
- Drag bullets to reorder them.

## 6) Spellcheck
- Settings -> Enable Spellcheck turns spellcheck on/off.
- Right-click a misspelled word for suggestions.
- Choose Ignore All to ignore a word in this project.
- Settings -> Edit Document Spellcheck Data... lets you manage the ignore list.

## 7) Preview and Export
- Preview shows how the resume will render.
- Use File -> Export LaTeX / Word / PDF when ready.
- PDF export creates clickable links for email and header links.

## 8) Undo/Redo and Saving
- Use Edit -> Undo / Redo.
- Use Ctrl+S to save (or Save in the File menu).
"""


DETAILED_HELP_TEXT = """# Detailed Help

## Overview
This program builds a resume from structured data (Header + Sections + Entries + Bullets) and can export to multiple formats.
The key idea is: you edit data in the UI, the program keeps an internal JSON-like state, and Preview/Export render from that state.

## Projects (.json files)
- A "project" is a single JSON file that contains all resume data.
- Saving writes the current in-memory state to the project file.
- Loading replaces the current in-memory state with the file contents.

## Header
- The Header is your contact block (name, phone, email, links).
- The preview uses Header fields to build the top section of the resume.
- Header edits are committed automatically, but saving is what persists them to disk.

### Platform links
- Platform dropdowns let you choose what each link represents (LinkedIn/GitHub/Portfolio/etc.).
- None omits the link from Preview/Export and disables the corresponding URL field.
- Custom prompts for a label.
- The "Link Label" fields control the text shown for the link in Preview/Export.
- PDF export outputs clickable links.

## Sections
Sections are the major resume blocks (Education, Experience, Projects, Skills, etc.).

### Section order
- The section order determines the order in Preview and in exports.
- You can drag sections to reorder them.

### Section type (kind)
- The kind controls which fields show up in entry editors and how exports format the content.
- For example: Experience entries use a role/org/location/dates model and a bullet list.

## Entries
Entries are items inside a section.

### Entry order
- Entry order affects Preview/Export order.
- You can drag entries to reorder them.

### Editing entries
- Double-click or use the edit button to open the entry editor.
- Different section kinds have different entry fields.

## Bullets
Bullets are primarily used in Experience and Projects.

### Bullet editor
- Bullets are edited in a dedicated dialog with a list on the left and a rich text editor on the right.
- Drag bullets to reorder.
- If a bullet is empty, it will render as empty and is usually not useful in exports.

## Rich text vs single-line fields
- Some fields are single-line (they behave like normal entry fields).
- Bullets and some other areas use a rich text editor to allow bold/italic/underline and other formatting.

## Spellcheck logic
- Spellcheck runs on text input widgets and underlines misspelled words.
- Right-click on a misspelled word to get suggestions.
- Ignore All adds the normalized word to the project-specific ignore list.
- The ignore list is saved inside the project file as spellcheck_ignore_all.
- Settings -> Edit Document Spellcheck Data... opens a dialog where you can add/remove ignored words.

### Notes about ignored words
- Ignored words are stored in lowercase.
- Add supports entering multiple words separated by spaces.

## Undo/Redo
- Undo/Redo tracks structural and text changes.
- Some operations (like typing) are grouped so you don't get a new undo step per keystroke.
- If you undo/redo, the UI is refreshed from the restored state.

## Preview
- Preview renders the current state into a read-only view.
- Use Preview to sanity-check spacing, ordering, and formatting before exporting.

## Export
- Export functions convert the current state into the requested format.
- If an export fails, the program shows an error describing what happened.

## Tips
- Save early so you always have a project file to recover from.
- Keep Skills concise and scannable.
- Use consistent tense and parallel structure across bullets.
"""


class ToolTip:
    def __init__(self, widget: tk.Widget, text: str, *, delay_ms: int = 600):
        self.widget = widget
        self.text = text
        self.delay_ms = delay_ms
        self._after_id: Optional[str] = None
        self._tip: Optional[tk.Toplevel] = None

        self.widget.bind("<Enter>", self._schedule, add="+")
        self.widget.bind("<Leave>", self._hide, add="+")
        self.widget.bind("<ButtonPress>", self._hide, add="+")

    def _schedule(self, _event=None):
        self._cancel()
        try:
            self._after_id = self.widget.after(self.delay_ms, self._show)
        except Exception:
            self._after_id = None

    def _cancel(self):
        if self._after_id is None:
            return
        try:
            self.widget.after_cancel(self._after_id)
        except Exception:
            pass
        self._after_id = None

    def _show(self):
        if self._tip is not None and self._tip.winfo_exists():
            return
        if not self.text:
            return
        try:
            x = int(self.widget.winfo_pointerx()) + 12
            y = int(self.widget.winfo_pointery()) + 12
        except Exception:
            try:
                x = int(self.widget.winfo_rootx()) + 12
                y = int(self.widget.winfo_rooty()) + 12
            except Exception:
                return

        tip = tk.Toplevel(self.widget)
        tip.overrideredirect(True)
        try:
            tip.attributes("-topmost", True)
        except Exception:
            pass

        lbl = tk.Label(
            tip,
            text=self.text,
            bg="#ffffe0",
            fg="#000000",
            bd=1,
            relief="solid",
            padx=6,
            pady=3,
        )
        lbl.pack()
        tip.geometry(f"+{x}+{y}")
        self._tip = tip

    def _hide(self, _event=None):
        self._cancel()
        try:
            if self._tip is not None and self._tip.winfo_exists():
                self._tip.destroy()
        except Exception:
            pass
        self._tip = None


class SingleLineTextField(ttk.Frame):
    def __init__(self, master, *, textvariable: Optional[tk.StringVar] = None):
        super().__init__(master)

        self.var = textvariable or tk.StringVar()
        self._syncing = False

        self.text = tk.Text(
            self,
            height=1,
            wrap="none",
            undo=True,
            borderwidth=1,
            relief="solid",
            padx=4,
            pady=2,
            highlightthickness=0,
        )
        self.text.pack(fill=tk.X, expand=True)

        try:
            self.text.insert("1.0", str(self.var.get() or ""))
            self.text.edit_modified(False)
        except Exception:
            pass

        try:
            self.var.trace_add("write", lambda *_a: self._sync_from_var())
        except Exception:
            pass

        self.text.bind("<<Modified>>", self._on_modified, add="+")
        self.text.bind("<KeyRelease>", self._on_modified, add="+")

        self.text.bind("<Return>", lambda _e: "break", add="+")
        self.text.bind("<KP_Enter>", lambda _e: "break", add="+")
        self.text.bind("<Control-j>", lambda _e: "break", add="+")
        self.text.bind("<Control-m>", lambda _e: "break", add="+")

        self.text.bind("<Tab>", self._on_tab, add="+")
        self.text.bind("<Shift-Tab>", self._on_shift_tab, add="+")

        self.text.bind("<<Paste>>", self._on_paste, add="+")
        self.text.bind("<Control-v>", self._on_paste, add="+")

        self.text.bind("<Control-a>", self._select_all, add="+")

    def _select_all(self, _event=None):
        try:
            self.text.tag_add("sel", "1.0", "end-1c")
            self.text.mark_set("insert", "end-1c")
            self.text.see("insert")
        except Exception:
            pass
        return "break"

    def _on_tab(self, _event=None):
        try:
            nxt = self.text.tk_focusNext()
            if nxt is not None:
                nxt.focus_set()
        except Exception:
            pass
        return "break"

    def _on_shift_tab(self, _event=None):
        try:
            prev = self.text.tk_focusPrev()
            if prev is not None:
                prev.focus_set()
        except Exception:
            pass
        return "break"

    def _sanitize(self, s: str) -> str:
        try:
            s = str(s or "")
        except Exception:
            return ""
        s = s.replace("\r", " ").replace("\n", " ")
        return s

    def _on_paste(self, _event=None):
        try:
            clip = self.text.clipboard_get()
        except Exception:
            clip = ""
        clip = self._sanitize(clip)

        try:
            if self.text.tag_ranges("sel"):
                self.text.delete("sel.first", "sel.last")
        except Exception:
            pass

        try:
            self.text.insert("insert", clip)
            self.text.edit_modified(True)
        except Exception:
            pass

        self._on_modified(None)
        return "break"

    def _sync_from_var(self) -> None:
        if self._syncing:
            return
        self._syncing = True
        try:
            target = self._sanitize(self.var.get())
            try:
                cur = self.text.get("1.0", "end-1c")
            except Exception:
                cur = ""

            if cur != target:
                try:
                    self.text.delete("1.0", "end")
                    self.text.insert("1.0", target)
                    self.text.edit_modified(False)
                except Exception:
                    pass
        finally:
            self._syncing = False

    def _on_modified(self, _event=None):
        if self._syncing:
            return
        try:
            if self.text.edit_modified():
                self.text.edit_modified(False)
        except Exception:
            pass

        try:
            raw = self.text.get("1.0", "end-1c")
        except Exception:
            return
        val = self._sanitize(raw)

        if val != raw:
            try:
                self.text.delete("1.0", "end")
                self.text.insert("1.0", val)
                self.text.edit_modified(False)
            except Exception:
                pass

        self._syncing = True
        try:
            try:
                self.var.set(val)
            except Exception:
                pass
        finally:
            self._syncing = False


class SpellCheckManager:
    def __init__(self, root: tk.Misc):
        self.root = root
        self.enabled = False

        self.ignore_all: set = set()

        self._spellchecker = None
        self._entry_menu: Optional[tk.Menu] = None
        self._text_menu: Optional[tk.Menu] = None

        self._text_widgets: List[tk.Text] = []
        self._entry_widgets: List[ttk.Entry] = []
        self._jobs: Dict[str, str] = {}
        self._entry_original_style: Dict[str, str] = {}

        self._entry_underlines: Dict[str, tk.Frame] = {}

        try:
            ttk.Style(self.root).configure(
                "SpellError.TEntry",
                fieldbackground="#ffecec",
            )
        except Exception:
            pass

    def _ensure_spellchecker(self) -> bool:
        if self._spellchecker is not None:
            return True
        try:
            from spellchecker import SpellChecker  # type: ignore

            self._spellchecker = SpellChecker(language="en")
            return True
        except Exception:
            return False

    def set_enabled(self, enabled: bool) -> None:
        self.enabled = bool(enabled)
        self.refresh_all()

    def set_ignore_all(self, words: Any) -> None:
        out = set()
        if isinstance(words, (list, tuple, set)):
            for w in words:
                try:
                    s = str(w).strip()
                except Exception:
                    continue
                if not s:
                    continue
                out.add(s.lower())
        self.ignore_all = out
        self.refresh_all()

    def get_ignore_all(self) -> List[str]:
        try:
            return sorted({str(w) for w in (self.ignore_all or set()) if str(w)})
        except Exception:
            return []

    def register_rich_text_editor(self, rte: "RichTextEditor") -> None:
        try:
            self.register_text(rte.text)
        except Exception:
            return

    def register_text(self, text: tk.Text) -> None:
        key = str(text)
        if any(str(t) == key for t in self._text_widgets):
            return
        self._text_widgets.append(text)

        try:
            text.tag_configure("__spell_error__", underline=True, underlinefg="#ff0000")
        except Exception:
            try:
                text.tag_configure("__spell_error__", underline=True, foreground="#ff0000")
            except Exception:
                pass

        text.bind("<<Modified>>", lambda _e, w=text: self._schedule_text_check(w), add="+")
        text.bind("<KeyRelease>", lambda _e, w=text: self._schedule_text_check(w), add="+")
        text.bind("<ButtonRelease-1>", lambda _e, w=text: self._schedule_text_check(w), add="+")
        text.bind("<Button-3>", lambda e, w=text: self._on_text_right_click(e, w), add="+")

        self._schedule_text_check(text, delay_ms=10)

    def register_entry(self, entry: ttk.Entry) -> None:
        key = str(entry)
        if any(str(e) == key for e in self._entry_widgets):
            return
        self._entry_widgets.append(entry)

        try:
            self._entry_original_style[key] = str(entry.cget("style") or "")
        except Exception:
            self._entry_original_style[key] = ""

        entry.bind("<KeyRelease>", lambda _e, w=entry: self._schedule_entry_check(w), add="+")
        entry.bind("<FocusOut>", lambda _e, w=entry: self._schedule_entry_check(w, delay_ms=10), add="+")
        entry.bind("<Button-3>", lambda e, w=entry: self._on_entry_right_click(e, w), add="+")

        try:
            ul = tk.Frame(entry.master, height=2, bg="#ff0000", bd=0, highlightthickness=0)
            ul.place_forget()
            self._entry_underlines[key] = ul

            def _reposition(_e=None, w=entry):
                self._position_entry_underline(w)

            entry.bind("<Configure>", _reposition, add="+")
            try:
                entry.master.bind("<Configure>", _reposition, add="+")
            except Exception:
                pass

            entry.bind(
                "<Destroy>",
                lambda _e, ww=ul, k=key: self._on_entry_destroyed(k, ww),
                add="+",
            )
        except Exception:
            pass

        self._schedule_entry_check(entry, delay_ms=10)

    def _on_entry_destroyed(self, key: str, underline: tk.Frame) -> None:
        try:
            self._entry_underlines.pop(key, None)
        except Exception:
            pass
        try:
            underline.destroy()
        except Exception:
            pass

    def _position_entry_underline(self, entry: ttk.Entry) -> None:
        key = str(entry)
        ul = self._entry_underlines.get(key)
        if ul is None:
            return
        try:
            if not ul.winfo_exists():
                return
        except Exception:
            return

        try:
            if ul.winfo_manager() != "place":
                return
        except Exception:
            return

        try:
            x = int(entry.winfo_x())
            y = int(entry.winfo_y()) + int(entry.winfo_height()) - 1
            w = int(entry.winfo_width())
            if w <= 1:
                w = int(entry.winfo_reqwidth())
            ul.place(x=x, y=y, width=w)
            ul.lift()
        except Exception:
            pass

    def _show_entry_underline(self, entry: ttk.Entry) -> None:
        key = str(entry)
        ul = self._entry_underlines.get(key)
        if ul is None:
            return

        def _do() -> None:
            try:
                x = int(entry.winfo_x())
                y = int(entry.winfo_y()) + int(entry.winfo_height()) - 1
                w = int(entry.winfo_width())
                if w <= 1:
                    w = int(entry.winfo_reqwidth())
                ul.place(x=x, y=y, width=w)
                ul.lift()
            except Exception:
                pass

        try:
            entry.after_idle(_do)
        except Exception:
            _do()

    def _hide_entry_underline(self, entry: ttk.Entry) -> None:
        key = str(entry)
        ul = self._entry_underlines.get(key)
        if ul is None:
            return
        try:
            ul.place_forget()
        except Exception:
            pass

    def _ignore_word_all(self, word: str) -> None:
        w = self._normalize_word(word)
        if not w:
            return
        if not self._is_word_candidate(w):
            return
        try:
            self.ignore_all.add(w.lower())
        except Exception:
            return
        self.refresh_all()

    def _schedule_text_check(self, text: tk.Text, *, delay_ms: int = 350) -> None:
        key = str(text)
        prev = self._jobs.get(key)
        if prev:
            try:
                text.after_cancel(prev)
            except Exception:
                pass

        try:
            jid = text.after(delay_ms, lambda w=text: self._check_text_now(w))
            self._jobs[key] = str(jid)
        except Exception:
            pass

    def _schedule_entry_check(self, entry: ttk.Entry, *, delay_ms: int = 350) -> None:
        key = str(entry)
        prev = self._jobs.get(key)
        if prev:
            try:
                entry.after_cancel(prev)
            except Exception:
                pass

        try:
            jid = entry.after(delay_ms, lambda w=entry: self._check_entry_now(w))
            self._jobs[key] = str(jid)
        except Exception:
            pass

    def _should_check_header_entry(self, entry: ttk.Entry) -> bool:
        return False

    def _normalize_word(self, w: str) -> str:
        return (w or "").strip("'\"-–—()[]{}.,;:!?/\\").strip()

    def _is_word_candidate(self, w: str) -> bool:
        if not w:
            return False
        if "@" in w:
            return False
        if w.startswith("http") or w.startswith("www"):
            return False
        if any(ch.isdigit() for ch in w):
            return False
        if w.isupper() and len(w) <= 4:
            return False
        return any(ch.isalpha() for ch in w)

    def _misspelled(self, word: str) -> bool:
        if not self._ensure_spellchecker():
            return False
        w = self._normalize_word(word)
        if not self._is_word_candidate(w):
            return False
        if w.lower() in (self.ignore_all or set()):
            return False
        try:
            return w.lower() in self._spellchecker.unknown([w.lower()])
        except Exception:
            return False

    def _candidates(self, word: str) -> List[str]:
        if not self._ensure_spellchecker():
            return []
        w = self._normalize_word(word)
        if not self._is_word_candidate(w):
            return []
        if w.lower() in (self.ignore_all or set()):
            return []
        try:
            cands = list(self._spellchecker.candidates(w.lower()) or [])
        except Exception:
            cands = []
        out: List[str] = []
        for c in cands:
            if not c:
                continue
            s = str(c)
            if s.lower() == w.lower():
                continue
            if s not in out:
                out.append(s)
            if len(out) >= 8:
                break
        return out

    def _check_text_now(self, text: tk.Text) -> None:
        if not self.enabled:
            try:
                text.tag_remove("__spell_error__", "1.0", "end")
            except Exception:
                pass
            return

        if not self._ensure_spellchecker():
            try:
                text.tag_remove("__spell_error__", "1.0", "end")
            except Exception:
                pass
            return

        try:
            s = text.get("1.0", "end-1c")
        except Exception:
            return

        try:
            text.tag_remove("__spell_error__", "1.0", "end")
        except Exception:
            pass

        for m in re.finditer(r"[A-Za-z][A-Za-z']*", s):
            w = m.group(0)
            if not self._misspelled(w):
                continue
            try:
                start = f"1.0+{m.start()}c"
                end = f"1.0+{m.end()}c"
                text.tag_add("__spell_error__", start, end)
            except Exception:
                continue

    def _check_entry_now(self, entry: ttk.Entry) -> None:
        if not self.enabled:
            self._set_entry_ok(entry)
            return
        if not self._ensure_spellchecker():
            self._set_entry_ok(entry)
            return

        try:
            s = str(entry.get() or "")
        except Exception:
            return

        miss = False
        for m in re.finditer(r"[A-Za-z][A-Za-z']*", s):
            if self._misspelled(m.group(0)):
                miss = True
                break

        if miss:
            self._set_entry_error(entry)
        else:
            self._set_entry_ok(entry)

    def _set_entry_error(self, entry: ttk.Entry) -> None:
        try:
            entry.configure(style="SpellError.TEntry")
        except Exception:
            pass
        self._show_entry_underline(entry)

    def _set_entry_ok(self, entry: ttk.Entry) -> None:
        key = str(entry)
        style = self._entry_original_style.get(key, "")
        try:
            entry.configure(style=style)
        except Exception:
            pass
        self._hide_entry_underline(entry)

    def refresh_all(self) -> None:
        for t in list(self._text_widgets):
            try:
                if t.winfo_exists():
                    self._schedule_text_check(t, delay_ms=10)
            except Exception:
                pass

        for e in list(self._entry_widgets):
            try:
                if e.winfo_exists():
                    self._schedule_entry_check(e, delay_ms=10)
            except Exception:
                pass

    def _text_word_at_event(self, text: tk.Text, event) -> Tuple[Optional[str], Optional[str], Optional[str]]:
        try:
            idx = text.index(f"@{event.x},{event.y}")
        except Exception:
            return (None, None, None)

        try:
            start = text.index(f"{idx} wordstart")
            end = text.index(f"{idx} wordend")
            w = text.get(start, end)
        except Exception:
            return (None, None, None)

        w = str(w or "")
        if not w:
            return (None, None, None)
        return (w, start, end)

    def _entry_word_at_event(self, entry: ttk.Entry, event) -> Tuple[Optional[str], Optional[int], Optional[int]]:
        try:
            i = int(entry.index(f"@{event.x}"))
        except Exception:
            try:
                i = int(entry.index("insert"))
            except Exception:
                return (None, None, None)

        try:
            s = str(entry.get() or "")
        except Exception:
            return (None, None, None)

        if not s:
            return (None, None, None)

        i = max(0, min(i, len(s)))
        left = i
        while left > 0 and re.fullmatch(r"[A-Za-z']", s[left - 1]):
            left -= 1
        right = i
        while right < len(s) and re.fullmatch(r"[A-Za-z']", s[right]):
            right += 1

        if right <= left:
            return (None, None, None)
        return (s[left:right], left, right)

    def _on_text_right_click(self, event, text: tk.Text):
        if not self.enabled:
            return
        if not self._ensure_spellchecker():
            return

        w, start, end = self._text_word_at_event(text, event)
        if not w or not start or not end:
            return
        if not self._misspelled(w):
            return

        cands = self._candidates(w)

        m = self._text_menu
        try:
            if m is None or not m.winfo_exists():
                m = None
        except Exception:
            m = None

        if m is None:
            try:
                m = tk.Menu(self.root, tearoff=False)
            except Exception:
                m = tk.Menu(text, tearoff=False)
            self._text_menu = m

        try:
            m.delete(0, tk.END)
        except Exception:
            try:
                m = tk.Menu(self.root, tearoff=False)
            except Exception:
                m = tk.Menu(text, tearoff=False)
            self._text_menu = m
            try:
                m.delete(0, tk.END)
            except Exception:
                return

        for c in cands:
            try:
                m.add_command(
                    label=c,
                    command=lambda rep=c, s=start, e=end, tw=text: self._replace_text_word(tw, s, e, rep),
                )
            except Exception:
                return

        if cands:
            m.add_separator()
        m.add_command(label="Ignore All", command=lambda ww=w: self._ignore_word_all(ww))

        try:
            m.tk_popup(event.x_root, event.y_root)
        except Exception:
            pass

    def _replace_text_word(self, text: tk.Text, start: str, end: str, replacement: str) -> None:
        try:
            text.delete(start, end)
            text.insert(start, replacement)
        except Exception:
            return
        self._schedule_text_check(text, delay_ms=10)

    def _on_entry_right_click(self, event, entry: ttk.Entry):
        if not self.enabled:
            return
        if not self._ensure_spellchecker():
            return

        w, start, end = self._entry_word_at_event(entry, event)
        if w is None or start is None or end is None:
            return
        if not self._misspelled(w):
            return

        cands = self._candidates(w)

        m = self._entry_menu
        try:
            if m is None or not m.winfo_exists():
                m = None
        except Exception:
            m = None

        if m is None:
            try:
                m = tk.Menu(self.root, tearoff=False)
            except Exception:
                m = tk.Menu(entry, tearoff=False)
            self._entry_menu = m

        try:
            m.delete(0, tk.END)
        except Exception:
            try:
                m = tk.Menu(self.root, tearoff=False)
            except Exception:
                m = tk.Menu(entry, tearoff=False)
            self._entry_menu = m
            try:
                m.delete(0, tk.END)
            except Exception:
                return

        for c in cands:
            try:
                m.add_command(
                    label=c,
                    command=lambda rep=c, a=start, b=end, ew=entry: self._replace_entry_word(ew, a, b, rep),
                )
            except Exception:
                return

        if cands:
            m.add_separator()
        m.add_command(label="Ignore All", command=lambda ww=w: self._ignore_word_all(ww))

        try:
            m.tk_popup(event.x_root, event.y_root)
        except Exception:
            pass

    def _replace_entry_word(self, entry: ttk.Entry, start: int, end: int, replacement: str) -> None:
        try:
            s = str(entry.get() or "")
        except Exception:
            return
        new_s = s[:start] + replacement + s[end:]
        try:
            entry.delete(0, tk.END)
            entry.insert(0, new_s)
        except Exception:
            return
        self._schedule_entry_check(entry, delay_ms=10)


def _find_spellcheck_manager(widget: tk.Misc) -> Optional[SpellCheckManager]:
    w: Any = widget
    while w is not None:
        try:
            if hasattr(w, "spellcheck_manager"):
                mgr = getattr(w, "spellcheck_manager")
                if isinstance(mgr, SpellCheckManager):
                    return mgr
        except Exception:
            pass

        try:
            w = getattr(w, "master", None)
        except Exception:
            w = None
    return None


class SpellcheckIgnoreAllDialog(tk.Toplevel):
    def __init__(
        self,
        master: tk.Misc,
        *,
        words: List[str],
        on_apply: Callable[[List[str]], None],
    ):
        super().__init__(master)
        self.title("Edit Document Spellcheck Data")
        try:
            self.transient(master)
        except Exception:
            pass
        try:
            self.grab_set()
        except Exception:
            pass

        self._on_apply = on_apply
        self._words = sorted({str(w).strip().lower() for w in (words or []) if str(w).strip()})

        root = ttk.Frame(self)
        root.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        ttk.Label(root, text="Ignored words (Ignore All):").pack(anchor="w")

        list_frame = ttk.Frame(root)
        list_frame.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

        scroll = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.listbox = tk.Listbox(list_frame, height=12, yscrollcommand=scroll.set)
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll.config(command=self.listbox.yview)

        btns = ttk.Frame(root)
        btns.pack(fill=tk.X, pady=(8, 0))
        ttk.Button(btns, text="Add", command=self.add_word).pack(side=tk.LEFT)
        ttk.Button(btns, text="Delete", command=self.delete_selected).pack(side=tk.LEFT, padx=(6, 0))

        action = ttk.Frame(root)
        action.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(action, text="Apply", command=self.apply_changes).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(action, text="Close", command=self.close).pack(side=tk.RIGHT)

        self.refresh_list()
        self.minsize(420, 320)
        try:
            self.after(0, lambda: _center_window(self))
        except Exception:
            pass

        try:
            self.protocol("WM_DELETE_WINDOW", self.close)
        except Exception:
            pass

        try:
            self.bind("<Escape>", lambda _e: self.close())
        except Exception:
            pass

    def refresh_list(self) -> None:
        try:
            self.listbox.delete(0, tk.END)
        except Exception:
            return
        for w in self._words:
            try:
                self.listbox.insert(tk.END, w)
            except Exception:
                pass

    def add_word(self) -> None:
        raw = _ask_string_centered(self, "Add Ignored Word", "Word(s) to ignore:")
        if not raw:
            return

        tokens = [t.strip().lower() for t in str(raw).strip().split() if t.strip()]
        if not tokens:
            return

        changed = False
        for t in tokens:
            if t not in self._words:
                self._words.append(t)
                changed = True

        if changed:
            self._words = sorted(self._words)
            self.refresh_list()

    def delete_selected(self) -> None:
        try:
            sel = list(self.listbox.curselection() or [])
        except Exception:
            sel = []
        if not sel:
            return
        idx = int(sel[0])
        if idx < 0 or idx >= len(self._words):
            return
        try:
            self._words.pop(idx)
        except Exception:
            return
        self.refresh_list()

    def apply_changes(self) -> None:
        try:
            self._on_apply(sorted({str(w).strip().lower() for w in (self._words or []) if str(w).strip()}))
        except Exception:
            return

    def close(self) -> None:
        try:
            self.destroy()
        except Exception:
            pass


class HelpDialog(tk.Toplevel):
    def __init__(self, master: tk.Misc, *, title: str, content: str):
        super().__init__(master)
        self.title(title)
        try:
            self.transient(master)
        except Exception:
            pass

        container = ttk.Frame(self)
        container.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        text_frame = ttk.Frame(container)
        text_frame.pack(fill=tk.BOTH, expand=True)

        scroll = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
        scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.text = tk.Text(
            text_frame,
            wrap="word",
            yscrollcommand=scroll.set,
            padx=10,
            pady=10,
            borderwidth=0,
            highlightthickness=0,
        )
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll.config(command=self.text.yview)

        try:
            base = tkfont.nametofont("TkDefaultFont")
            base_family = str(base.cget("family"))
            base_size = int(base.cget("size"))
        except Exception:
            base_family = "Arial"
            base_size = 10

        self._font_h1 = tkfont.Font(family=base_family, size=base_size + 6, weight="bold")
        self._font_h2 = tkfont.Font(family=base_family, size=base_size + 3, weight="bold")
        self._font_body = tkfont.Font(family=base_family, size=base_size)

        self.text.tag_configure("h1", font=self._font_h1, spacing1=6, spacing3=6)
        self.text.tag_configure("h2", font=self._font_h2, spacing1=4, spacing3=4)
        self.text.tag_configure("body", font=self._font_body)
        self.text.tag_configure("bullet", font=self._font_body, lmargin1=16, lmargin2=28)

        self._insert_markdownish(content)

        try:
            self.text.configure(state="disabled")
        except Exception:
            pass

        action = ttk.Frame(container)
        action.pack(fill=tk.X, pady=(10, 0))
        ttk.Button(action, text="Close", command=self.close).pack(side=tk.RIGHT)

        self.minsize(720, 520)
        try:
            self.after(0, lambda: _center_window(self))
        except Exception:
            pass

        try:
            self.protocol("WM_DELETE_WINDOW", self.close)
        except Exception:
            pass

        try:
            self.bind("<Escape>", lambda _e: self.close())
        except Exception:
            pass

    def _insert_markdownish(self, content: str) -> None:
        lines = str(content or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
        for line in lines:
            raw = line.rstrip("\n")
            if raw.startswith("# "):
                self.text.insert("end", raw[2:].strip() + "\n", ("h1",))
            elif raw.startswith("## "):
                self.text.insert("end", raw[3:].strip() + "\n", ("h2",))
            elif raw.startswith("- "):
                self.text.insert("end", "- " + raw[2:].strip() + "\n", ("bullet",))
            else:
                self.text.insert("end", raw + "\n", ("body",))

    def close(self) -> None:
        try:
            self.destroy()
        except Exception:
            pass


def _center_window(win: tk.Toplevel) -> None:
    try:
        if win is None or not win.winfo_exists():
            return
    except Exception:
        return

    try:
        win.update_idletasks()
    except Exception:
        pass

    try:
        width = int(win.winfo_width())
        height = int(win.winfo_height())
    except Exception:
        width, height = 0, 0

    if width <= 1 or height <= 1:
        try:
            width = int(win.winfo_reqwidth())
            height = int(win.winfo_reqheight())
        except Exception:
            return

    try:
        sw = int(win.winfo_screenwidth())
        sh = int(win.winfo_screenheight())
    except Exception:
        return

    x = max(0, int((sw - width) / 2))
    y = max(0, int((sh - height) / 2))

    try:
        win.geometry(f"{width}x{height}+{x}+{y}")
    except Exception:
        try:
            win.geometry(f"+{x}+{y}")
        except Exception:
            pass


def _ask_string_centered(
    parent: tk.Misc,
    title: str,
    prompt: str,
    *,
    initialvalue: str = "",
) -> Optional[str]:
    result: Dict[str, Optional[str]] = {"value": None}

    dlg = tk.Toplevel(parent)
    dlg.title(title)
    try:
        dlg.transient(parent.winfo_toplevel())
    except Exception:
        pass

    root = ttk.Frame(dlg)
    root.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)

    ttk.Label(root, text=prompt).pack(anchor="w")
    var = tk.StringVar(value=str(initialvalue or ""))
    ent = ttk.Entry(root, textvariable=var)
    ent.pack(fill=tk.X, pady=(6, 10))

    btns = ttk.Frame(root)
    btns.pack(fill=tk.X)

    def _ok() -> None:
        result["value"] = str(var.get() or "")
        try:
            dlg.destroy()
        except Exception:
            pass

    def _cancel() -> None:
        result["value"] = None
        try:
            dlg.destroy()
        except Exception:
            pass

    ttk.Button(btns, text="Cancel", command=_cancel).pack(side=tk.RIGHT, padx=(6, 0))
    ttk.Button(btns, text="OK", command=_ok).pack(side=tk.RIGHT)

    try:
        dlg.protocol("WM_DELETE_WINDOW", _cancel)
    except Exception:
        pass
    try:
        dlg.bind("<Return>", lambda _e: _ok())
        dlg.bind("<Escape>", lambda _e: _cancel())
    except Exception:
        pass

    try:
        ent.focus_set()
        ent.selection_range(0, tk.END)
    except Exception:
        pass

    try:
        dlg.grab_set()
    except Exception:
        pass

    try:
        dlg.update_idletasks()
        _center_window(dlg)
    except Exception:
        pass

    try:
        parent.wait_window(dlg)
    except Exception:
        try:
            dlg.wait_window(dlg)
        except Exception:
            pass

    val = result.get("value")
    if val is None:
        return None
    val = str(val)
    if val.strip() == "":
        return None
    return val


def apply_modern_theme(root: tk.Tk) -> None:
    style = ttk.Style(root)
    try:
        windowing = root.tk.call("tk", "windowingsystem")
    except Exception:
        windowing = ""

    preferred: List[str] = []
    if windowing == "win32":
        preferred = ["vista", "xpnative"]
    elif windowing == "aqua":
        preferred = ["aqua"]
    else:
        preferred = ["clam", "alt"]

    for theme in preferred:
        if theme in style.theme_names():
            style.theme_use(theme)
            break

    try:
        default_font = tkfont.nametofont("TkDefaultFont")
        base_family = default_font.cget("family")
        base_size = int(default_font.cget("size"))
        style.configure("TButton", padding=(10, 6))
        style.configure("TEntry", padding=(6, 4))
        style.configure("TCombobox", padding=(6, 4))
        style.configure("TLabelframe.Label", font=(base_family, base_size, "bold"))
        style.configure("Heading.Treeview", font=(base_family, base_size, "bold"))
        style.configure("Treeview", rowheight=max(24, base_size + 12))
    except Exception:
        pass


DEFAULT_DATA = {
    "header": {
        "name": "",
        "phone": "",
        "email": "",
        "linkedin_kind": "LinkedIn",
        "linkedin": "",
        "linkedin_display": "",
        "github_kind": "GitHub",
        "github": "",
        "github_display": "",
    },
    "sections": [
        {"id": "education", "title": "Education", "kind": "education", "entries": []},
        {"id": "experience", "title": "Experience", "kind": "experience", "entries": []},
        {"id": "projects", "title": "Projects", "kind": "projects", "entries": []},
        {"id": "technical_skills", "title": "Technical Skills", "kind": "skills", "entries": []},
    ],
}


DEMO_DATA = {
    "header": {
        "name": "Jake Ryan",
        "phone": "123-456-7890",
        "email": "jake@su.edu",
        "linkedin_kind": "LinkedIn",
        "linkedin": "https://linkedin.com/in/...",
        "linkedin_display": "linkedin.com/in/jake",
        "github_kind": "GitHub",
        "github": "https://github.com/...",
        "github_display": "github.com/jake",
    },
    "sections": [
        {
            "id": "education",
            "title": "Education",
            "kind": "education",
            "entries": [
                {
                    "school": "Southwestern University",
                    "location": "Georgetown, TX",
                    "degree": "Bachelor of Arts in Computer Science, Minor in Business",
                    "dates": "Aug. 2018 -- May 2021",
                },
                {
                    "school": "Blinn College",
                    "location": "Bryan, TX",
                    "degree": "Associate's in Liberal Arts",
                    "dates": "Aug. 2014 -- May 2018",
                },
            ],
        },
        {
            "id": "experience",
            "title": "Experience",
            "kind": "experience",
            "entries": [
                {
                    "role": "Undergraduate Research Assistant",
                    "dates": "June 2020 -- Present",
                    "org": "Texas A&M University",
                    "location": "College Station, TX",
                    "bullets": [
                        [{"text": "Developed a REST API using FastAPI and PostgreSQL to store data from learning management systems"}],
                        [{"text": "Developed a full-stack web application using Flask, React, PostgreSQL and Docker to analyze GitHub data"}],
                        [{"text": "Explored ways to visualize GitHub collaboration in a classroom setting"}],
                    ],
                },
                {
                    "role": "Information Technology Support Specialist",
                    "dates": "Sep. 2018 -- Present",
                    "org": "Southwestern University",
                    "location": "Georgetown, TX",
                    "bullets": [
                        [{"text": "Communicate with managers to set up campus computers used on campus"}],
                        [{"text": "Assess and troubleshoot computer problems brought by students, faculty and staff"}],
                        [{"text": "Maintain upkeep of computers, classroom equipment, and 200 printers across campus"}],
                    ],
                },
                {
                    "role": "Artificial Intelligence Research Assistant",
                    "dates": "May 2019 -- July 2019",
                    "org": "Southwestern University",
                    "location": "Georgetown, TX",
                    "bullets": [
                        [
                            {
                                "text": "Explored methods to generate video game dungeons based off of ",
                            },
                            {"text": "The Legend of Zelda", "i": True},
                        ],
                        [{"text": "Developed a game in Java to test the generated dungeons"}],
                        [{"text": "Contributed 50K+ lines of code to an established codebase via Git"}],
                        [{"text": "Conducted a human subject study to determine which video game dungeon generation technique is enjoyable"}],
                        [{"text": "Wrote an 8-page paper and gave multiple presentations on-campus"}],
                        [{"text": "Presented virtually to the World Conference on Computational Intelligence"}],
                    ],
                },
            ],
        },
        {
            "id": "projects",
            "title": "Projects",
            "kind": "projects",
            "entries": [
                {
                    "title": "Gitlytics",
                    "stack": "Python, Flask, React, PostgreSQL, Docker",
                    "dates": "June 2020 -- Present",
                    "bullets": [
                        [{"text": "Developed a full-stack web application using with Flask serving a REST API with React as the frontend"}],
                        [{"text": "Implemented GitHub OAuth to get data from user's repositories"}],
                        [{"text": "Visualized GitHub data to show collaboration"}],
                        [{"text": "Used Celery and Redis for asynchronous tasks"}],
                    ],
                },
                {
                    "title": "Simple Paintball",
                    "stack": "Spigot API, Java, Maven, TravisCI, Git",
                    "dates": "May 2018 -- May 2020",
                    "bullets": [
                        [{"text": "Developed a Minecraft server plugin to entertain kids during free time for a previous job"}],
                        [{"text": "Published plugin to websites gaining 2K+ downloads and an average 4.5/5-star review"}],
                        [{"text": "Implemented continuous delivery using TravisCI to build the plugin upon new a release"}],
                        [{"text": "Collaborated with Minecraft server administrators to suggest features and get feedback about the plugin"}],
                    ],
                },
            ],
        },
        {
            "id": "technical_skills",
            "title": "Technical Skills",
            "kind": "skills",
            "entries": [
                {
                    "label": "Languages",
                    "value": [
                        {"text": "Java, Python, C/C++, SQL (Postgres), JavaScript, HTML/CSS, R"}
                    ],
                },
                {
                    "label": "Frameworks",
                    "value": [
                        {"text": "React, Node.js, Flask, JUnit, WordPress, Material-UI, FastAPI"}
                    ],
                },
                {
                    "label": "Developer Tools",
                    "value": [
                        {
                            "text": "Git, Docker, TravisCI, Google Cloud Platform, VS Code, Visual Studio, PyCharm, IntelliJ, Eclipse"
                        }
                    ],
                },
                {
                    "label": "Libraries",
                    "value": [{"text": "pandas, NumPy, Matplotlib"}],
                },
            ],
        },
    ],
}


LATEX_PREAMBLE = r"""%-------------------------
% Resume in Latex
% Author : Jake Gutierrez
% Based off of: https://github.com/sb2nov/resume
% License : MIT
%------------------------

\documentclass[letterpaper,11pt]{article}

\usepackage{latexsym}
\usepackage[empty]{fullpage}
\usepackage{titlesec}
\usepackage{marvosym}
\usepackage[usenames,dvipsnames]{color}
\usepackage[usenames,dvipsnames]{xcolor}
\usepackage{verbatim}
\usepackage{enumitem}
\usepackage[hidelinks]{hyperref}
\usepackage{fancyhdr}
\usepackage[english]{babel}
\usepackage{tabularx}
\input{glyphtounicode}


%----------FONT OPTIONS----------
% sans-serif
% \usepackage[sfdefault]{FiraSans}
% \usepackage[sfdefault]{roboto}
% \usepackage[sfdefault]{noto-sans}
% \usepackage[default]{sourcesanspro}

% serif
% \usepackage{CormorantGaramond}
% \usepackage{charter}


\pagestyle{fancy}
\fancyhf{} % clear all header and footer fields
\fancyfoot{}
\renewcommand{\headrulewidth}{0pt}
\renewcommand{\footrulewidth}{0pt}

% Adjust margins
\addtolength{\oddsidemargin}{-0.5in}
\addtolength{\evensidemargin}{-0.5in}
\addtolength{\textwidth}{1in}
\addtolength{\topmargin}{-.5in}
\addtolength{\textheight}{1.0in}

\urlstyle{same}

\raggedbottom
\raggedright
\setlength{\tabcolsep}{0in}

% Sections formatting
\titleformat{\section}{
  \vspace{-4pt}\scshape\raggedright\large
}{}{0em}{}[\color{black}\titlerule \vspace{-5pt}]

% Ensure that generate pdf is machine readable/ATS parsable
\pdfgentounicode=1

%-------------------------
% Custom commands
\newcommand{\resumeItem}[1]{
  \item\small{
    {#1 \vspace{-2pt}}
  }
}

\newcommand{\resumeSubheading}[4]{
  \vspace{-2pt}\item
    \begin{tabular*}{0.97\textwidth}[t]{l@{\extracolsep{\fill}}r}
      \textbf{#1} & #2 \\
      \textit{\small#3} & \textit{\small #4} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubSubheading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \textit{\small#1} & \textit{\small #2} \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeProjectHeading}[2]{
    \item
    \begin{tabular*}{0.97\textwidth}{l@{\extracolsep{\fill}}r}
      \small#1 & #2 \\
    \end{tabular*}\vspace{-7pt}
}

\newcommand{\resumeSubItem}[1]{\resumeItem{#1}\vspace{-4pt}}

\renewcommand\labelitemii{$\vcenter{\hbox{\tiny$\bullet$}}$}

\newcommand{\resumeSubHeadingListStart}{\begin{itemize}[leftmargin=0.15in, label={}]}
\newcommand{\resumeSubHeadingListEnd}{\end{itemize}}
\newcommand{\resumeItemListStart}{\begin{itemize}}
\newcommand{\resumeItemListEnd}{\end{itemize}\vspace{-5pt}}

%-------------------------------------------
%%%%%%  RESUME STARTS HERE  %%%%%%%%%%%%%%%%%%%%%%%%%%%%


\begin{document}
"""


LATEX_END = r"""
%-------------------------------------------
\end{document}
"""


def _workspace_path() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _help_docs_dir() -> str:
    return os.path.join(_workspace_path(), "Program Help Docs")


def _app_data_dir() -> str:
    base = os.getenv("APPDATA") or os.path.expanduser("~")
    path = os.path.join(base, "ResumeTemplateMaker")
    try:
        os.makedirs(path, exist_ok=True)
    except Exception:
        pass
    return path


def _data_file_path() -> str:
    if getattr(sys, "frozen", False):
        return os.path.join(_app_data_dir(), DATA_FILE_NAME)
    return os.path.join(_workspace_path(), DATA_FILE_NAME)


def _export_file_path() -> str:
    return os.path.join(_workspace_path(), EXPORT_FILE_NAME)


def deep_copy(obj):
    return json.loads(json.dumps(obj))


_LATEX_ESCAPE_REPLACEMENTS = {
    "\\": r"\textbackslash{}",
    "{": r"\{",
    "}": r"\}",
    "&": r"\&",
    "%": r"\%",
    "$": r"\$",
    "#": r"\#",
    "_": r"\_",
    "~": r"\textasciitilde{}",
    "^": r"\textasciicircum{}",
}


def latex_escape(s: str) -> str:
    return "".join(_LATEX_ESCAPE_REPLACEMENTS.get(ch, ch) for ch in s)


def _normalize_href(url: str) -> str:
    u = ("" if url is None else str(url)).strip()
    if not u:
        return ""
    try:
        if re.match(r"^[a-zA-Z][a-zA-Z0-9+.\-]*:", u):
            return u
    except Exception:
        return u
    return "https://" + u


def _hex_to_rgb_floats(hex_color: str) -> Optional[Tuple[float, float, float]]:
    m = re.fullmatch(r"#?([0-9a-fA-F]{6})", hex_color or "")
    if not m:
        return None
    v = m.group(1)
    r = int(v[0:2], 16) / 255.0
    g = int(v[2:4], 16) / 255.0
    b = int(v[4:6], 16) / 255.0
    return (r, g, b)


def rich_segments_to_latex(segments) -> str:
    out_parts: List[str] = []
    for seg in segments or []:
        raw_text = seg.get("text", "")
        if raw_text == "":
            continue
        t = latex_escape(raw_text)
        if seg.get("i"):
            t = rf"\emph{{{t}}}"
        if seg.get("u"):
            t = rf"\underline{{{t}}}"
        if seg.get("b"):
            t = rf"\textbf{{{t}}}"

        fg = seg.get("fg")
        bg = seg.get("bg")

        fg_rgb = _hex_to_rgb_floats(fg) if fg else None
        bg_rgb = _hex_to_rgb_floats(bg) if bg else None

        if fg_rgb is not None:
            t = rf"\textcolor[rgb]{{{fg_rgb[0]:.3f},{fg_rgb[1]:.3f},{fg_rgb[2]:.3f}}}{{{t}}}"
        if bg_rgb is not None:
            t = rf"\colorbox[rgb]{{{bg_rgb[0]:.3f},{bg_rgb[1]:.3f},{bg_rgb[2]:.3f}}}{{{t}}}"

        size = seg.get("size")
        if size is not None:
            try:
                s_int = int(size)
            except Exception:
                s_int = 0
            if s_int > 0:
                baseline = max(int(round(s_int * 1.2)), s_int + 1)
                t = rf"{{\fontsize{{{s_int}}}{{{baseline}}}\selectfont {t}}}"

        out_parts.append(t)
    return "".join(out_parts)


def load_resume_data() -> dict:
    path = _data_file_path()
    if not os.path.exists(path):
        return deep_copy(DEFAULT_DATA)
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return deep_copy(DEFAULT_DATA)


def save_resume_data(data: dict) -> None:
    path = _data_file_path()
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def generate_latex(data: dict) -> str:
    header = data.get("header", {})
    name = latex_escape(header.get("name", ""))
    phone = latex_escape(header.get("phone", ""))
    email = header.get("email", "")
    linkedin_kind = str(header.get("linkedin_kind", "LinkedIn") or "").strip()
    linkedin = "" if linkedin_kind.lower() == "none" else header.get("linkedin", "")
    linkedin_display_override = header.get("linkedin_display", "")
    github_kind = str(header.get("github_kind", "GitHub") or "").strip()
    github = "" if github_kind.lower() == "none" else header.get("github", "")
    github_display_override = header.get("github_display", "")

    email_display = latex_escape(email)
    email_link = latex_escape(email)

    linkedin = _normalize_href(linkedin)
    github = _normalize_href(github)

    linkedin_display = latex_escape(
        (linkedin_display_override or linkedin.replace("https://", "").replace("http://", ""))
    )
    github_display = latex_escape(
        (github_display_override or github.replace("https://", "").replace("http://", ""))
    )

    parts: List[str] = [LATEX_PREAMBLE]

    contact_parts: List[str] = []
    if phone:
        contact_parts.append(phone)
    if email:
        contact_parts.append(rf"\href{{mailto:{email_link}}}{{\underline{{{email_display}}}}}")
    if linkedin:
        contact_parts.append(
            rf"\href{{{latex_escape(linkedin)}}}{{\underline{{{linkedin_display}}}}}"
        )
    if github:
        contact_parts.append(rf"\href{{{latex_escape(github)}}}{{\underline{{{github_display}}}}}")

    contact_line = " $|$ ".join(contact_parts)
    if contact_line:
        contact_line = "    \\small " + contact_line
    else:
        contact_line = "    \\small"

    parts.append(
        "\n".join(
            [
                "\\begin{center}",
                f"    \\textbf{{\\Huge \\scshape {name}}} \\\\ \\vspace{{1pt}}",
                contact_line,
                "\\end{center}",
                "",
            ]
        )
    )

    for section in data.get("sections", []):
        title = latex_escape(section.get("title", ""))
        kind = section.get("kind")
        entries = section.get("entries", [])

        parts.append(f"\\section{{{title}}}")

        if kind in ("education", "experience", "projects"):
            parts.append("  \\resumeSubHeadingListStart")

        if kind == "education":
            for e in entries:
                parts.append("    \\resumeSubheading")
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('school',''))}}}"
                    + rf"{{{latex_escape(e.get('location',''))}}}"
                )
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('degree',''))}}}"
                    + rf"{{{latex_escape(e.get('dates',''))}}}"
                )

                body = e.get("body", [])
                if body:
                    parts.append("      \\resumeItemListStart")
                    parts.append("        " + rf"\resumeItem{{{rich_segments_to_latex(body)}}}")
                    parts.append("      \\resumeItemListEnd")
            parts.append("  \\resumeSubHeadingListEnd")
            parts.append("")
            continue

        if kind == "experience":
            for e in entries:
                parts.append("    \\resumeSubheading")
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('role',''))}}}"
                    + rf"{{{latex_escape(e.get('dates',''))}}}"
                )
                parts.append(
                    "      "
                    + rf"{{{latex_escape(e.get('org',''))}}}"
                    + rf"{{{latex_escape(e.get('location',''))}}}"
                )

                bullets = e.get("bullets", [])
                if bullets:
                    parts.append("      \\resumeItemListStart")
                    for b in bullets:
                        parts.append(
                            "        " + rf"\resumeItem{{{rich_segments_to_latex(b)}}}"
                        )
                    parts.append("      \\resumeItemListEnd")
                    parts.append("")
            parts.append("  \\resumeSubHeadingListEnd")
            parts.append("")
            continue

        if kind == "projects":
            for e in entries:
                title_text = latex_escape(e.get("title", ""))
                stack_text = latex_escape(e.get("stack", ""))
                dates_text = latex_escape(e.get("dates", ""))
                left = rf"\textbf{{{title_text}}} $|$ \emph{{{stack_text}}}"
                parts.append("      \\resumeProjectHeading")
                parts.append("          " + rf"{{{left}}}{{{dates_text}}}")

                bullets = e.get("bullets", [])
                if bullets:
                    parts.append("          \\resumeItemListStart")
                    for b in bullets:
                        parts.append(
                            "            "
                            + rf"\resumeItem{{{rich_segments_to_latex(b)}}}"
                        )
                    parts.append("          \\resumeItemListEnd")
            parts.append("    \\resumeSubHeadingListEnd")
            parts.append("")
            continue

        if kind == "skills":
            parts.append(" \\begin{itemize}[leftmargin=0.15in, label={}]")
            parts.append("    \\small{\\item{")
            for idx, e in enumerate(entries):
                label = latex_escape(e.get("label", ""))
                value = rich_segments_to_latex(e.get("value", []))
                line = rf"     \textbf{{{label}}}{{: {value}}}"
                if idx != len(entries) - 1:
                    line += r" \\\\"
                parts.append(line)
            parts.append("    }}")
            parts.append(" \\end{itemize}")
            parts.append("")
            continue

        for e in entries:
            title_text = latex_escape(e.get("title", ""))
            body_segments = e.get("body", [])
            if title_text:
                parts.append(rf"\textbf{{{title_text}}}\\")
            if body_segments:
                parts.append(rich_segments_to_latex(body_segments) + r"\\")
            parts.append("")

    parts.append(LATEX_END)
    return "\n".join(parts)


def export_latex(data: dict) -> None:
    out = generate_latex(data)
    path = _export_file_path()
    with open(path, "w", encoding="utf-8") as f:
        f.write(out)


@dataclass
class StyleState:
    family: str
    size: int
    bold: bool
    italic: bool
    underline: bool
    fg: Optional[str]
    bg: Optional[str]


class RichTextEditor(ttk.Frame):
    def __init__(self, master, *, default_family: Optional[str] = None, default_size: Optional[int] = None):
        super().__init__(master)

        try:
            df = tkfont.nametofont("TkDefaultFont")
            sys_family = df.cget("family")
            sys_size = int(df.cget("size"))
        except Exception:
            sys_family = "Arial"
            sys_size = 10

        self.default_family = default_family or sys_family
        self.default_size = int(default_size or sys_size)
        self.font_cache: Dict[Tuple[Any, ...], tkfont.Font] = {}

        toolbar = ttk.Frame(self)
        toolbar.pack(fill=tk.X, padx=6, pady=(6, 0))

        self.btn_bold = ttk.Button(toolbar, text="B", width=3, command=self.toggle_bold)
        self.btn_italic = ttk.Button(toolbar, text="I", width=3, command=self.toggle_italic)
        self.btn_underline = ttk.Button(toolbar, text="U", width=3, command=self.toggle_underline)

        self.btn_bold.pack(side=tk.LEFT, padx=(0, 4))
        self.btn_italic.pack(side=tk.LEFT, padx=(0, 4))
        self.btn_underline.pack(side=tk.LEFT, padx=(0, 8))

        families = sorted(
            {
                self.default_family,
                "Arial",
                "Helvetica",
                "Times New Roman",
                "Courier New",
                "DejaVu Sans",
            }
        )
        self.font_family_var = tk.StringVar(value=self.default_family)
        self.font_family_combo = ttk.Combobox(
            toolbar,
            textvariable=self.font_family_var,
            values=families,
            width=18,
            state="readonly",
        )
        self.font_family_combo.pack(side=tk.LEFT, padx=(0, 6))
        self.font_family_combo.bind("<<ComboboxSelected>>", lambda _e: self.apply_font_family())

        self.font_size_var = tk.IntVar(value=self.default_size)
        self.font_size_spin = ttk.Spinbox(
            toolbar,
            from_=8,
            to=48,
            textvariable=self.font_size_var,
            width=5,
            command=self.apply_font_size,
        )
        self.font_size_spin.pack(side=tk.LEFT, padx=(0, 6))

        ttk.Button(toolbar, text="Text Color", command=self.choose_fg).pack(
            side=tk.LEFT, padx=(0, 6)
        )
        ttk.Button(toolbar, text="Highlight", command=self.choose_bg).pack(
            side=tk.LEFT, padx=(0, 6)
        )

        text_frame = ttk.Frame(self)
        text_frame.pack(fill=tk.BOTH, expand=True, padx=6, pady=6)

        text_scroll = ttk.Scrollbar(text_frame, orient=tk.VERTICAL)
        text_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self._last_selection: Optional[Tuple[str, str]] = None
        self.text = tk.Text(
            text_frame,
            wrap="word",
            undo=True,
            yscrollcommand=text_scroll.set,
            exportselection=False,
        )
        self.text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        text_scroll.config(command=self.text.yview)

        self.text.bind("<ButtonRelease-1>", lambda _e: self._capture_selection())
        self.text.bind("<B1-Motion>", lambda _e: self._capture_selection())
        self.text.bind("<KeyRelease>", lambda _e: self._capture_selection())

        mgr = _find_spellcheck_manager(self)
        if mgr is not None:
            try:
                mgr.register_rich_text_editor(self)
            except Exception:
                pass

    def _mark_modified(self) -> None:
        try:
            self.text.edit_modified(True)
        except Exception:
            return
        try:
            self.text.event_generate("<<Modified>>")
        except Exception:
            pass

    def _capture_selection(self) -> None:
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
            self._last_selection = (start, end)
        except tk.TclError:
            self._last_selection = None

    def _selection_range(self) -> Optional[Tuple[str, str]]:
        try:
            start = self.text.index("sel.first")
            end = self.text.index("sel.last")
            self._last_selection = (start, end)
            return (start, end)
        except tk.TclError:
            return self._last_selection

    def _index_for_style_lookup(self) -> str:
        sel = self._selection_range()
        if sel:
            return sel[0]
        return self.text.index("insert")

    def _font_key_tag(self, st: StyleState) -> str:
        b = 1 if st.bold else 0
        i = 1 if st.italic else 0
        u = 1 if st.underline else 0
        return f"f:{st.family}|s:{st.size}|b:{b}|i:{i}|u:{u}"

    def _get_or_create_font(self, family: str, size: int, bold: bool, italic: bool, underline: bool) -> tkfont.Font:
        key = (family, size, bold, italic, underline)
        if key in self.font_cache:
            return self.font_cache[key]
        weight = "bold" if bold else "normal"
        slant = "italic" if italic else "roman"
        f = tkfont.Font(family=family, size=size, weight=weight, slant=slant, underline=underline)
        self.font_cache[key] = f
        return f

    def _style_at(self, index: str) -> StyleState:
        tags = set(self.text.tag_names(index))
        family = self.default_family
        size = self.default_size
        bold = False
        italic = False
        underline = False
        fg = None
        bg = None

        for t in tags:
            if t.startswith("f:"):
                try:
                    m = re.fullmatch(r"f:(.*)\|s:(\d+)\|b:(\d)\|i:(\d)\|u:(\d)", t)
                    if m:
                        family = m.group(1)
                        size = int(m.group(2))
                        bold = m.group(3) == "1"
                        italic = m.group(4) == "1"
                        underline = m.group(5) == "1"
                except Exception:
                    pass
            elif t.startswith("fg:"):
                fg = t[3:]
            elif t.startswith("bg:"):
                bg = t[3:]

        return StyleState(family=family, size=size, bold=bold, italic=italic, underline=underline, fg=fg, bg=bg)

    def _remove_tag_prefix_in_range(self, prefix: str, start: str, end: str) -> None:
        for tag in self.text.tag_names():
            if tag.startswith(prefix):
                self.text.tag_remove(tag, start, end)

    def _apply_style(self, *, family=None, size=None, bold=None, italic=None, underline=None, fg=None, bg=None) -> None:
        sel = self._selection_range()
        if not sel:
            return
        start, end = sel

        try:
            if self.text.compare(start, "==", end):
                return
        except Exception:
            pass

        want_font_change = any(v is not None for v in (family, size, bold, italic, underline))
        want_fg_change = fg is not None
        want_bg_change = bg is not None

        runs: List[Tuple[str, str, StyleState]] = []
        idx = start
        run_start = start
        cur_style = self._style_at(start)

        while True:
            try:
                if not self.text.compare(idx, "<", end):
                    break
            except Exception:
                break

            next_idx = self.text.index(f"{idx}+1c")
            try:
                if self.text.compare(next_idx, ">", end):
                    next_idx = end
            except Exception:
                next_idx = end

            if next_idx == end:
                break

            st_next = self._style_at(next_idx)
            if st_next != cur_style:
                runs.append((run_start, next_idx, cur_style))
                run_start = next_idx
                cur_style = st_next

            idx = next_idx

        runs.append((run_start, end, cur_style))

        for rstart, rend, base in runs:
            if want_font_change:
                st = StyleState(
                    family=family if family is not None else base.family,
                    size=size if size is not None else base.size,
                    bold=bold if bold is not None else base.bold,
                    italic=italic if italic is not None else base.italic,
                    underline=underline if underline is not None else base.underline,
                    fg=base.fg,
                    bg=base.bg,
                )
                self._remove_tag_prefix_in_range("f:", rstart, rend)
                font_tag = self._font_key_tag(st)
                fnt = self._get_or_create_font(st.family, st.size, st.bold, st.italic, st.underline)
                self.text.tag_configure(font_tag, font=fnt)
                self.text.tag_add(font_tag, rstart, rend)

            if want_fg_change:
                self._remove_tag_prefix_in_range("fg:", rstart, rend)
                fg_tag = f"fg:{fg}"
                self.text.tag_configure(fg_tag, foreground=fg)
                self.text.tag_add(fg_tag, rstart, rend)

            if want_bg_change:
                self._remove_tag_prefix_in_range("bg:", rstart, rend)
                bg_tag = f"bg:{bg}"
                self.text.tag_configure(bg_tag, background=bg)
                self.text.tag_add(bg_tag, rstart, rend)

        self._mark_modified()

    def toggle_bold(self):
        idx = self._index_for_style_lookup()
        st = self._style_at(idx)
        self._apply_style(bold=not st.bold)

    def toggle_italic(self):
        idx = self._index_for_style_lookup()
        st = self._style_at(idx)
        self._apply_style(italic=not st.italic)

    def toggle_underline(self):
        idx = self._index_for_style_lookup()
        st = self._style_at(idx)
        self._apply_style(underline=not st.underline)

    def apply_font_family(self):
        self._apply_style(family=self.font_family_var.get())

    def apply_font_size(self):
        try:
            s = int(self.font_size_var.get())
        except Exception:
            return
        self._apply_style(size=s)

    def choose_fg(self):
        c = colorchooser.askcolor(title="Choose text color")
        if not c or not c[1]:
            return
        if not self._selection_range():
            try:
                start = self.text.index("insert wordstart")
                end = self.text.index("insert wordend")
                if self.text.compare(start, "<", end):
                    self.text.tag_add("sel", start, end)
                    self._last_selection = (start, end)
            except Exception:
                pass
        self._apply_style(fg=c[1])

    def choose_bg(self):
        c = colorchooser.askcolor(title="Choose highlight color")
        if not c or not c[1]:
            return
        if not self._selection_range():
            try:
                start = self.text.index("insert wordstart")
                end = self.text.index("insert wordend")
                if self.text.compare(start, "<", end):
                    self.text.tag_add("sel", start, end)
                    self._last_selection = (start, end)
            except Exception:
                pass
        self._apply_style(bg=c[1])

    def set_segments(self, segments):
        self.text.delete("1.0", "end")
        self.text.insert("1.0", "")
        if not segments:
            return

        for seg in segments:
            txt = seg.get("text", "")
            if txt == "":
                continue
            start_index = self.text.index("end-1c")
            self.text.insert("end", txt)
            end_index = self.text.index("end-1c")

            st = StyleState(
                family=seg.get("font") or self.default_family,
                size=int(seg.get("size") or self.default_size),
                bold=bool(seg.get("b")),
                italic=bool(seg.get("i")),
                underline=bool(seg.get("u")),
                fg=seg.get("fg"),
                bg=seg.get("bg"),
            )

            font_tag = self._font_key_tag(st)
            f = self._get_or_create_font(st.family, st.size, st.bold, st.italic, st.underline)
            self.text.tag_configure(font_tag, font=f)
            self.text.tag_add(font_tag, start_index, end_index)

            if st.fg:
                fg_tag = f"fg:{st.fg}"
                self.text.tag_configure(fg_tag, foreground=st.fg)
                self.text.tag_add(fg_tag, start_index, end_index)

            if st.bg:
                bg_tag = f"bg:{st.bg}"
                self.text.tag_configure(bg_tag, background=st.bg)
                self.text.tag_add(bg_tag, start_index, end_index)

    def get_segments(self):
        text = self.text.get("1.0", "end-1c")
        if text == "":
            return []

        intervals_add: Dict[int, List[Tuple[str, str]]] = {}
        intervals_remove: Dict[int, List[Tuple[str, str]]] = {}

        def _add_event(offset: int, kind: str, value: str):
            intervals_add.setdefault(offset, []).append((kind, value))

        def _remove_event(offset: int, kind: str, value: str):
            intervals_remove.setdefault(offset, []).append((kind, value))

        def _offset_for_index(idx: str) -> int:
            try:
                c = self.text.count("1.0", idx, "chars")
                if c and c[0] is not None:
                    return int(c[0])
            except Exception:
                pass

            try:
                return len(self.text.get("1.0", idx))
            except Exception:
                return 0

        for tag in self.text.tag_names():
            if tag.startswith("f:") or tag.startswith("fg:") or tag.startswith("bg:"):
                ranges = self.text.tag_ranges(tag)
                for i in range(0, len(ranges), 2):
                    start = str(ranges[i])
                    end = str(ranges[i + 1])
                    start_off = _offset_for_index(start)
                    end_off = _offset_for_index(end)
                    if start_off == end_off:
                        continue
                    if tag.startswith("f:"):
                        _add_event(start_off, "font", tag)
                        _remove_event(end_off, "font", tag)
                    elif tag.startswith("fg:"):
                        _add_event(start_off, "fg", tag[3:])
                        _remove_event(end_off, "fg", tag[3:])
                    elif tag.startswith("bg:"):
                        _add_event(start_off, "bg", tag[3:])
                        _remove_event(end_off, "bg", tag[3:])

        active_font_tag = None
        active_fg = None
        active_bg = None

        def _style_from_font_tag(font_tag: Optional[str]):
            if not font_tag:
                return (self.default_family, self.default_size, False, False, False)
            m = re.fullmatch(r"f:(.*)\|s:(\d+)\|b:(\d)\|i:(\d)\|u:(\d)", font_tag)
            if not m:
                return (self.default_family, self.default_size, False, False, False)
            family = m.group(1)
            size = int(m.group(2))
            b = m.group(3) == "1"
            i = m.group(4) == "1"
            u = m.group(5) == "1"
            return (family, size, b, i, u)

        segments: List[Dict[str, Any]] = []
        cur = {
            "family": None,
            "size": None,
            "b": None,
            "i": None,
            "u": None,
            "fg": None,
            "bg": None,
        }
        cur_text_parts: List[str] = []

        def _flush():
            if not cur_text_parts:
                return
            seg = {"text": "".join(cur_text_parts)}
            if cur["family"] is not None:
                seg["font"] = cur["family"]
            if cur["size"] is not None:
                seg["size"] = cur["size"]
            if cur["b"]:
                seg["b"] = True
            if cur["i"]:
                seg["i"] = True
            if cur["u"]:
                seg["u"] = True
            if cur["fg"]:
                seg["fg"] = cur["fg"]
            if cur["bg"]:
                seg["bg"] = cur["bg"]
            segments.append(seg)
            cur_text_parts.clear()

        def _set_style_from_active():
            family, size, b, i, u = _style_from_font_tag(active_font_tag)
            cur["family"] = family
            cur["size"] = size
            cur["b"] = b
            cur["i"] = i
            cur["u"] = u
            cur["fg"] = active_fg
            cur["bg"] = active_bg

        _set_style_from_active()

        for pos in range(0, len(text)):
            if pos in intervals_add:
                for kind, value in intervals_add[pos]:
                    if kind == "font":
                        active_font_tag = value
                    elif kind == "fg":
                        active_fg = value
                    elif kind == "bg":
                        active_bg = value

            new_family, new_size, new_b, new_i, new_u = _style_from_font_tag(active_font_tag)
            new_style = {
                "family": new_family,
                "size": new_size,
                "b": new_b,
                "i": new_i,
                "u": new_u,
                "fg": active_fg,
                "bg": active_bg,
            }

            if (
                cur_text_parts
                and (
                    cur["family"] != new_style["family"]
                    or cur["size"] != new_style["size"]
                    or cur["b"] != new_style["b"]
                    or cur["i"] != new_style["i"]
                    or cur["u"] != new_style["u"]
                    or cur["fg"] != new_style["fg"]
                    or cur["bg"] != new_style["bg"]
                )
            ):
                _flush()
                cur.update(new_style)

            if not cur_text_parts:
                cur.update(new_style)

            cur_text_parts.append(text[pos])

            if (pos + 1) in intervals_remove:
                for kind, value in intervals_remove[pos + 1]:
                    if kind == "font" and active_font_tag == value:
                        active_font_tag = None
                    elif kind == "fg" and active_fg == value:
                        active_fg = None
                    elif kind == "bg" and active_bg == value:
                        active_bg = None

        _flush()

        compacted = []
        for seg in segments:
            if seg.get("text") == "":
                continue
            compacted.append(seg)
        return compacted


class BulletEditorDialog(tk.Toplevel):
    def __init__(self, master, title: str, bullets: list, *, editor_height=12):
        super().__init__(master)
        self.title(title)
        self.transient(master)
        self.grab_set()

        self.result = None
        self.bullets = deep_copy(bullets or [])

        self.active_index = -1
        self.dirty = False
        self._loading = False
        self._dragging = False
        self._drag_tip_win: Optional[tk.Toplevel] = None
        self._drag_tip_label: Optional[tk.Label] = None

        main = ttk.Frame(self)
        main.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left = ttk.Frame(main)
        left.pack(side=tk.LEFT, fill=tk.Y)

        list_frame = ttk.Frame(left)
        list_frame.pack(fill=tk.BOTH, expand=True)

        list_scroll_v = ttk.Scrollbar(list_frame, orient=tk.VERTICAL)
        list_scroll_v.pack(side=tk.RIGHT, fill=tk.Y)

        list_scroll_h = ttk.Scrollbar(list_frame, orient=tk.HORIZONTAL)
        list_scroll_h.pack(side=tk.BOTTOM, fill=tk.X)

        self.listbox = tk.Listbox(
            list_frame,
            height=editor_height,
            width=38,
            yscrollcommand=list_scroll_v.set,
            xscrollcommand=list_scroll_h.set,
        )
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        list_scroll_v.config(command=self.listbox.yview)
        list_scroll_h.config(command=self.listbox.xview)

        btns = ttk.Frame(left)
        btns.pack(fill=tk.X, pady=(6, 0))
        self.btn_add_bullet = ttk.Button(btns, text="Add", command=self.add_bullet)
        self.btn_add_bullet.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_delete_bullet = ttk.Button(btns, text="Delete", command=self.delete_bullet)
        self.btn_delete_bullet.pack(side=tk.LEFT)

        right = ttk.Frame(main)
        right.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(10, 0))

        self.editor = RichTextEditor(right)
        self.editor.pack(fill=tk.BOTH, expand=True)

        try:
            self.editor.text.edit_modified(False)
        except Exception:
            pass
        self.editor.text.bind("<<Modified>>", self.on_editor_modified)

        action = ttk.Frame(self)
        action.pack(fill=tk.X, padx=10, pady=(0, 10))
        self.btn_save = ttk.Button(action, text="Save", command=self.save_active)
        self.btn_save.pack(side=tk.LEFT)
        ttk.Button(action, text="OK", command=self.on_ok).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(action, text="Cancel", command=self.on_cancel).pack(side=tk.RIGHT)

        self.listbox.bind("<<ListboxSelect>>", self.on_select)
        self.listbox.bind("<ButtonPress-1>", self._drag_start)
        self.listbox.bind("<B1-Motion>", self._drag_motion)
        self.listbox.bind("<ButtonRelease-1>", self._drag_drop)
        self.listbox.bind("<ButtonPress-3>", self._drag_start)
        self.listbox.bind("<B3-Motion>", self._drag_motion)
        self.listbox.bind("<ButtonRelease-3>", self._drag_drop)

        self.refresh_list()
        if self.bullets:
            self.listbox.selection_set(0)
            self.listbox.activate(0)
            self.on_select(None)
        else:
            self._show_no_selection()

        self._update_save_button()
        self._update_bullet_buttons()

        self.minsize(700, 420)
        try:
            self.after(0, lambda: _center_window(self))
        except Exception:
            pass
        self.protocol("WM_DELETE_WINDOW", self.on_cancel)

    def _show_drag_tip(self, text: str, x: int, y: int) -> None:
        try:
            if self._drag_tip_win is None or not self._drag_tip_win.winfo_exists():
                tip = tk.Toplevel(self)
                tip.overrideredirect(True)
                try:
                    tip.attributes("-topmost", True)
                except Exception:
                    pass
                lbl = tk.Label(
                    tip,
                    text=text,
                    bg="#ffffe0",
                    fg="#000000",
                    bd=1,
                    relief="solid",
                    padx=6,
                    pady=3,
                )
                lbl.pack()
                self._drag_tip_win = tip
                self._drag_tip_label = lbl
            else:
                if self._drag_tip_label is not None:
                    self._drag_tip_label.configure(text=text)
            self._drag_tip_win.geometry(f"+{x}+{y}")
        except Exception:
            return

    def _move_drag_tip(self, x: int, y: int) -> None:
        try:
            if self._drag_tip_win is None or not self._drag_tip_win.winfo_exists():
                return
            self._drag_tip_win.geometry(f"+{x}+{y}")
        except Exception:
            return

    def _hide_drag_tip(self) -> None:
        try:
            if self._drag_tip_win is not None and self._drag_tip_win.winfo_exists():
                self._drag_tip_win.destroy()
        except Exception:
            pass
        self._drag_tip_win = None
        self._drag_tip_label = None

    def refresh_list(self):
        self.listbox.delete(0, tk.END)
        for i, b in enumerate(self.bullets):
            preview = rich_segments_to_latex(b)
            if preview.strip() == "":
                preview = "(empty)"
            self.listbox.insert(tk.END, f"Bullet {i + 1}: {preview}")
        self._update_bullet_buttons()

    def _set_editor_enabled(self, enabled: bool) -> None:
        def _apply(w):
            try:
                if isinstance(w, tk.Text):
                    w.configure(state=("normal" if enabled else "disabled"))
                elif isinstance(w, ttk.Combobox):
                    w.configure(state=("readonly" if enabled else "disabled"))
                elif isinstance(w, ttk.Spinbox):
                    w.configure(state=("normal" if enabled else "disabled"))
                elif isinstance(w, (ttk.Button, ttk.Entry, ttk.Scrollbar)):
                    w.configure(state=("normal" if enabled else "disabled"))
            except Exception:
                pass

            for c in w.winfo_children():
                _apply(c)

        _apply(self.editor)

    def _show_no_selection(self) -> None:
        self.active_index = -1
        self.dirty = False
        self._update_save_button()

        self._loading = True
        try:
            self._set_editor_enabled(True)
            self.editor.set_segments([])
            try:
                self.editor.text.edit_modified(False)
            except Exception:
                pass
            self._set_editor_enabled(False)
        finally:
            self._loading = False
        self._update_bullet_buttons()

    def _update_bullet_buttons(self) -> None:
        if hasattr(self, "btn_delete_bullet"):
            self.btn_delete_bullet.configure(state=("normal" if len(self.bullets) > 0 else "disabled"))

    def selected_index(self) -> int:
        try:
            return int(self.listbox.curselection()[0])
        except Exception:
            return -1

    def _update_save_button(self) -> None:
        if hasattr(self, "btn_save"):
            can_save = self.dirty and (0 <= int(self.active_index) < len(self.bullets))
            self.btn_save.configure(state=("normal" if can_save else "disabled"))

    def save_active(self):
        idx = int(self.active_index)
        if not (0 <= idx < len(self.bullets)):
            return
        self.bullets[idx] = self.editor.get_segments()
        self.dirty = False
        self._update_save_button()

        sel = self.selected_index()
        self.refresh_list()
        if 0 <= sel < len(self.bullets):
            self.listbox.selection_set(sel)
            self.listbox.activate(sel)
        else:
            self.listbox.selection_set(idx)
            self.listbox.activate(idx)

    def on_select(self, _event):
        if getattr(self, "_dragging", False):
            return
        idx = self.selected_index()
        if idx < 0:
            if not self.bullets:
                self._show_no_selection()
            return
        idx = max(0, min(idx, len(self.bullets) - 1))

        if self.dirty:
            self.save_active()

        self.active_index = idx
        self._loading = True
        try:
            self._set_editor_enabled(True)
            self.editor.set_segments(self.bullets[idx])
            try:
                self.editor.text.edit_modified(False)
            except Exception:
                pass
        finally:
            self._loading = False

        self.dirty = False
        self._update_save_button()

    def on_editor_modified(self, _event):
        if self._loading:
            try:
                self.editor.text.edit_modified(False)
            except Exception:
                pass
            return

        try:
            if not self.editor.text.edit_modified():
                return
            self.editor.text.edit_modified(False)
        except Exception:
            pass

        self.dirty = True
        self._update_save_button()

    def add_bullet(self):
        if self.dirty:
            self.save_active()
        self.bullets.append([{"text": ""}])
        self.refresh_list()
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(len(self.bullets) - 1)
        self.listbox.activate(len(self.bullets) - 1)
        self.on_select(None)

    def delete_bullet(self):
        if len(self.bullets) <= 0:
            return
        idx = self.selected_index()
        if self.dirty:
            self.save_active()
        if idx < 0:
            idx = int(self.active_index)
        if 0 <= idx < len(self.bullets):
            self.bullets.pop(idx)
        self.refresh_list()

        if not self.bullets:
            self.listbox.selection_clear(0, tk.END)
            self._show_no_selection()
            return

        new_idx = min(max(idx, 0), len(self.bullets) - 1)
        self.listbox.selection_clear(0, tk.END)
        self.listbox.selection_set(new_idx)
        self.listbox.activate(new_idx)
        self.on_select(None)

    def on_ok(self):
        idx = int(self.active_index)
        if not (0 <= idx < len(self.bullets)):
            idx = int(self.selected_index())
        if 0 <= idx < len(self.bullets):
            self.bullets[idx] = self.editor.get_segments()
            self.dirty = False
        self._hide_drag_tip()
        self.result = self.bullets
        self.destroy()

    def on_cancel(self):
        self._hide_drag_tip()
        self.result = None
        self.destroy()

    def _drag_start(self, event):
        self._dragging = False
        try:
            self._drag_from = int(self.listbox.nearest(event.y))
        except Exception:
            self._drag_from = -1

        if 0 <= getattr(self, "_drag_from", -1) < len(self.bullets):
            text = str(self.bullets[self._drag_from])
            text = text.replace("\n", " ").strip()
            if len(text) > 80:
                text = text[:77] + "..."
            self._show_drag_tip(text, event.x_root + 12, event.y_root + 12)
        else:
            self._hide_drag_tip()

    def _drag_motion(self, event):
        self._dragging = True
        self._move_drag_tip(event.x_root + 12, event.y_root + 12)
        try:
            idx = int(self.listbox.nearest(event.y))
        except Exception:
            return
        if idx < 0:
            return
        try:
            self.listbox.selection_clear(0, tk.END)
            self.listbox.selection_set(idx)
            self.listbox.activate(idx)
        except Exception:
            return

    def _drag_drop(self, event):
        try:
            from_idx = int(getattr(self, "_drag_from", -1))
            to_idx = int(self.listbox.nearest(event.y))
        except Exception:
            return

        self._dragging = False
        self._hide_drag_tip()

        if not (0 <= from_idx < len(self.bullets)):
            return
        if not (0 <= to_idx < len(self.bullets)):
            to_idx = len(self.bullets) - 1
        if from_idx == to_idx:
            return

        if self.dirty:
            self.save_active()

        item = self.bullets.pop(from_idx)
        if to_idx > from_idx:
            to_idx -= 1
        self.bullets.insert(to_idx, item)

        self.active_index = to_idx
        self.refresh_list()
        try:
            self.listbox.selection_clear(0, tk.END)
            self.listbox.selection_set(to_idx)
            self.listbox.activate(to_idx)
        except Exception:
            pass
        self.on_select(None)


class PreviewDialog(tk.Toplevel):
    def __init__(self, master, data: dict, latex_source: str):
        super().__init__(master)
        self.title("Preview")
        self.transient(master)

        self.data = data
        self.latex_source = latex_source
        self.font_cache: Dict[Tuple[Any, ...], tkfont.Font] = {}

        root = ttk.Frame(self)
        root.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        nb = ttk.Notebook(root)
        nb.pack(fill=tk.BOTH, expand=True)

        rendered_tab = ttk.Frame(nb)
        latex_tab = ttk.Frame(nb)

        nb.add(rendered_tab, text="Rendered")
        nb.add(latex_tab, text="LaTeX")

        rendered_frame = ttk.Frame(rendered_tab)
        rendered_frame.pack(fill=tk.BOTH, expand=True)

        rendered_scroll = ttk.Scrollbar(rendered_frame)
        rendered_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.rendered_text = tk.Text(
            rendered_frame,
            wrap="word",
            yscrollcommand=rendered_scroll.set,
            padx=16,
            pady=16,
            borderwidth=0,
            highlightthickness=0,
        )
        self.rendered_text.pack(fill=tk.BOTH, expand=True)
        rendered_scroll.config(command=self.rendered_text.yview)

        self._build_rendered_preview()

        latex_text_frame = ttk.Frame(latex_tab)
        latex_text_frame.pack(fill=tk.BOTH, expand=True)

        latex_scroll = ttk.Scrollbar(latex_text_frame)
        latex_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.latex_text = tk.Text(latex_text_frame, wrap="none", yscrollcommand=latex_scroll.set)
        self.latex_text.pack(fill=tk.BOTH, expand=True)
        latex_scroll.config(command=self.latex_text.yview)

        self.latex_text.insert("1.0", latex_source)
        self.latex_text.configure(state="disabled")

        actions = ttk.Frame(self)
        actions.pack(fill=tk.X, padx=10, pady=(0, 10))
        ttk.Button(actions, text="Copy LaTeX Code to Clipboard", command=self.copy_latex_to_clipboard).pack(
            side=tk.LEFT
        )
        ttk.Button(actions, text="Close", command=self.on_close).pack(side=tk.RIGHT)

        self.minsize(900, 600)
        try:
            self.after(0, lambda: _center_window(self))
        except Exception:
            pass
        self.protocol("WM_DELETE_WINDOW", self.on_close)

    def copy_latex_to_clipboard(self):
        try:
            txt = self.latex_text.get("1.0", "end-1c")
        except Exception:
            txt = self.latex_source

        try:
            self.clipboard_clear()
            self.clipboard_append(txt)
            self.update_idletasks()
        except Exception as e:
            messagebox.showerror("Copy Failed", str(e))
            return

        messagebox.showinfo("Copied", "LaTeX copied to clipboard")

    def _get_font(self, family: str, size: int, bold: bool, italic: bool, underline: bool) -> tkfont.Font:
        key = (family, size, bold, italic, underline)
        if key in self.font_cache:
            return self.font_cache[key]
        weight = "bold" if bold else "normal"
        slant = "italic" if italic else "roman"
        f = tkfont.Font(family=family, size=size, weight=weight, slant=slant, underline=underline)
        self.font_cache[key] = f
        return f

    def _insert_segments(self, segments: List[Dict[str, Any]], *, base_family: str, base_size: int):
        for idx, seg in enumerate(segments or []):
            txt = seg.get("text", "")
            if txt == "":
                continue
            family = seg.get("font") or base_family
            size = int(seg.get("size") or base_size)
            bold = bool(seg.get("b"))
            italic = bool(seg.get("i"))
            underline = bool(seg.get("u"))
            fg = seg.get("fg")
            bg = seg.get("bg")

            tag = f"seg_{self.rendered_text.index('end')}__{idx}"
            f = self._get_font(family, size, bold, italic, underline)
            self.rendered_text.tag_configure(tag, font=f)
            if fg:
                self.rendered_text.tag_configure(tag, foreground=fg)
            if bg:
                self.rendered_text.tag_configure(tag, background=bg)

            start = self.rendered_text.index("end-1c")
            self.rendered_text.insert("end", txt)
            end = self.rendered_text.index("end-1c")
            self.rendered_text.tag_add(tag, start, end)

    def _build_rendered_preview(self):
        self.rendered_text.configure(state="normal")
        self.rendered_text.delete("1.0", "end")

        try:
            df = tkfont.nametofont("TkDefaultFont")
            base_family = df.cget("family")
            base_size = int(df.cget("size"))
        except Exception:
            base_family = "Arial"
            base_size = 10

        name_font = self._get_font(base_family, base_size + 10, True, False, False)
        section_font = self._get_font(base_family, base_size + 2, True, False, False)
        subhead_font = self._get_font(base_family, base_size, True, False, False)
        italic_font = self._get_font(base_family, base_size, False, True, False)

        self.rendered_text.tag_configure("center", justify="center")
        self.rendered_text.tag_configure("name", font=name_font, justify="center")
        self.rendered_text.tag_configure("contact", font=self._get_font(base_family, base_size, False, False, False), justify="center")
        self.rendered_text.tag_configure("section", font=section_font)
        self.rendered_text.tag_configure("subhead", font=subhead_font)
        self.rendered_text.tag_configure("italic", font=italic_font)
        self.rendered_text.tag_configure("spacer", spacing1=8, spacing3=8)

        header = self.data.get("header", {})
        name = header.get("name", "")
        phone = header.get("phone", "")
        email = header.get("email", "")
        linkedin_kind = str(header.get("linkedin_kind", "LinkedIn") or "").strip()
        github_kind = str(header.get("github_kind", "GitHub") or "").strip()
        linkedin_text = header.get("linkedin_display", "") or header.get("linkedin", "")
        github_text = header.get("github_display", "") or header.get("github", "")

        contact_parts: List[str] = []
        if phone:
            contact_parts.append(phone)
        if email:
            contact_parts.append(email)
        if linkedin_kind.lower() != "none" and linkedin_text:
            contact_parts.append(linkedin_text)
        if github_kind.lower() != "none" and github_text:
            contact_parts.append(github_text)

        self.rendered_text.insert("end", (name + "\n"), ("name", "center"))
        if contact_parts:
            self.rendered_text.insert("end", (" | ".join(contact_parts) + "\n"), ("contact", "center"))
        self.rendered_text.insert("end", "\n")

        for section in self.data.get("sections", []):
            title = section.get("title", "")
            kind = section.get("kind")
            entries = section.get("entries", [])

            self.rendered_text.insert("end", title + "\n", ("section",))
            self.rendered_text.insert("end", "―" * 60 + "\n")

            if kind == "education":
                for e in entries:
                    school = e.get("school", "")
                    location = e.get("location", "")
                    degree = e.get("degree", "")
                    dates = e.get("dates", "")
                    body = e.get("body", [])

                    line1 = school
                    if location:
                        line1 += "    " + location
                    self.rendered_text.insert("end", line1 + "\n", ("subhead",))

                    line2 = degree
                    if dates:
                        line2 += "    " + dates
                    self.rendered_text.insert("end", line2 + "\n", ("italic",))

                    if body:
                        self.rendered_text.insert("end", "  • ")
                        self._insert_segments(body, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

            elif kind == "experience":
                for e in entries:
                    role = e.get("role", "")
                    dates = e.get("dates", "")
                    org = e.get("org", "")
                    location = e.get("location", "")

                    line1 = role
                    if dates:
                        line1 += "    " + dates
                    self.rendered_text.insert("end", line1 + "\n", ("subhead",))

                    line2 = org
                    if location:
                        line2 += "    " + location
                    self.rendered_text.insert("end", line2 + "\n", ("italic",))

                    for b in e.get("bullets", []) or []:
                        self.rendered_text.insert("end", "  • ")
                        self._insert_segments(b, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

            elif kind == "projects":
                for e in entries:
                    ptitle = e.get("title", "")
                    stack = e.get("stack", "")
                    dates = e.get("dates", "")

                    header_line = ptitle
                    if stack:
                        header_line += " | " + stack
                    if dates:
                        header_line += "    " + dates
                    self.rendered_text.insert("end", header_line + "\n", ("subhead",))

                    for b in e.get("bullets", []) or []:
                        self.rendered_text.insert("end", "  • ")
                        self._insert_segments(b, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

            elif kind == "skills":
                for e in entries:
                    label = e.get("label", "")
                    self.rendered_text.insert("end", f"{label}: ", ("subhead",))
                    self._insert_segments(e.get("value", []), base_family=base_family, base_size=base_size)
                    self.rendered_text.insert("end", "\n")
                self.rendered_text.insert("end", "\n")

            else:
                for e in entries:
                    etitle = e.get("title", "")
                    if etitle:
                        self.rendered_text.insert("end", etitle + "\n", ("subhead",))
                    body = e.get("body", [])
                    if body:
                        self._insert_segments(body, base_family=base_family, base_size=base_size)
                        self.rendered_text.insert("end", "\n")
                    self.rendered_text.insert("end", "\n")

        self.rendered_text.configure(state="disabled")

    def on_close(self):
        self.destroy()


class EntryEditorDialog(tk.Toplevel):
    def __init__(self, master, section_kind: str, entry: Optional[dict]):
        super().__init__(master)
        self.section_kind = section_kind
        self.original = deep_copy(entry) if entry is not None else None
        self.entry = deep_copy(entry) if entry is not None else {}
        self.result = None

        self.title("Edit Entry")
        self.transient(master)
        self.grab_set()
        try:
            self.after(0, lambda: _center_window(self))
        except Exception:
            pass

        root = ttk.Frame(self)
        root.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.fields_frame = ttk.Frame(root)
        self.fields_frame.pack(fill=tk.X)

        self.field_vars: Dict[str, tk.StringVar] = {}

        def add_field(label: str, key: str):
            row = ttk.Frame(self.fields_frame)
            row.pack(fill=tk.X, pady=3)
            ttk.Label(row, text=label, width=18).pack(side=tk.LEFT)
            var = tk.StringVar(value=str(self.entry.get(key, "")))
            self.field_vars[key] = var
            fld = SingleLineTextField(row, textvariable=var)
            fld.pack(side=tk.LEFT, fill=tk.X, expand=True)
            mgr = _find_spellcheck_manager(self)
            if mgr is not None:
                try:
                    mgr.register_text(fld.text)
                except Exception:
                    pass

        if section_kind == "education":
            add_field("School", "school")
            add_field("Location", "location")
            add_field("Degree", "degree")
            add_field("Dates", "dates")
        elif section_kind == "experience":
            add_field("Role", "role")
            add_field("Dates", "dates")
            add_field("Organization", "org")
            add_field("Location", "location")
        elif section_kind == "projects":
            add_field("Title", "title")
            add_field("Stack", "stack")
            add_field("Dates", "dates")
        elif section_kind == "skills":
            add_field("Label", "label")
        else:
            add_field("Title", "title")

        self.body_frame = ttk.Frame(root)
        self.body_frame.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

        self.editor = None
        self.bullets_button = None

        if section_kind in ("experience", "projects"):
            self.bullets = deep_copy(self.entry.get("bullets", []))
            if not self.bullets:
                self.bullets = [[{"text": ""}]]

            self.bullets_button = ttk.Button(
                self.body_frame,
                text=f"Edit Bullets ({len(self.bullets)})",
                command=self.edit_bullets,
            )
            self.bullets_button.pack(anchor="w")

        elif section_kind == "skills":
            ttk.Label(self.body_frame, text="Value").pack(anchor="w")
            self.editor = RichTextEditor(self.body_frame)
            self.editor.pack(fill=tk.BOTH, expand=True)
            self.editor.set_segments(self.entry.get("value", []))

        else:
            ttk.Label(self.body_frame, text="Body").pack(anchor="w")
            self.editor = RichTextEditor(self.body_frame)
            self.editor.pack(fill=tk.BOTH, expand=True)
            self.editor.set_segments(self.entry.get("body", []))

        actions = ttk.Frame(self)
        actions.pack(fill=tk.X, padx=10, pady=(0, 10))
        ttk.Button(actions, text="Cancel", command=self.on_cancel).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(actions, text="OK", command=self.on_ok).pack(side=tk.RIGHT)

        self.minsize(720, 420)
        self.protocol("WM_DELETE_WINDOW", self.on_cancel)

    def edit_bullets(self):
        dlg = BulletEditorDialog(self, "Edit Bullets", self.bullets)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self.bullets = dlg.result
        if self.bullets_button is not None:
            self.bullets_button.configure(text=f"Edit Bullets ({len(self.bullets)})")

    def on_ok(self):
        out = deep_copy(self.entry)
        for k, v in self.field_vars.items():
            out[k] = v.get().strip()

        if self.section_kind in ("experience", "projects"):
            out["bullets"] = deep_copy(self.bullets)
        elif self.section_kind == "skills":
            out["value"] = self.editor.get_segments() if self.editor else []
        else:
            out["body"] = self.editor.get_segments() if self.editor else []

        self.result = out
        self.destroy()

    def on_cancel(self):
        self.result = None
        self.destroy()


class ResumeApp(ttk.Frame):
    def __init__(self, master):
        super().__init__(master)
        self.master = master
        self.data = deep_copy(DEFAULT_DATA)

        self.current_project_path: Optional[str] = None
        self._file_menu_save_index: Optional[int] = None
        self._file_menu: Optional[tk.Menu] = None

        self.spellcheck_manager = SpellCheckManager(self.master)
        self.master.spellcheck_manager = self.spellcheck_manager
        self.var_spellcheck = tk.BooleanVar(value=True)

        self._drag_section_iid: Optional[str] = None
        self._drag_entry_iid: Optional[str] = None
        self._drag_tip_win: Optional[tk.Toplevel] = None
        self._drag_tip_label: Optional[tk.Label] = None

        self.undo_stack: List[dict] = []
        self.redo_stack: List[dict] = []
        self._suspend_undo = False
        self._header_typing_in_progress = False
        self._header_commit_job = None

        self.pack(fill=tk.BOTH, expand=True)

        self.spellcheck_manager.set_enabled(True)

        self._build_ui()
        self._ensure_help_docs_written()
        self._bind_undo_redo_shortcuts()
        self._update_undo_redo_buttons()
        self._update_save_project_buttons()
        self.refresh_sections()

    def _open_spellcheck_data_editor(self) -> None:
        def _apply(words: List[str]) -> None:
            try:
                self.spellcheck_manager.set_ignore_all(words)
            except Exception:
                return

        try:
            words = self.spellcheck_manager.get_ignore_all()
        except Exception:
            words = []

        dlg = SpellcheckIgnoreAllDialog(self.master, words=words, on_apply=_apply)
        try:
            dlg.grab_set()
        except Exception:
            pass

    def _ensure_help_docs_written(self) -> None:
        try:
            out_dir = _help_docs_dir()
            os.makedirs(out_dir, exist_ok=True)
        except Exception:
            return

        try:
            with open(os.path.join(out_dir, "Quick Start Guide.txt"), "w", encoding="utf-8") as f:
                f.write(str(QUICK_START_GUIDE_TEXT))
        except Exception:
            pass

        try:
            with open(os.path.join(out_dir, "Detailed Help.txt"), "w", encoding="utf-8") as f:
                f.write(str(DETAILED_HELP_TEXT))
        except Exception:
            pass

    def _open_quick_start_guide(self) -> None:
        self._ensure_help_docs_written()
        HelpDialog(self.master, title="Quick Start Guide", content=QUICK_START_GUIDE_TEXT)

    def _open_detailed_help(self) -> None:
        self._ensure_help_docs_written()
        HelpDialog(self.master, title="Detailed Help", content=DETAILED_HELP_TEXT)

    def _normalize_loaded_state(self, loaded: Any) -> dict:
        if not isinstance(loaded, dict):
            raise ValueError("Invalid file format (expected a JSON object at the top level).")

        state = deep_copy(DEFAULT_DATA)

        def _norm_key(k: Any) -> str:
            try:
                s = str(k)
            except Exception:
                return ""
            s = s.strip().lower()
            s = s.replace("_", " ")
            s = re.sub(r"\s+", " ", s)
            return s

        header_key_map = {
            "name": "name",
            "phone": "phone",
            "email": "email",
            "linkedin kind": "linkedin_kind",
            "linkedin_kind": "linkedin_kind",
            "linkedin type": "linkedin_kind",
            "linkedin_type": "linkedin_kind",
            "linkedin": "linkedin",
            "linked in": "linkedin",
            "li": "linkedin",
            "li url": "linkedin",
            "li link": "linkedin",
            "linkedin url": "linkedin",
            "linkedin link": "linkedin",
            "li text": "linkedin_display",
            "linkedin text": "linkedin_display",
            "linkedin display": "linkedin_display",
            "linkedin display text": "linkedin_display",
            "linkedin_display": "linkedin_display",
            "github kind": "github_kind",
            "github_kind": "github_kind",
            "github type": "github_kind",
            "github_type": "github_kind",
            "github": "github",
            "git hub": "github",
            "gh": "github",
            "gh url": "github",
            "gh link": "github",
            "github url": "github",
            "github link": "github",
            "gh text": "github_display",
            "github text": "github_display",
            "github display": "github_display",
            "github display text": "github_display",
            "github_display": "github_display",
        }

        def _apply_header_values(src: Any) -> None:
            if not isinstance(src, dict):
                return
            for k, v in src.items():
                nk = _norm_key(k)
                dest = header_key_map.get(nk)
                if dest is None and isinstance(k, str) and k in state.get("header", {}):
                    dest = k
                if not dest:
                    continue
                state["header"][dest] = str(v) if v is not None else ""

        header = loaded.get("header")
        _apply_header_values(header)
        if isinstance(loaded.get("data"), dict):
            _apply_header_values(loaded.get("data").get("header"))
        _apply_header_values(loaded)

        sections_any: Any = loaded.get("sections")
        sections_list: List[Any] = []
        if isinstance(sections_any, list):
            sections_list = sections_any
        elif isinstance(sections_any, dict):
            sections_list = list(sections_any.values())

        if not sections_list and isinstance(loaded.get("data"), dict):
            data_wrap = loaded.get("data")
            sections_any = data_wrap.get("sections")
            if isinstance(sections_any, list):
                sections_list = sections_any
            elif isinstance(sections_any, dict):
                sections_list = list(sections_any.values())

        defaults_by_id = {
            s.get("id"): s
            for s in deep_copy(DEFAULT_DATA).get("sections", [])
            if isinstance(s, dict) and s.get("id")
        }

        if not sections_list:
            inferred_sections: List[dict] = []
            for sec_id, default_sec in defaults_by_id.items():
                if sec_id in loaded:
                    v = loaded.get(sec_id)
                    if isinstance(v, dict):
                        s = deep_copy(v)
                        if not isinstance(s.get("entries"), list):
                            s["entries"] = []
                        inferred_sections.append(s)
                    elif isinstance(v, list):
                        s = deep_copy(default_sec)
                        s["entries"] = v
                        inferred_sections.append(s)
            if inferred_sections:
                sections_list = inferred_sections

        def _coerce_segments(val: Any) -> List[dict]:
            if val is None:
                return []
            if isinstance(val, str):
                return [{"text": val}]
            if isinstance(val, dict) and "text" in val:
                return [deep_copy(val)]
            if isinstance(val, list):
                if all(isinstance(x, dict) and "text" in x for x in val):
                    return deep_copy(val)
                if all(isinstance(x, str) for x in val):
                    return [{"text": "\n".join(val)}]
            return []

        def _coerce_bullets(val: Any) -> List[List[dict]]:
            if val is None:
                return []
            if isinstance(val, str):
                return [[{"text": val}]]
            if isinstance(val, list):
                if not val:
                    return []
                if all(isinstance(b, list) for b in val):
                    out: List[List[dict]] = []
                    for b in val:
                        segs = _coerce_segments(b)
                        if segs:
                            out.append(segs)
                        else:
                            out.append([{"text": ""}])
                    return out
                if all(isinstance(b, str) for b in val):
                    return [[{"text": b}] for b in val]
                if all(isinstance(b, dict) for b in val):
                    segs = _coerce_segments(val)
                    return [segs] if segs else []
            return []

        if sections_list:
            normalized_sections: List[dict] = []
            for sec in sections_list:
                if not isinstance(sec, dict):
                    continue
                s = deep_copy(sec)
                sec_id = s.get("id")
                if sec_id in defaults_by_id:
                    default_sec = defaults_by_id[sec_id]
                    for k in ("title", "kind"):
                        if not s.get(k) and default_sec.get(k):
                            s[k] = default_sec.get(k)

                if not isinstance(s.get("entries"), list):
                    s["entries"] = []

                kind = s.get("kind")
                normalized_entries: List[dict] = []
                for e in s.get("entries", []) or []:
                    if not isinstance(e, dict):
                        continue
                    ee = deep_copy(e)
                    if kind == "education":
                        ee["body"] = _coerce_segments(ee.get("body"))
                    elif kind in ("experience", "projects"):
                        if "bullets" in ee:
                            ee["bullets"] = _coerce_bullets(ee.get("bullets"))
                        elif "body" in ee:
                            body_segs = _coerce_segments(ee.get("body"))
                            ee["bullets"] = [body_segs] if body_segs else []
                            ee.pop("body", None)
                        else:
                            ee["bullets"] = _coerce_bullets([])
                    elif kind == "skills":
                        ee["value"] = _coerce_segments(ee.get("value"))
                    else:
                        ee["body"] = _coerce_segments(ee.get("body"))

                    normalized_entries.append(ee)

                s["entries"] = normalized_entries
                normalized_sections.append(s)
            if normalized_sections:
                state["sections"] = normalized_sections

        return state

    def _build_ui(self):
        self.master.title(f"{APP_NAME} v.{APP_VERSION}")
        self.master.minsize(980, 600)

        topbar = ttk.Frame(self)
        topbar.pack(fill=tk.X, padx=10, pady=10)

        topbar_row1 = ttk.Frame(topbar)
        topbar_row1.pack(fill=tk.X)

        topbar_row2 = ttk.Frame(topbar)
        topbar_row2.pack(fill=tk.X, pady=(6, 0))

        file_btn = ttk.Menubutton(topbar_row1, text="File")
        file_menu = tk.Menu(file_btn, tearoff=False)
        self._file_menu = file_menu

        file_menu.add_command(label="Save", command=self.save_project, state="disabled")
        self._file_menu_save_index = file_menu.index("end")
        file_menu.add_command(label="Save project as .json...", command=self.save_as)
        file_menu.add_command(label="Load project from .json file...", command=self.load_from_file)
        file_menu.add_separator()
        file_menu.add_command(label="Export LaTeX as .tex...", command=self.export_as_tex)
        file_menu.add_command(label="Export as Word Document...", command=self.export_as_docx)
        file_menu.add_command(label="Export as PDF...", command=self.export_as_pdf)
        file_btn.configure(menu=file_menu)
        file_btn.pack(side=tk.LEFT, padx=(0, 6))

        edit_btn = ttk.Menubutton(topbar_row1, text="Edit")
        edit_menu = tk.Menu(edit_btn, tearoff=False)
        edit_menu.add_command(label="Undo", command=self.undo)
        edit_menu.add_command(label="Redo", command=self.redo)
        edit_menu.add_separator()
        edit_menu.add_command(label="Delete all", command=self.delete_all)
        edit_btn.configure(menu=edit_menu)
        edit_btn.pack(side=tk.LEFT, padx=(0, 6))

        settings_btn = ttk.Menubutton(topbar_row1, text="Settings")
        settings_menu = tk.Menu(settings_btn, tearoff=False)

        def _toggle_spellcheck():
            enabled = bool(self.var_spellcheck.get())
            if enabled and not self.spellcheck_manager._ensure_spellchecker():
                messagebox.showerror(
                    "Spellcheck Unavailable",
                    "Spellcheck requires the 'pyspellchecker' package.\n\nInstall it with: pip install pyspellchecker",
                )
                self.var_spellcheck.set(False)
                enabled = False

            self.spellcheck_manager.set_enabled(enabled)

        settings_menu.add_checkbutton(
            label="Enable Spellcheck",
            variable=self.var_spellcheck,
            command=_toggle_spellcheck,
        )

        settings_menu.add_command(
            label="Edit Document Spellcheck Data...",
            command=self._open_spellcheck_data_editor,
        )

        settings_btn.configure(menu=settings_menu)
        settings_btn.pack(side=tk.LEFT, padx=(0, 6))

        help_btn = ttk.Menubutton(topbar_row1, text="Help")
        help_menu = tk.Menu(help_btn, tearoff=False)
        help_menu.add_command(label="Quick Start Guide", command=self._open_quick_start_guide)
        help_menu.add_command(label="Detailed Help", command=self._open_detailed_help)
        help_btn.configure(menu=help_menu)
        help_btn.pack(side=tk.LEFT, padx=(0, 6))

        self.btn_save_project = ttk.Button(topbar_row2, text="💾", command=self.save_project, state="disabled")
        self.btn_save_project.pack(side=tk.LEFT, padx=(0, 6))
        ToolTip(self.btn_save_project, "Save (Ctrl+S)")
        self.btn_undo = ttk.Button(topbar_row2, text="Undo", command=self.undo)
        self.btn_undo.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_redo = ttk.Button(topbar_row2, text="Redo", command=self.redo)
        self.btn_redo.pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(topbar_row2, text="Delete all", command=self.delete_all).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(topbar_row2, text="Preview", command=self.preview).pack(side=tk.LEFT, padx=(0, 6))

        ttk.Separator(topbar_row2, orient=tk.VERTICAL).pack(side=tk.LEFT, fill=tk.Y, padx=8)

        ttk.Button(topbar_row2, text="Demo mode", command=self.demo_mode).pack(
            side=tk.LEFT, padx=(0, 6)
        )

        outer = ttk.Panedwindow(self, orient=tk.HORIZONTAL)
        outer.pack(fill=tk.BOTH, expand=True, padx=10, pady=(0, 10))

        left = ttk.Frame(outer)
        right = ttk.Frame(outer)
        outer.add(left, weight=1)
        outer.add(right, weight=3)

        header = ttk.Labelframe(right, text="Header")
        header.pack(fill=tk.X, pady=(0, 10))

        self.header_vars: Dict[str, tk.StringVar] = {}
        self._header_widgets: Dict[str, Any] = {}
        self._header_last_kind: Dict[str, str] = {}
        self._header_label_vars: Dict[str, tk.StringVar] = {}

        header_left_col_chars = 20
        header_left_col_padx = 6
        header_left_col_minsize = 160
        try:
            _f = tkfont.nametofont("TkDefaultFont")
            header_left_col_minsize = int(_f.measure("0" * header_left_col_chars)) + 24 + header_left_col_padx
        except Exception:
            pass

        try:
            tmp_cb = ttk.Combobox(header, values=("X (Twitter)",), state="readonly", width=header_left_col_chars)
            tmp_cb.pack()
            try:
                self.master.update_idletasks()
            except Exception:
                pass
            cb_px = int(tmp_cb.winfo_reqwidth() or 0)
            try:
                tmp_cb.destroy()
            except Exception:
                pass
            if cb_px > 0:
                header_left_col_minsize = max(int(header_left_col_minsize), cb_px + header_left_col_padx)
        except Exception:
            pass

        linkedin_kind_values = (
            "None",
            "Custom",
            "LinkedIn",
            "Facebook",
            "Instagram",
            "TikTok",
            "Snapchat",
            "X (Twitter)",
            "Slack",
            "Discord",
            "Indeed",
            "Monster",
            "Glassdoor",
            "Handshake",
            "Wellfound",
            "Stack Overflow",
        )
        github_kind_values = ("None", "Custom", "GitHub", "Portfolio", "Webpage")

        def add_header_field(row_parent, label, key):
            row = ttk.Frame(row_parent)
            row.pack(fill=tk.X, padx=8, pady=4)
            row.columnconfigure(1, weight=1)
            row.columnconfigure(0, minsize=header_left_col_minsize)

            if key in ("linkedin_display", "github_display"):
                lbl_var = tk.StringVar(value=str(label))
                self._header_label_vars[key] = lbl_var
                ttk.Label(row, textvariable=lbl_var, width=header_left_col_chars, anchor="w").grid(
                    row=0, column=0, sticky="w", padx=(0, header_left_col_padx)
                )
            else:
                ttk.Label(row, text=label, width=header_left_col_chars, anchor="w").grid(
                    row=0, column=0, sticky="w", padx=(0, header_left_col_padx)
                )
            var = tk.StringVar(value=str(self.data.get("header", {}).get(key, "")))
            self.header_vars[key] = var
            excluded_keys = {
                "name",
                "phone",
                "email",
                "linkedin",
                "linkedin_display",
                "github",
                "github_display",
            }

            if key in excluded_keys:
                ent = ttk.Entry(row, textvariable=var)
                ent.grid(row=0, column=1, sticky="ew")
                self._header_widgets[key] = ent
            else:
                fld = SingleLineTextField(row, textvariable=var)
                fld.grid(row=0, column=1, sticky="ew")
                self._header_widgets[key] = fld
                try:
                    self.spellcheck_manager.register_text(fld.text)
                except Exception:
                    pass

            var.trace_add("write", lambda *_a: self._on_header_var_changed())

        def add_header_link_field(row_parent, kind_key: str, url_key: str, kind_values: Tuple[str, ...], default_kind: str):
            row = ttk.Frame(row_parent)
            row.pack(fill=tk.X, padx=8, pady=4)
            row.columnconfigure(1, weight=1)
            row.columnconfigure(0, minsize=header_left_col_minsize)

            kind_var = tk.StringVar(value=str(self.data.get("header", {}).get(kind_key, default_kind)))
            self.header_vars[kind_key] = kind_var
            self._header_last_kind[kind_key] = str(kind_var.get() or "")
            cb = ttk.Combobox(
                row,
                textvariable=kind_var,
                values=kind_values,
                state="readonly",
                width=header_left_col_chars,
            )
            cb.grid(row=0, column=0, sticky="w", padx=(0, header_left_col_padx))
            self._header_widgets[kind_key] = cb

            var = tk.StringVar(value=str(self.data.get("header", {}).get(url_key, "")))
            self.header_vars[url_key] = var
            ent = ttk.Entry(row, textvariable=var)
            ent.grid(row=0, column=1, sticky="ew")
            self._header_widgets[url_key] = ent

            def _on_kind_selected(_event: Any = None) -> None:
                selected = str(kind_var.get() or "").strip()
                prev = str(self._header_last_kind.get(kind_key, "") or "").strip()

                if selected.lower() == "custom":
                    prompt_title = "Custom"
                    prompt_label = "Enter label:"
                    if kind_key == "linkedin_kind":
                        prompt_title = "Custom LinkedIn Label"
                        prompt_label = "Enter label for this link:"
                    elif kind_key == "github_kind":
                        prompt_title = "Custom GitHub Label"
                        prompt_label = "Enter label for this link:"

                    custom = simpledialog.askstring(prompt_title, prompt_label, parent=self)
                    custom = "" if custom is None else str(custom).strip()
                    if not custom:
                        kind_var.set(prev)
                        return
                    else:
                        kind_var.set(custom)
                        selected = custom

                        display_key = ""
                        if kind_key == "linkedin_kind":
                            display_key = "linkedin_display"
                        elif kind_key == "github_kind":
                            display_key = "github_display"

                        if display_key:
                            dv = self.header_vars.get(display_key)
                            if dv is not None:
                                was_suspend = self._suspend_undo
                                self._suspend_undo = True
                                try:
                                    dv.set(custom)
                                finally:
                                    self._suspend_undo = was_suspend

                if selected == prev:
                    self._update_header_link_field_states()
                    return

                self._header_last_kind[kind_key] = selected
                self._update_header_link_field_states()
                self._update_header_platform_text_labels()
                self._on_header_var_changed()

            cb.bind("<<ComboboxSelected>>", _on_kind_selected)
            var.trace_add("write", lambda *_a: self._on_header_var_changed())

        add_header_field(header, "Name", "name")
        add_header_field(header, "Phone", "phone")
        add_header_field(header, "Email", "email")
        add_header_link_field(header, "linkedin_kind", "linkedin", linkedin_kind_values, "LinkedIn")
        add_header_field(header, "LI Text", "linkedin_display")
        add_header_link_field(header, "github_kind", "github", github_kind_values, "GitHub")
        add_header_field(header, "GH Text", "github_display")

        self._update_header_link_field_states()
        self._update_header_platform_text_labels()

        sections_frame = ttk.Labelframe(left, text="Sections")
        sections_frame.pack(fill=tk.BOTH, expand=True)

        sec_list_frame = ttk.Frame(sections_frame)
        sec_list_frame.pack(fill=tk.BOTH, expand=True, padx=8, pady=(8, 0))

        sec_scroll = ttk.Scrollbar(sec_list_frame, orient=tk.VERTICAL)
        sec_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.sections_tree = ttk.Treeview(
            sec_list_frame,
            columns=("title", "kind"),
            show="headings",
            selectmode="browse",
            height=12,
            yscrollcommand=sec_scroll.set,
        )
        self.sections_tree.heading("title", text="Section")
        self.sections_tree.heading("kind", text="Type")
        self.sections_tree.column("title", width=220, stretch=True)
        self.sections_tree.column("kind", width=90, stretch=False)
        self.sections_tree.pack(fill=tk.BOTH, expand=True)
        sec_scroll.config(command=self.sections_tree.yview)
        self.sections_tree.bind("<<TreeviewSelect>>", self.on_section_selected)
        self.sections_tree.bind("<ButtonPress-1>", self._on_sections_drag_start, add="+")
        self.sections_tree.bind("<B1-Motion>", self._on_sections_drag_motion, add="+")
        self.sections_tree.bind("<ButtonRelease-1>", self._on_sections_drag_drop, add="+")
        self.sections_tree.bind("<ButtonPress-3>", self._on_sections_drag_start, add="+")
        self.sections_tree.bind("<B3-Motion>", self._on_sections_drag_motion, add="+")
        self.sections_tree.bind("<ButtonRelease-3>", self._on_sections_drag_drop, add="+")

        sec_btns = ttk.Frame(sections_frame)
        sec_btns.pack(fill=tk.X, padx=8, pady=8)
        ttk.Button(sec_btns, text="Add", command=self.add_section).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(sec_btns, text="Rename", command=self.rename_section).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(sec_btns, text="Delete", command=self.delete_section).pack(side=tk.LEFT)

        entries_frame = ttk.Labelframe(right, text="Entries")
        entries_frame.pack(fill=tk.BOTH, expand=True)

        ent_list_frame = ttk.Frame(entries_frame)
        entries_frame.columnconfigure(0, weight=1)
        entries_frame.rowconfigure(0, weight=1)
        entries_frame.rowconfigure(1, weight=0)
        ent_list_frame.grid(row=0, column=0, sticky="nsew", padx=8, pady=(8, 0))

        ent_scroll = ttk.Scrollbar(ent_list_frame, orient=tk.VERTICAL)
        ent_scroll.pack(side=tk.RIGHT, fill=tk.Y)

        self.entries_tree = ttk.Treeview(
            ent_list_frame,
            columns=("entry",),
            show="headings",
            selectmode="browse",
            height=14,
            yscrollcommand=ent_scroll.set,
        )
        self.entries_tree.heading("entry", text="Entry")
        self.entries_tree.column("entry", width=640, stretch=True)
        self.entries_tree.pack(fill=tk.BOTH, expand=True)
        ent_scroll.config(command=self.entries_tree.yview)

        self.entries_tree.bind("<Double-Button-1>", lambda _e: self.edit_entry())
        self.entries_tree.bind("<Return>", lambda _e: self.edit_entry())
        self.entries_tree.bind("<Delete>", lambda _e: self.delete_entry())
        self.entries_tree.bind("<<TreeviewSelect>>", lambda _e: self._update_entry_action_buttons())
        self.entries_tree.bind("<ButtonPress-1>", self._on_entries_drag_start, add="+")
        self.entries_tree.bind("<B1-Motion>", self._on_entries_drag_motion, add="+")
        self.entries_tree.bind("<ButtonRelease-1>", self._on_entries_drag_drop, add="+")
        self.entries_tree.bind("<ButtonPress-3>", self._on_entries_drag_start, add="+")
        self.entries_tree.bind("<B3-Motion>", self._on_entries_drag_motion, add="+")
        self.entries_tree.bind("<ButtonRelease-3>", self._on_entries_drag_drop, add="+")

        entry_btns = ttk.Frame(entries_frame)
        entry_btns.grid(row=1, column=0, sticky="ew", padx=8, pady=8)

        self.btn_entry_add = ttk.Button(entry_btns, text="Add", command=self.add_entry)
        self.btn_entry_add.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_entry_edit = ttk.Button(entry_btns, text="Edit", command=self.edit_entry)
        self.btn_entry_edit.pack(side=tk.LEFT, padx=(0, 6))
        self.btn_entry_delete = ttk.Button(entry_btns, text="Delete", command=self.delete_entry)
        self.btn_entry_delete.pack(side=tk.LEFT)

        self._update_entry_action_buttons()

    def _selected_section_index(self) -> int:
        try:
            sel = self.sections_tree.selection()
            if not sel:
                return -1
            return int(sel[0])
        except Exception:
            return -1

    def _selected_entry_index(self) -> int:
        try:
            sel = self.entries_tree.selection()
            if not sel:
                return -1
            return int(sel[0])
        except Exception:
            return -1

    def refresh_sections(self):
        prev_idx = -1
        try:
            sel = self.sections_tree.selection()
            if sel:
                prev_idx = int(sel[0])
        except Exception:
            prev_idx = -1

        for iid in self.sections_tree.get_children(""):
            self.sections_tree.delete(iid)

        sections = self.data.get("sections", [])
        for idx, sec in enumerate(sections):
            kind = sec.get("kind", "")
            title = sec.get("title", "")
            self.sections_tree.insert("", "end", iid=str(idx), values=(title, kind))

        if sections:
            target = -1
            if 0 <= prev_idx < len(sections):
                target = prev_idx
            else:
                for i, sec in enumerate(sections):
                    if sec.get("entries"):
                        target = i
                        break
            if target < 0:
                target = 0

            if not sections[target].get("entries"):
                for i, sec in enumerate(sections):
                    if sec.get("entries"):
                        target = i
                        break

            self.sections_tree.selection_set(str(target))
            self.on_section_selected(None)
        else:
            for iid in self.entries_tree.get_children(""):
                self.entries_tree.delete(iid)
            self._update_entry_action_buttons()

    def refresh_entries(self):
        for iid in self.entries_tree.get_children(""):
            self.entries_tree.delete(iid)
        sidx = self._selected_section_index()
        if sidx < 0:
            self._update_entry_action_buttons()
            return
        sec = self.data.get("sections", [])[sidx]
        kind = sec.get("kind")
        entries = sec.get("entries", [])

        for idx, e in enumerate(entries):
            self.entries_tree.insert(
                "", "end", iid=str(idx), values=(self._entry_summary(kind, e),)
            )

        if self.entries_tree.get_children(""):
            sel = self.entries_tree.selection()
            if not sel:
                self.entries_tree.selection_set("0")

        self._update_entry_action_buttons()

    def _update_entry_action_buttons(self) -> None:
        if not hasattr(self, "btn_entry_edit") or not hasattr(self, "btn_entry_delete"):
            return

        has_entries = len(self.entries_tree.get_children("")) > 0
        has_selection = self._selected_entry_index() >= 0

        self.btn_entry_edit.configure(state=("normal" if has_selection else "disabled"))
        self.btn_entry_delete.configure(state=("normal" if has_entries else "disabled"))

    def _entry_summary(self, kind: str, e: dict) -> str:
        if kind == "education":
            return f"{e.get('school','')} - {e.get('degree','')}"
        if kind == "experience":
            return f"{e.get('role','')} - {e.get('org','')}"
        if kind == "projects":
            return f"{e.get('title','')} - {e.get('stack','')}"
        if kind == "skills":
            return f"{e.get('label','')}"
        return e.get("title", "(entry)")

    def on_section_selected(self, _event):
        self.refresh_entries()

    def _move_item_in_list(self, items: list, from_idx: int, to_idx: int) -> int:
        if not (0 <= from_idx < len(items)):
            return from_idx
        if not (0 <= to_idx < len(items)):
            to_idx = len(items) - 1
        if from_idx == to_idx:
            return from_idx
        item = items.pop(from_idx)
        if to_idx > from_idx:
            to_idx -= 1
        items.insert(to_idx, item)
        return to_idx

    def _on_sections_drag_start(self, event):
        self._hide_drag_tip()
        try:
            if self.sections_tree.identify_region(event.x, event.y) != "cell":
                self._drag_section_iid = None
                return
            iid = self.sections_tree.identify_row(event.y)
            self._drag_section_iid = iid if iid else None
        except Exception:
            self._drag_section_iid = None

        if self._drag_section_iid is not None:
            try:
                idx = int(self._drag_section_iid)
                sections = self.data.get("sections", [])
                if 0 <= idx < len(sections):
                    title = str(sections[idx].get("title", ""))
                    kind = str(sections[idx].get("kind", ""))
                    text = title if not kind else f"{title} ({kind})"
                    text = text.replace("\n", " ").strip()
                    if len(text) > 80:
                        text = text[:77] + "..."
                    self._show_drag_tip(text, event.x_root + 12, event.y_root + 12)
            except Exception:
                pass
        else:
            self._hide_drag_tip()

    def _on_sections_drag_motion(self, event):
        self._move_drag_tip(event.x_root + 12, event.y_root + 12)
        iid = None
        try:
            iid = self.sections_tree.identify_row(event.y)
        except Exception:
            iid = None
        if not iid:
            return
        try:
            self.sections_tree.selection_set(iid)
        except Exception:
            return

    def _on_sections_drag_drop(self, event):
        from_iid = getattr(self, "_drag_section_iid", None)
        self._drag_section_iid = None
        self._hide_drag_tip()
        if not from_iid:
            return

        try:
            to_iid = self.sections_tree.identify_row(event.y)
        except Exception:
            to_iid = ""
        if not to_iid:
            to_idx = len(self.data.get("sections", []) or []) - 1
        else:
            try:
                to_idx = int(to_iid)
            except Exception:
                return
        try:
            from_idx = int(from_iid)
        except Exception:
            return

        sections = self.data.get("sections", [])
        if not isinstance(sections, list):
            return

        if from_idx == to_idx:
            return

        self._checkpoint_before_change()
        new_idx = self._move_item_in_list(sections, from_idx, to_idx)
        self.refresh_sections()
        try:
            self.sections_tree.selection_set(str(new_idx))
        except Exception:
            pass
        self.on_section_selected(None)
        self._update_undo_redo_buttons()

    def _on_entries_drag_start(self, event):
        self._hide_drag_tip()
        try:
            if self.entries_tree.identify_region(event.x, event.y) != "cell":
                self._drag_entry_iid = None
                return
            iid = self.entries_tree.identify_row(event.y)
            self._drag_entry_iid = iid if iid else None
        except Exception:
            self._drag_entry_iid = None

        if self._drag_entry_iid is not None:
            try:
                sidx = self._selected_section_index()
                if sidx >= 0:
                    sec = self.data.get("sections", [])[sidx]
                    entries = sec.get("entries", []) if isinstance(sec, dict) else []
                    idx = int(self._drag_entry_iid)
                    if 0 <= idx < len(entries):
                        kind = str(sec.get("kind", ""))
                        text = self._entry_summary(kind, entries[idx])
                        text = str(text).replace("\n", " ").strip()
                        if len(text) > 80:
                            text = text[:77] + "..."
                        self._show_drag_tip(text, event.x_root + 12, event.y_root + 12)
            except Exception:
                pass
        else:
            self._hide_drag_tip()

    def _on_entries_drag_motion(self, event):
        self._move_drag_tip(event.x_root + 12, event.y_root + 12)
        iid = None
        try:
            iid = self.entries_tree.identify_row(event.y)
        except Exception:
            iid = None
        if not iid:
            return
        try:
            self.entries_tree.selection_set(iid)
        except Exception:
            return

    def _on_entries_drag_drop(self, event):
        from_iid = getattr(self, "_drag_entry_iid", None)
        self._drag_entry_iid = None
        self._hide_drag_tip()
        if not from_iid:
            return

        sidx = self._selected_section_index()
        if sidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        if not isinstance(sec, dict):
            return
        entries = sec.get("entries", [])
        if not isinstance(entries, list):
            return

        try:
            to_iid = self.entries_tree.identify_row(event.y)
        except Exception:
            to_iid = ""
        if not to_iid:
            to_idx = len(entries) - 1
        else:
            try:
                to_idx = int(to_iid)
            except Exception:
                return
        try:
            from_idx = int(from_iid)
        except Exception:
            return

        if from_idx == to_idx:
            return

        self._checkpoint_before_change()
        new_idx = self._move_item_in_list(entries, from_idx, to_idx)
        self.refresh_entries()
        try:
            self.entries_tree.selection_set(str(new_idx))
        except Exception:
            pass
        self._update_entry_action_buttons()
        self._update_undo_redo_buttons()

    def _show_drag_tip(self, text: str, x: int, y: int) -> None:
        try:
            if self._drag_tip_win is None or not self._drag_tip_win.winfo_exists():
                tip = tk.Toplevel(self.master)
                tip.overrideredirect(True)
                try:
                    tip.attributes("-topmost", True)
                except Exception:
                    pass
                lbl = tk.Label(
                    tip,
                    text=text,
                    bg="#ffffe0",
                    fg="#000000",
                    bd=1,
                    relief="solid",
                    padx=6,
                    pady=3,
                )
                lbl.pack()
                self._drag_tip_win = tip
                self._drag_tip_label = lbl
            else:
                if self._drag_tip_label is not None:
                    self._drag_tip_label.configure(text=text)
            self._drag_tip_win.geometry(f"+{x}+{y}")
        except Exception:
            return

    def _move_drag_tip(self, x: int, y: int) -> None:
        try:
            if self._drag_tip_win is None or not self._drag_tip_win.winfo_exists():
                return
            self._drag_tip_win.geometry(f"+{x}+{y}")
        except Exception:
            return

    def _hide_drag_tip(self) -> None:
        try:
            if self._drag_tip_win is not None and self._drag_tip_win.winfo_exists():
                self._drag_tip_win.destroy()
        except Exception:
            pass
        self._drag_tip_win = None
        self._drag_tip_label = None

    def sync_header_from_ui(self):
        header = self.data.setdefault("header", {})
        for k, var in self.header_vars.items():
            try:
                header[k] = str(var.get() or "").strip()
            except Exception:
                header[k] = ""

    def _update_header_link_field_states(self) -> None:
        def _get_var(k: str) -> str:
            try:
                v = self.header_vars.get(k)
                return "" if v is None else str(v.get() or "")
            except Exception:
                return ""

        def _set_state(k: str, enabled: bool) -> None:
            w = self._header_widgets.get(k)
            if w is None:
                return
            state = "normal" if enabled else "disabled"
            try:
                w.configure(state=state)
                return
            except Exception:
                pass
            try:
                if enabled:
                    w.state(["!disabled"])
                else:
                    w.state(["disabled"])
            except Exception:
                return

        linkedin_kind = _get_var("linkedin_kind").strip().lower()
        github_kind = _get_var("github_kind").strip().lower()

        linkedin_enabled = linkedin_kind != "none"
        github_enabled = github_kind != "none"

        _set_state("linkedin", linkedin_enabled)
        _set_state("linkedin_display", linkedin_enabled)
        _set_state("github", github_enabled)
        _set_state("github_display", github_enabled)

    def _update_header_platform_text_labels(self) -> None:
        def _get_var(k: str) -> str:
            try:
                v = self.header_vars.get(k)
                return "" if v is None else str(v.get() or "")
            except Exception:
                return ""

        def _platform_label(kind: str) -> str:
            k = (kind or "").strip()
            if not k:
                return "Link Label"
            if k.strip().lower() == "none":
                return "Link Label"
            return f"{k} Link Label"

        def _set_label(display_key: str, txt: str) -> None:
            try:
                v = getattr(self, "_header_label_vars", {}).get(display_key)
                if v is not None:
                    v.set(txt)
            except Exception:
                return

        linkedin_kind = _get_var("linkedin_kind").strip().lower()
        github_kind = _get_var("github_kind").strip().lower()

        _set_label("linkedin_display", _platform_label(_get_var("linkedin_kind")))
        _set_label("github_display", _platform_label(_get_var("github_kind")))

    def _push_undo_snapshot(self, snapshot: dict) -> None:
        self.undo_stack.append(deep_copy(snapshot))
        if len(self.undo_stack) > 100:
            self.undo_stack.pop(0)
        self.redo_stack.clear()
        self._update_undo_redo_buttons()

    def _checkpoint_before_change(self) -> None:
        self._commit_header_typing()
        self.sync_header_from_ui()
        self._push_undo_snapshot(self.data)

    def _clear_header_typing_state(self) -> None:
        if self._header_commit_job is not None:
            try:
                self.master.after_cancel(self._header_commit_job)
            except Exception:
                pass
        self._header_commit_job = None
        self._header_typing_in_progress = False

    def _commit_header_typing(self) -> None:
        if self._header_commit_job is not None:
            try:
                self.master.after_cancel(self._header_commit_job)
            except Exception:
                pass
            self._header_commit_job = None

        if self._header_typing_in_progress:
            self.sync_header_from_ui()
            self._header_typing_in_progress = False

    def _on_header_var_changed(self) -> None:
        if self._suspend_undo:
            return

        if not self._header_typing_in_progress:
            self._push_undo_snapshot(self.data)
            self._header_typing_in_progress = True

        if self._header_commit_job is not None:
            try:
                self.master.after_cancel(self._header_commit_job)
            except Exception:
                pass
            self._header_commit_job = None

        self._header_commit_job = self.master.after(600, self._commit_header_typing)

    def _apply_state(self, state: dict) -> None:
        self._clear_header_typing_state()
        self._suspend_undo = True
        try:
            st = deep_copy(state)
            try:
                if isinstance(st, dict) and "spellcheck_ignore_all" in st:
                    self.spellcheck_manager.set_ignore_all(st.get("spellcheck_ignore_all"))
                    st.pop("spellcheck_ignore_all", None)
            except Exception:
                pass

            self.data = st
            for k, var in self.header_vars.items():
                var.set(str(self.data.get("header", {}).get(k, "")))

            try:
                for k in ("linkedin_kind", "github_kind"):
                    v = self.header_vars.get(k)
                    if v is not None:
                        self._header_last_kind[k] = str(v.get() or "")
            except Exception:
                pass

            self._update_header_link_field_states()
            self._update_header_platform_text_labels()
            self.refresh_sections()
            try:
                self.master.update_idletasks()
            except Exception:
                pass
        finally:
            self._suspend_undo = False

    def _update_undo_redo_buttons(self) -> None:
        if hasattr(self, "btn_undo"):
            self.btn_undo.configure(state=("normal" if self.undo_stack else "disabled"))
        if hasattr(self, "btn_redo"):
            self.btn_redo.configure(state=("normal" if self.redo_stack else "disabled"))

    def _update_save_project_buttons(self) -> None:
        enabled = bool(self.current_project_path)

        if hasattr(self, "btn_save_project"):
            self.btn_save_project.configure(state=("normal" if enabled else "disabled"))

        try:
            if self._file_menu_save_index is not None:
                if self._file_menu is not None:
                    self._file_menu.entryconfigure(
                        self._file_menu_save_index,
                        state=("normal" if enabled else "disabled"),
                    )
        except Exception:
            pass

    def undo(self) -> None:
        self._commit_header_typing()
        self.sync_header_from_ui()
        if not self.undo_stack:
            return
        current = deep_copy(self.data)
        prev = self.undo_stack.pop()
        self.redo_stack.append(current)
        self._apply_state(prev)
        self._update_undo_redo_buttons()

    def redo(self) -> None:
        self._commit_header_typing()
        self.sync_header_from_ui()
        if not self.redo_stack:
            return
        current = deep_copy(self.data)
        nxt = self.redo_stack.pop()
        self.undo_stack.append(current)
        self._apply_state(nxt)
        self._update_undo_redo_buttons()

    def _handle_undo_shortcut(self, _event, which: str):
        w = self.master.focus_get()
        if w is not None:
            try:
                if w.winfo_toplevel() is not self.master:
                    return
            except Exception:
                return
            if isinstance(w, tk.Text):
                return

        if which == "undo":
            self.undo()
        else:
            self.redo()
        return "break"

    def _bind_undo_redo_shortcuts(self) -> None:
        self.master.bind_all(
            "<Control-z>", lambda e: self._handle_undo_shortcut(e, "undo"), add="+"
        )
        self.master.bind_all(
            "<Control-y>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )
        self.master.bind_all(
            "<Control-Shift-z>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )
        self.master.bind_all(
            "<Command-z>", lambda e: self._handle_undo_shortcut(e, "undo"), add="+"
        )
        self.master.bind_all(
            "<Command-y>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )
        self.master.bind_all(
            "<Command-Shift-z>", lambda e: self._handle_undo_shortcut(e, "redo"), add="+"
        )

        self.master.bind_all("<Control-KeyPress-s>", self._handle_save_shortcut, add="+")
        self.master.bind_all("<Command-KeyPress-s>", self._handle_save_shortcut, add="+")

    def _handle_save_shortcut(self, event):
        try:
            if str(getattr(event, "keysym", "")).lower() != "s":
                return
        except Exception:
            return

        try:
            state = int(getattr(event, "state", 0) or 0)
        except Exception:
            state = 0

        if (state & 0x4) == 0:
            return

        w = self.master.focus_get()
        if w is not None:
            try:
                if w.winfo_toplevel() is not self.master:
                    return
            except Exception:
                return

        self.save_project()
        return "break"

    def save(self):
        self.sync_header_from_ui()
        save_resume_data(self.data)
        messagebox.showinfo("Saved", f"Saved to {DATA_FILE_NAME}")

    def save_as(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Save JSON",
            defaultextension=".json",
            initialdir=_workspace_path(),
            initialfile=DATA_FILE_NAME,
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                payload = deep_copy(self.data)
                try:
                    payload["spellcheck_ignore_all"] = self.spellcheck_manager.get_ignore_all()
                except Exception:
                    pass
                json.dump(payload, f, indent=2, ensure_ascii=False)
        except Exception as e:
            messagebox.showerror("Save Failed", str(e))
            return
        self.current_project_path = path
        self._update_save_project_buttons()

    def save_project(self):
        if not self.current_project_path:
            self.save_as()
            return
        self._commit_header_typing()
        self.sync_header_from_ui()
        try:
            with open(self.current_project_path, "w", encoding="utf-8") as f:
                payload = deep_copy(self.data)
                try:
                    payload["spellcheck_ignore_all"] = self.spellcheck_manager.get_ignore_all()
                except Exception:
                    pass
                json.dump(payload, f, indent=2, ensure_ascii=False)
        except Exception as e:
            messagebox.showerror("Save Failed", str(e))
            return

    def load(self):
        self._checkpoint_before_change()
        data = load_resume_data()
        self._apply_state(data)
        self._update_undo_redo_buttons()
        messagebox.showinfo("Loaded", f"Loaded from {DATA_FILE_NAME}")

    def load_from_file(self):
        path = filedialog.askopenfilename(
            parent=self,
            title="Load JSON",
            initialdir=_workspace_path(),
            filetypes=[("JSON files", "*.json"), ("All files", "*.*")],
        )
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                raw = json.load(f)
        except Exception as e:
            messagebox.showerror("Load Failed", str(e))
            return

        try:
            if isinstance(raw, dict):
                self.spellcheck_manager.set_ignore_all(raw.get("spellcheck_ignore_all", []))
                raw.pop("spellcheck_ignore_all", None)
        except Exception:
            pass

        try:
            data = self._normalize_loaded_state(raw)
        except Exception as e:
            messagebox.showerror("Load Failed", str(e))
            return

        self._checkpoint_before_change()

        try:
            for iid in self.sections_tree.selection():
                self.sections_tree.selection_remove(iid)
        except Exception:
            pass

        self._apply_state(data)
        self._update_undo_redo_buttons()
        self.refresh_sections()
        self.current_project_path = path
        self._update_save_project_buttons()
        try:
            self.master.update_idletasks()
        except Exception:
            pass

        try:
            total_entries = sum(len(sec.get("entries", []) or []) for sec in (self.data.get("sections", []) or []) if isinstance(sec, dict))
        except Exception:
            total_entries = 0

        h = self.data.get("header", {}) if isinstance(self.data.get("header"), dict) else {}
        name_preview = str(h.get("name", ""))
        email_preview = str(h.get("email", ""))

        try:
            ui_sections = len(self.sections_tree.get_children(""))
        except Exception:
            ui_sections = -1
        try:
            ui_entries = len(self.entries_tree.get_children(""))
        except Exception:
            ui_entries = -1
        try:
            ui_name = self.header_vars.get("name").get() if self.header_vars.get("name") else ""
            ui_email = self.header_vars.get("email").get() if self.header_vars.get("email") else ""
        except Exception:
            ui_name = ""
            ui_email = ""

        messagebox.showinfo(
            "Loaded",
            f"Loaded from {path}\n\nName: {name_preview}\nEmail: {email_preview}\nSections: {len(self.data.get('sections', []) or [])}\nEntries: {total_entries}\n\nUI Name: {ui_name}\nUI Email: {ui_email}\nUI Sections: {ui_sections}\nUI Entries: {ui_entries}",
        )

    def export(self):
        self.sync_header_from_ui()
        export_latex(self.data)
        messagebox.showinfo("Exported", f"Wrote {EXPORT_FILE_NAME}")

    def export_as_tex(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Export LaTeX",
            defaultextension=".tex",
            initialdir=_workspace_path(),
            initialfile=EXPORT_FILE_NAME,
            filetypes=[("TeX files", "*.tex"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            latex_source = generate_latex(self.data)
            with open(path, "w", encoding="utf-8") as f:
                f.write(latex_source)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))
            return

        messagebox.showinfo("Exported", f"Wrote {path}")

    def export_as_docx(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Export Word Document",
            defaultextension=".docx",
            initialdir=_workspace_path(),
            initialfile=EXPORT_DOCX_FILE_NAME,
            filetypes=[("Word Documents", "*.docx"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Inches, Pt, RGBColor
        except Exception:
            messagebox.showerror(
                "Export Failed",
                "Missing dependency: python-docx\n\nInstall it with: pip install python-docx",
            )
            return

        def _sanitize_text(t: Any) -> str:
            return "" if t is None else str(t)

        def _norm_hex_color(c: Any) -> Optional[str]:
            s = ("" if c is None else str(c)).strip()
            if not s:
                return None
            if re.fullmatch(r"#?[0-9a-fA-F]{6}", s):
                return ("#" + s.lstrip("#")).lower()
            return None

        def _hex_to_rgb(c: Any) -> Optional[RGBColor]:
            s = _norm_hex_color(c)
            if not s:
                return None
            try:
                return RGBColor.from_string(s.lstrip("#"))
            except Exception:
                return None

        def _apply_run_bg(run, c: Any) -> None:
            s = _norm_hex_color(c)
            if not s:
                return
            try:
                rpr = run._r.get_or_add_rPr()
                shd = rpr.find(qn("w:shd"))
                if shd is None:
                    shd = OxmlElement("w:shd")
                    rpr.append(shd)
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), s.lstrip("#"))
            except Exception:
                return

        def _apply_paragraph_bottom_rule(p) -> None:
            try:
                ppr = p._p.get_or_add_pPr()
                pbdr = ppr.find(qn("w:pBdr"))
                if pbdr is None:
                    pbdr = OxmlElement("w:pBdr")
                    ppr.append(pbdr)
                bottom = pbdr.find(qn("w:bottom"))
                if bottom is None:
                    bottom = OxmlElement("w:bottom")
                    pbdr.append(bottom)
                bottom.set(qn("w:val"), "single")
                bottom.set(qn("w:sz"), "8")
                bottom.set(qn("w:space"), "1")
                bottom.set(qn("w:color"), "000000")
            except Exception:
                return

        def _clear_table_borders(table) -> None:
            try:
                tbl = table._tbl
                tblPr = tbl.tblPr
                tblBorders = tblPr.first_child_found_in("w:tblBorders")
                if tblBorders is None:
                    tblBorders = OxmlElement("w:tblBorders")
                    tblPr.append(tblBorders)
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    el = tblBorders.find(qn(f"w:{edge}"))
                    if el is None:
                        el = OxmlElement(f"w:{edge}")
                        tblBorders.append(el)
                    el.set(qn("w:val"), "nil")
            except Exception:
                return

        def _map_docx_font(face: Any) -> str:
            f = ("" if face is None else str(face)).strip().lower()
            if not f:
                return "Times New Roman"
            if "courier" in f or "mono" in f:
                return "Courier New"
            if "arial" in f or "helvetica" in f or "sans" in f:
                return "Arial"
            if "times" in f or "serif" in f or "georgia" in f:
                return "Times New Roman"
            return "Times New Roman"

        def segments_to_plain(segments: Any) -> str:
            if not isinstance(segments, list):
                return ""
            parts: List[str] = []
            for seg in segments:
                if isinstance(seg, dict):
                    t = seg.get("text", "")
                    if isinstance(t, str) and t:
                        parts.append(t)
            return "".join(parts)

        def _strip_bullet_prefix(segments: Any) -> Any:
            prefixes = ("- ", "– ", "— ", "• ", "* ")

            if isinstance(segments, str):
                s = segments.lstrip()
                for p in prefixes:
                    if s.startswith(p):
                        return s[len(p) :].lstrip()
                return s

            if not isinstance(segments, list):
                return segments

            out: List[Any] = []
            stripped = False
            for seg in segments:
                if not stripped and isinstance(seg, dict):
                    t = seg.get("text", "")
                    if isinstance(t, str) and t:
                        s = t.lstrip()
                        for p in prefixes:
                            if s.startswith(p):
                                s = s[len(p) :].lstrip()
                                break
                        new_seg = dict(seg)
                        new_seg["text"] = s
                        out.append(new_seg)
                        stripped = True
                        continue
                if not stripped and isinstance(seg, str) and seg:
                    s = seg.lstrip()
                    for p in prefixes:
                        if s.startswith(p):
                            s = s[len(p) :].lstrip()
                            break
                    out.append(s)
                    stripped = True
                    continue
                out.append(seg)
            return out

        def _add_rich_runs(p, segments: Any) -> None:
            if not isinstance(segments, list):
                if isinstance(segments, str) and segments:
                    p.add_run(_sanitize_text(segments))
                return
            for seg in segments:
                if not isinstance(seg, dict):
                    continue
                t = seg.get("text", "")
                if not isinstance(t, str) or t == "":
                    continue
                run = p.add_run(t)
                run.bold = bool(seg.get("b"))
                run.italic = bool(seg.get("i"))
                run.underline = bool(seg.get("u"))
                if seg.get("font"):
                    try:
                        run.font.name = _map_docx_font(seg.get("font"))
                    except Exception:
                        pass
                if seg.get("size"):
                    try:
                        run.font.size = Pt(float(seg.get("size")))
                    except Exception:
                        pass
                rgb = _hex_to_rgb(seg.get("fg"))
                if rgb is not None:
                    try:
                        run.font.color.rgb = rgb
                    except Exception:
                        pass
                _apply_run_bg(run, seg.get("bg"))

        def add_rich_paragraph(doc, segments: Any, *, style: Optional[str] = None, bullet: bool = False):
            p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
            if bullet:
                try:
                    p.style = doc.styles["List Bullet"]
                except Exception:
                    pass
            _add_rich_runs(p, segments)
            return p

        header = self.data.get("header", {}) if isinstance(self.data.get("header"), dict) else {}
        name = str(header.get("name", "")).strip()
        phone = str(header.get("phone", "")).strip()
        email = str(header.get("email", "")).strip()
        linkedin_kind = str(header.get("linkedin_kind", "LinkedIn") or "").strip()
        github_kind = str(header.get("github_kind", "GitHub") or "").strip()
        linkedin = "" if linkedin_kind.lower() == "none" else str(header.get("linkedin", "")).strip()
        github = "" if github_kind.lower() == "none" else str(header.get("github", "")).strip()

        doc = Document()
        try:
            sec0 = doc.sections[0]
            sec0.left_margin = Inches(0.5)
            sec0.right_margin = Inches(0.5)
            sec0.top_margin = Inches(0.5)
            sec0.bottom_margin = Inches(0.5)
        except Exception:
            pass

        try:
            normal = doc.styles["Normal"].font
            normal.name = "Times New Roman"
            normal.size = Pt(10.5)
        except Exception:
            pass

        if name:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(_sanitize_text(name))
            r.bold = True
            r.font.size = Pt(22)
            r.font.name = "Times New Roman"

        contact_parts: List[str] = []
        if phone:
            contact_parts.append(phone)
        if email:
            contact_parts.append(email)
        if linkedin:
            contact_parts.append(linkedin)
        if github:
            contact_parts.append(github)
        if contact_parts:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(" | ".join([_sanitize_text(x) for x in contact_parts]))
            r.font.size = Pt(9.5)
            r.font.name = "Times New Roman"
            try:
                p.paragraph_format.space_after = Pt(10)
            except Exception:
                pass
        else:
            try:
                doc.add_paragraph().paragraph_format.space_after = Pt(10)
            except Exception:
                doc.add_paragraph()

        content_width_in = 7.5
        right_col_in = 2.2
        left_col_in = max(0.5, content_width_in - right_col_in)
        indent_in = 0.15

        def _add_two_col(left_text: Optional[str], right_text: Optional[str], *, left_bold: bool = False, left_italic: bool = False, right_italic: bool = False, left_runs: Any = None) -> None:
            t = doc.add_table(rows=1, cols=2)
            t.autofit = False
            _clear_table_borders(t)
            try:
                t.columns[0].width = Inches(left_col_in)
                t.columns[1].width = Inches(right_col_in)
            except Exception:
                pass

            c0 = t.cell(0, 0)
            c1 = t.cell(0, 1)
            p0 = c0.paragraphs[0]
            p1 = c1.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            try:
                p0.paragraph_format.space_after = Pt(0)
                p1.paragraph_format.space_after = Pt(0)
            except Exception:
                pass

            if left_runs is not None:
                _add_rich_runs(p0, left_runs)
            elif left_text:
                r0 = p0.add_run(_sanitize_text(left_text))
                r0.bold = bool(left_bold)
                r0.italic = bool(left_italic)
                r0.font.name = "Times New Roman"

            if right_text:
                r1 = p1.add_run(_sanitize_text(right_text))
                r1.italic = bool(right_italic)
                r1.font.name = "Times New Roman"

        def _add_section_header(title: str) -> None:
            t = (title or "").strip()
            if not t:
                return
            p = doc.add_paragraph()
            r = p.add_run(_sanitize_text(t))
            r.font.name = "Times New Roman"
            r.font.size = Pt(11)
            _apply_paragraph_bottom_rule(p)
            try:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(3)
            except Exception:
                pass

        for section in self.data.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title", "")).strip()
            kind = str(section.get("kind", "") or "").strip().lower()
            raw_entries = section.get("entries", []) or []
            entries = raw_entries if isinstance(raw_entries, list) else []

            if title:
                _add_section_header(title)

            for e in entries:
                if not isinstance(e, dict):
                    continue

                if kind == "education":
                    school = str(e.get("school", "")).strip()
                    location = str(e.get("location", "")).strip()
                    degree = str(e.get("degree", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    _add_two_col(school, location, left_bold=True)
                    _add_two_col(degree, dates, left_italic=True, right_italic=True)

                    body = e.get("body", [])
                    body_text = segments_to_plain(body)
                    if body_text:
                        p = add_rich_paragraph(doc, _strip_bullet_prefix(body), bullet=True)
                        try:
                            p.paragraph_format.left_indent = Inches(indent_in)
                        except Exception:
                            pass
                    try:
                        doc.add_paragraph().paragraph_format.space_after = Pt(3)
                    except Exception:
                        doc.add_paragraph()

                elif kind in ("experience", "projects"):
                    if kind == "experience":
                        role = str(e.get("role", "")).strip()
                        org = str(e.get("org", "")).strip()
                        location = str(e.get("location", "")).strip()
                        dates = str(e.get("dates", "")).strip()
                        _add_two_col(role, dates, left_bold=True)
                        _add_two_col(org, location, left_italic=True, right_italic=True)
                    else:
                        proj_title = str(e.get("title", "")).strip()
                        stack = str(e.get("stack", "")).strip()
                        dates = str(e.get("dates", "")).strip()
                        left_runs: List[dict] = []
                        if proj_title:
                            left_runs.append({"text": proj_title, "b": True})
                        if stack:
                            if left_runs:
                                left_runs.append({"text": " | "})
                            left_runs.append({"text": stack, "i": True})
                        _add_two_col(None, dates, left_runs=left_runs)

                    bullets = e.get("bullets", [])
                    if isinstance(bullets, list):
                        for b in bullets:
                            p = add_rich_paragraph(doc, _strip_bullet_prefix(b), bullet=True)
                            try:
                                p.paragraph_format.left_indent = Inches(indent_in)
                            except Exception:
                                pass
                    try:
                        doc.add_paragraph().paragraph_format.space_after = Pt(4)
                    except Exception:
                        doc.add_paragraph()

                elif kind == "skills":
                    break

                else:
                    entry_title = str(e.get("title", "")).strip()
                    if entry_title:
                        p = doc.add_paragraph()
                        try:
                            p.paragraph_format.left_indent = Inches(indent_in)
                            p.paragraph_format.space_after = Pt(0)
                        except Exception:
                            pass
                        r = p.add_run(_sanitize_text(entry_title))
                        r.bold = True
                        r.font.name = "Times New Roman"
                    body = e.get("body", [])
                    body_text = segments_to_plain(body)
                    if body_text:
                        p = add_rich_paragraph(doc, body)
                        try:
                            p.paragraph_format.left_indent = Inches(indent_in)
                        except Exception:
                            pass
                    try:
                        doc.add_paragraph().paragraph_format.space_after = Pt(4)
                    except Exception:
                        doc.add_paragraph()

            if kind == "skills":
                p = doc.add_paragraph()
                try:
                    p.paragraph_format.left_indent = Inches(indent_in)
                    p.paragraph_format.space_after = Pt(0)
                except Exception:
                    pass
                first_line = True
                for e in entries:
                    if not isinstance(e, dict):
                        continue
                    label = str(e.get("label", "")).strip()
                    value = e.get("value", [])
                    value_text = segments_to_plain(value)
                    if not (label or value_text):
                        continue
                    if not first_line:
                        try:
                            p.add_run().add_break()
                        except Exception:
                            doc.add_paragraph()
                    first_line = False
                    if label:
                        r = p.add_run(_sanitize_text(label))
                        r.bold = True
                        r.font.name = "Times New Roman"
                        p.add_run(": ")
                    _add_rich_runs(p, value)

        try:
            doc.save(path)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))
            return

        messagebox.showinfo("Exported", f"Wrote {path}")

    def export_as_pdf(self):
        self._commit_header_typing()
        self.sync_header_from_ui()
        path = filedialog.asksaveasfilename(
            parent=self,
            title="Export PDF",
            defaultextension=".pdf",
            initialdir=_workspace_path(),
            initialfile=EXPORT_PDF_FILE_NAME,
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
        )
        if not path:
            return

        try:
            from xml.sax.saxutils import escape

            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfbase.pdfmetrics import stringWidth
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.platypus import (
                HRFlowable,
                Indenter,
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )
        except Exception:
            messagebox.showerror(
                "Export Failed",
                "Missing dependency: reportlab\n\nInstall it with: pip install reportlab",
            )
            return

        def _sanitize_text(t: Any) -> str:
            s = "" if t is None else str(t)
            try:
                return s.encode("cp1252", errors="replace").decode("cp1252")
            except Exception:
                return s

        def _nbsp(s: str) -> str:
            return (s or "").replace(" ", "\u00A0")

        def _map_font_face(face: Any) -> str:
            f = ("" if face is None else str(face)).strip().lower()
            if not f:
                return "Times-Roman"
            if "courier" in f or "mono" in f:
                return "Courier"
            if "arial" in f or "helvetica" in f or "sans" in f:
                return "Helvetica"
            if "times" in f or "serif" in f or "georgia" in f:
                return "Times-Roman"
            return "Times-Roman"

        def _map_font_variant(base: str, *, bold: bool, italic: bool) -> str:
            b = bool(bold)
            i = bool(italic)
            if base.startswith("Times"):
                if b and i:
                    return "Times-BoldItalic"
                if b:
                    return "Times-Bold"
                if i:
                    return "Times-Italic"
                return "Times-Roman"
            if base.startswith("Helvetica"):
                if b and i:
                    return "Helvetica-BoldOblique"
                if b:
                    return "Helvetica-Bold"
                if i:
                    return "Helvetica-Oblique"
                return "Helvetica"
            if base.startswith("Courier"):
                if b and i:
                    return "Courier-BoldOblique"
                if b:
                    return "Courier-Bold"
                if i:
                    return "Courier-Oblique"
                return "Courier"
            return base

        def _normalize_color(c: Any) -> Optional[str]:
            s = ("" if c is None else str(c)).strip()
            if not s:
                return None
            if re.fullmatch(r"#?[0-9a-fA-F]{6}", s):
                return ("#" + s.lstrip("#")).lower()
            return s

        def _to_rl_color(c: Any):
            s = _normalize_color(c)
            if not s:
                return None
            try:
                return colors.toColor(s)
            except Exception:
                try:
                    if isinstance(s, str) and s.startswith("#") and len(s) == 7:
                        return colors.HexColor(s)
                except Exception:
                    pass
            return None

        def _common_bg(segments: Any) -> Optional[str]:
            if not isinstance(segments, list) or not segments:
                return None
            bg: Optional[str] = None
            saw = False
            for seg in segments:
                if not isinstance(seg, dict):
                    continue
                c = _normalize_color(seg.get("bg"))
                if c is None:
                    continue
                if not saw:
                    bg = c
                    saw = True
                elif c != bg:
                    return None
            return bg if saw else None

        def segments_to_markup(segments: Any) -> str:
            if segments is None:
                return ""
            if isinstance(segments, str):
                return escape(_sanitize_text(segments)).replace("\n", "<br/>")
            if not isinstance(segments, list):
                return ""
            out: List[str] = []
            for seg in segments:
                if isinstance(seg, dict):
                    raw = seg.get("text", "")
                    if not isinstance(raw, str) or raw == "":
                        continue
                    t = escape(_sanitize_text(raw)).replace("\n", "<br/>")

                    font_attrs: List[str] = []
                    face = seg.get("font")
                    if face:
                        is_b = bool(seg.get("b"))
                        is_i = bool(seg.get("i"))
                        mapped = _map_font_variant(_map_font_face(face), bold=is_b, italic=is_i)
                        font_attrs.append(f" face=\"{escape(mapped)}\"")
                        seg = dict(seg)
                        seg["b"] = False
                        seg["i"] = False
                    size = seg.get("size")
                    if size is not None:
                        try:
                            s_int = int(size)
                        except Exception:
                            s_int = 0
                        if s_int > 0:
                            font_attrs.append(f" size=\"{s_int}\"")
                    fg = _normalize_color(seg.get("fg"))
                    if fg:
                        font_attrs.append(f" color=\"{escape(fg)}\"")

                    bg = _normalize_color(seg.get("bg"))
                    if bg and _to_rl_color(bg) is not None:
                        font_attrs.append(f" backcolor=\"{escape(bg)}\"")

                    if seg.get("u"):
                        t = f"<u>{t}</u>"
                    if seg.get("i"):
                        t = f"<i>{t}</i>"
                    if seg.get("b"):
                        t = f"<b>{t}</b>"

                    if font_attrs:
                        t = f"<font{''.join(font_attrs)}>{t}</font>"
                    out.append(t)
                elif isinstance(seg, str) and seg:
                    out.append(escape(_sanitize_text(seg)).replace("\n", "<br/>"))
            return "".join(out)

        def segments_to_paragraph(segments: Any, style: ParagraphStyle, *, bullet_text: Optional[str] = None):
            markup = segments_to_markup(segments)
            if not markup:
                return None
            bg = _common_bg(segments)
            bg_color = _to_rl_color(bg) if bg else None
            st = style
            if bg and bg_color is not None:
                markup = f"<para backColor=\"{escape(str(bg))}\" borderPadding=\"1\">{markup}</para>"
            if bullet_text is None:
                return Paragraph(markup, st)
            return Paragraph(markup, st, bulletText=bullet_text)

        def _strip_bullet_prefix(segments: Any) -> Any:
            prefixes = ("- ", "– ", "— ", "• ", "* ")

            if isinstance(segments, str):
                s = segments.lstrip()
                for p in prefixes:
                    if s.startswith(p):
                        return s[len(p) :].lstrip()
                return s

            if not isinstance(segments, list):
                return segments

            out: List[Any] = []
            stripped = False
            for seg in segments:
                if not stripped and isinstance(seg, dict):
                    t = seg.get("text", "")
                    if isinstance(t, str) and t:
                        s = t.lstrip()
                        for p in prefixes:
                            if s.startswith(p):
                                s = s[len(p) :].lstrip()
                                break
                        new_seg = dict(seg)
                        new_seg["text"] = s
                        out.append(new_seg)
                        stripped = True
                        continue
                if not stripped and isinstance(seg, str) and seg:
                    s = seg.lstrip()
                    for p in prefixes:
                        if s.startswith(p):
                            s = s[len(p) :].lstrip()
                            break
                    out.append(s)
                    stripped = True
                    continue
                out.append(seg)
            return out

        def _display_url(url: str, display_override: str) -> str:
            if display_override:
                return display_override.strip()
            return (url or "").replace("https://", "").replace("http://", "").strip()

        styles = getSampleStyleSheet()
        base = ParagraphStyle(
            "ResumeBase",
            parent=styles["Normal"],
            fontName="Times-Roman",
            fontSize=10.5,
            leading=12.2,
            spaceBefore=0,
            spaceAfter=0,
        )
        name_style = ParagraphStyle(
            "ResumeName",
            parent=base,
            fontName="Times-Bold",
            fontSize=22,
            leading=24,
            alignment=1,
            spaceAfter=2,
        )
        contact_style = ParagraphStyle(
            "ResumeContact",
            parent=base,
            fontSize=9.5,
            leading=11,
            alignment=1,
            spaceAfter=10,
        )
        section_style = ParagraphStyle(
            "ResumeSection",
            parent=base,
            fontName="Times-Roman",
            fontSize=11,
            leading=13,
            spaceBefore=6,
            spaceAfter=1,
        )
        entry_left_bold = ParagraphStyle(
            "EntryLeftBold",
            parent=base,
            fontName="Times-Bold",
        )
        entry_left_italic = ParagraphStyle(
            "EntryLeftItalic",
            parent=base,
            fontName="Times-Italic",
            fontSize=9.8,
            leading=11.5,
        )
        entry_right = ParagraphStyle(
            "EntryRight",
            parent=base,
            alignment=2,
        )
        entry_right_italic = ParagraphStyle(
            "EntryRightItalic",
            parent=entry_right,
            fontName="Times-Italic",
            fontSize=9.8,
            leading=11.5,
        )
        bullet_style = ParagraphStyle(
            "ResumeBullet",
            parent=base,
            leftIndent=18,
            bulletIndent=10,
            spaceBefore=1,
            spaceAfter=0,
        )
        bullet_style.bulletFontName = "Helvetica"
        bullet_style.bulletFontSize = 7

        doc = SimpleDocTemplate(
            path,
            pagesize=letter,
            leftMargin=36,
            rightMargin=36,
            topMargin=36,
            bottomMargin=36,
        )

        current_left_indent = 0.0
        two_col_right_inset = 6.0

        def _two_col(left: Paragraph, right: Paragraph):
            right_plain = ""
            try:
                right_plain = right.getPlainText()
            except Exception:
                right_plain = str(getattr(right, "text", "") or "")

            try:
                right_font = str(getattr(right.style, "fontName", base.fontName))
                right_size = float(getattr(right.style, "fontSize", base.fontSize))
            except Exception:
                right_font = base.fontName
                right_size = base.fontSize

            available_width = max(1, doc.width - float(current_left_indent or 0.0) - float(two_col_right_inset or 0.0))
            w = stringWidth(_sanitize_text(right_plain), right_font, right_size) + 6
            right_w = max(available_width * 0.18, min(available_width * 0.42, w))
            left_w = max(1, available_width - right_w)

            t = Table(
                [[left, right]],
                colWidths=[left_w, right_w],
                hAlign="LEFT",
            )
            t.setStyle(
                TableStyle(
                    [
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("ALIGN", (0, 0), (0, 0), "LEFT"),
                        ("ALIGN", (1, 0), (1, 0), "RIGHT"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                        ("RIGHTPADDING", (1, 0), (1, 0), two_col_right_inset),
                        ("TOPPADDING", (0, 0), (-1, -1), 0),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ]
                )
            )
            return t

        def _add_section_header(story: List[Any], title: str) -> None:
            t = (title or "").strip()
            if not t:
                return
            story.append(Paragraph(escape(_sanitize_text(t)), section_style))
            story.append(
                HRFlowable(
                    width="100%",
                    thickness=1,
                    color=colors.black,
                    spaceBefore=0,
                    spaceAfter=3,
                )
            )

        def _add_bullets_to(dest: List[Any], bullets: Any) -> None:
            if not isinstance(bullets, list):
                return
            for b in bullets:
                p = segments_to_paragraph(_strip_bullet_prefix(b), bullet_style, bullet_text="\u2022")
                if p is not None:
                    dest.append(p)

        story: List[Any] = []

        header = self.data.get("header", {}) if isinstance(self.data.get("header"), dict) else {}
        name = str(header.get("name", "")).strip()
        phone = str(header.get("phone", "")).strip()
        email = str(header.get("email", "")).strip()
        linkedin_kind = str(header.get("linkedin_kind", "LinkedIn") or "").strip()
        github_kind = str(header.get("github_kind", "GitHub") or "").strip()
        linkedin = "" if linkedin_kind.lower() == "none" else str(header.get("linkedin", "")).strip()
        linkedin_display = str(header.get("linkedin_display", "")).strip()
        github = "" if github_kind.lower() == "none" else str(header.get("github", "")).strip()
        github_display = str(header.get("github_display", "")).strip()

        if name:
            story.append(Paragraph(escape(_sanitize_text(name)), name_style))

        def _is_omit_kind(k: str) -> bool:
            return (k or "").strip().lower() == "none"

        def _link_markup(href: str, text: str) -> str:
            href_attr = escape(_sanitize_text(href), {'"': "&quot;"})
            txt = escape(_sanitize_text(text))
            return f'<link href="{href_attr}"><u>{txt}</u></link>'

        contact_parts: List[str] = []
        if phone:
            contact_parts.append(escape(_sanitize_text(phone)))
        if email:
            contact_parts.append(_link_markup(f"mailto:{email}", email))
        if not _is_omit_kind(linkedin_kind) and linkedin:
            href = _normalize_href(linkedin)
            contact_parts.append(_link_markup(href, _display_url(href, linkedin_display)))
        if not _is_omit_kind(github_kind) and github:
            href = _normalize_href(github)
            contact_parts.append(_link_markup(href, _display_url(href, github_display)))
        if contact_parts:
            story.append(Paragraph(" | ".join(contact_parts), contact_style))
        else:
            story.append(Spacer(1, 10))

        for section in self.data.get("sections", []) or []:
            if not isinstance(section, dict):
                continue
            title = str(section.get("title", "")).strip()
            kind = str(section.get("kind", "") or "").strip().lower()
            raw_entries = section.get("entries", []) or []
            entries = raw_entries if isinstance(raw_entries, list) else []

            if title:
                _add_section_header(story, title)

            indent_started = False
            if kind in ("education", "experience", "projects", "skills"):
                story.append(Indenter(left=10.8, right=0))
                indent_started = True
                current_left_indent = 10.8

            for e in entries:
                if not isinstance(e, dict):
                    continue

                block: List[Any] = []

                if kind == "education":
                    school = str(e.get("school", "")).strip()
                    location = str(e.get("location", "")).strip()
                    degree = str(e.get("degree", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(school)), entry_left_bold),
                            Paragraph(escape(_sanitize_text(_nbsp(location))), entry_right),
                        )
                    )
                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(degree)), entry_left_italic),
                            Paragraph(escape(_sanitize_text(_nbsp(dates))), entry_right_italic),
                        )
                    )

                    p = segments_to_paragraph(
                        _strip_bullet_prefix(e.get("body", [])),
                        bullet_style,
                        bullet_text="\u2022",
                    )
                    if p is not None:
                        block.append(p)

                    block.append(Spacer(1, 3))

                elif kind == "experience":
                    role = str(e.get("role", "")).strip()
                    org = str(e.get("org", "")).strip()
                    location = str(e.get("location", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(role)), entry_left_bold),
                            Paragraph(escape(_sanitize_text(_nbsp(dates))), entry_right),
                        )
                    )
                    block.append(
                        _two_col(
                            Paragraph(escape(_sanitize_text(org)), entry_left_italic),
                            Paragraph(escape(_sanitize_text(_nbsp(location))), entry_right_italic),
                        )
                    )
                    _add_bullets_to(block, e.get("bullets", []))
                    block.append(Spacer(1, 4))

                elif kind == "projects":
                    proj_title = str(e.get("title", "")).strip()
                    stack = str(e.get("stack", "")).strip()
                    dates = str(e.get("dates", "")).strip()

                    left = f"<b>{escape(_sanitize_text(proj_title))}</b>"
                    if stack:
                        left += f" | <i>{escape(_sanitize_text(stack))}</i>"

                    block.append(
                        _two_col(
                            Paragraph(left, base),
                            Paragraph(escape(_sanitize_text(_nbsp(dates))), entry_right),
                        )
                    )
                    _add_bullets_to(block, e.get("bullets", []))
                    block.append(Spacer(1, 4))

                elif kind == "skills":
                    break

                else:
                    entry_title = str(e.get("title", "")).strip()
                    if entry_title:
                        block.append(Paragraph(escape(_sanitize_text(entry_title)), entry_left_bold))
                    p = segments_to_paragraph(e.get("body", []), base)
                    if p is not None:
                        block.append(p)
                    block.append(Spacer(1, 4))

                if block:
                    story.extend(block)

            if kind == "skills":
                lines: List[str] = []
                for e in entries:
                    if not isinstance(e, dict):
                        continue
                    label = str(e.get("label", "")).strip()
                    value_markup = segments_to_markup(e.get("value", []))
                    label_txt = escape(_sanitize_text(label))
                    if label_txt and value_markup:
                        lines.append(f"<b>{label_txt}</b>: {value_markup}")
                    elif label_txt:
                        lines.append(f"<b>{label_txt}</b>:")
                    elif value_markup:
                        lines.append(value_markup)
                if lines:
                    story.append(Paragraph("<br/>".join(lines), base))

            if indent_started:
                story.append(Indenter(left=-10.8, right=0))
                current_left_indent = 0.0

        try:
            doc.build(story)
        except Exception as e:
            messagebox.showerror("Export Failed", str(e))
            return

        messagebox.showinfo("Exported", f"Wrote {path}")

    def preview(self):
        self.sync_header_from_ui()
        latex_source = generate_latex(self.data)
        dlg = PreviewDialog(self.master, self.data, latex_source)
        dlg.grab_set()

    def delete_all(self):
        ok = messagebox.askyesno(
            "Delete all",
            "Delete all entered data and reset to a clean start?\n\nThis cannot be undone unless you use Undo.",
        )
        if not ok:
            return
        self._checkpoint_before_change()
        self._apply_state(DEFAULT_DATA)
        self._update_undo_redo_buttons()

    def demo_mode(self):
        ok = messagebox.askyesno(
            "Demo mode",
            "Do you want to load demo data?\n\nIf you select Yes, all current data in the program will be overwritten.",
        )
        if not ok:
            return
        self._checkpoint_before_change()
        self._apply_state(DEMO_DATA)
        self._update_undo_redo_buttons()

    def add_section(self):
        title = _ask_string_centered(self, "Add Section", "Section title:")
        if not title:
            return

        kind_result: Dict[str, Optional[str]] = {"value": None}

        dlg = tk.Toplevel(self)
        dlg.title("Add Section")
        try:
            dlg.transient(self.winfo_toplevel())
        except Exception:
            pass

        root = ttk.Frame(dlg)
        root.pack(fill=tk.BOTH, expand=True, padx=12, pady=12)

        ttk.Label(root, text="Section kind:").pack(anchor="w")
        kind_var = tk.StringVar(value="education")
        cb = ttk.Combobox(
            root,
            textvariable=kind_var,
            values=("education", "experience", "projects", "skills", "custom"),
            state="readonly",
        )
        cb.pack(fill=tk.X, pady=(4, 10))

        btns = ttk.Frame(root)
        btns.pack(fill=tk.X)

        def _ok() -> None:
            kind_result["value"] = str(kind_var.get() or "").strip().lower()
            try:
                dlg.destroy()
            except Exception:
                pass

        def _cancel() -> None:
            kind_result["value"] = None
            try:
                dlg.destroy()
            except Exception:
                pass

        ttk.Button(btns, text="Cancel", command=_cancel).pack(side=tk.RIGHT, padx=(6, 0))
        ttk.Button(btns, text="OK", command=_ok).pack(side=tk.RIGHT)

        try:
            dlg.protocol("WM_DELETE_WINDOW", _cancel)
        except Exception:
            pass
        try:
            dlg.bind("<Return>", lambda _e: _ok())
            dlg.bind("<Escape>", lambda _e: _cancel())
        except Exception:
            pass

        try:
            cb.focus_set()
        except Exception:
            pass

        try:
            dlg.grab_set()
        except Exception:
            pass

        try:
            dlg.update_idletasks()
            _center_window(dlg)
        except Exception:
            pass
        self.wait_window(dlg)

        kind = kind_result.get("value")
        if not kind:
            return

        self._checkpoint_before_change()
        sec = {"id": f"sec_{len(self.data.get('sections', [])) + 1}", "title": title.strip(), "kind": kind, "entries": []}
        self.data.setdefault("sections", []).append(sec)
        self.refresh_sections()
        self._update_undo_redo_buttons()

    def rename_section(self):
        idx = self._selected_section_index()
        if idx < 0:
            return
        sec = self.data.get("sections", [])[idx]
        new_title = _ask_string_centered(
            self,
            "Rename Section",
            "New title:",
            initialvalue=str(sec.get("title", "")),
        )
        if not new_title:
            return
        self._checkpoint_before_change()
        sec["title"] = new_title.strip()
        self.refresh_sections()
        self.sections_tree.selection_set(str(idx))
        self.on_section_selected(None)
        self._update_undo_redo_buttons()

    def delete_section(self):
        idx = self._selected_section_index()
        if idx < 0:
            return
        sec = self.data.get("sections", [])[idx]
        ok = messagebox.askyesno("Delete Section", f"Delete section '{sec.get('title','')}'?")
        if not ok:
            return
        self._checkpoint_before_change()
        self.data.get("sections", []).pop(idx)
        self.refresh_sections()
        self._update_undo_redo_buttons()

    def add_entry(self):
        sidx = self._selected_section_index()
        if sidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        kind = sec.get("kind")
        new_entry = {}
        if kind == "education":
            new_entry = {"school": "", "location": "", "degree": "", "dates": ""}
        elif kind == "experience":
            new_entry = {"role": "", "dates": "", "org": "", "location": "", "bullets": [[{"text": ""}]]}
        elif kind == "projects":
            new_entry = {"title": "", "stack": "", "dates": "", "bullets": [[{"text": ""}]]}
        elif kind == "skills":
            new_entry = {"label": "", "value": [{"text": ""}]}
        else:
            new_entry = {"title": "", "body": [{"text": ""}]}

        dlg = EntryEditorDialog(self, kind, new_entry)
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self._checkpoint_before_change()
        sec.setdefault("entries", []).append(dlg.result)
        self.refresh_entries()
        self._update_undo_redo_buttons()

    def edit_entry(self):
        sidx = self._selected_section_index()
        eidx = self._selected_entry_index()
        if sidx < 0 or eidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        kind = sec.get("kind")
        entries = sec.get("entries", [])
        if eidx >= len(entries):
            return

        dlg = EntryEditorDialog(self, kind, entries[eidx])
        self.wait_window(dlg)
        if dlg.result is None:
            return
        self._checkpoint_before_change()
        entries[eidx] = dlg.result
        self.refresh_entries()
        self.entries_tree.selection_set(str(eidx))
        self._update_undo_redo_buttons()

    def delete_entry(self):
        sidx = self._selected_section_index()
        eidx = self._selected_entry_index()
        if sidx < 0 or eidx < 0:
            return
        sec = self.data.get("sections", [])[sidx]
        entries = sec.get("entries", [])
        if eidx >= len(entries):
            return
        ok = messagebox.askyesno("Delete Entry", "Delete selected entry?")
        if not ok:
            return
        self._checkpoint_before_change()
        entries.pop(eidx)
        self.refresh_entries()
        self._update_undo_redo_buttons()


def main():
    root = tk.Tk()
    apply_modern_theme(root)
    app = ResumeApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
